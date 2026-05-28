# -*- coding: utf-8 -*-
"""
自动识别签名位/签字位（用于前端自动勾选角色）。

实现遵循 .cursor/skills/signature-field-identification 的规则要点：
- 同一区域出现「角色 + 日期」 => 高置信度
- 出现「编制/审核/批准」组合 => 高置信度
- 仅单个 token（如“签字/盖章”）=> 低置信度候选
"""
from __future__ import annotations

import os
import re
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Sequence, Tuple

from sign_handlers.config import ROLE_ID_TO_KEYWORD, canonical_sign_role_id, role_keywords
from sign_handlers.docx_revision_text import cell_effective_text, paragraph_effective_text
from sign_handlers.label_match import cell_text_matches_keyword, xlsx_cell_has_leading_role_keyword

_DATE_TOKENS = (
    "日期",
    "签署日期",
    "签字日期",
    "填报日期",
    "年月日",
    "年 月 日",
    "Date",
    "Signed Date",
    "Sign Date",
)
_ACTION_TOKENS = ("签字", "签名", "签章", "盖章", "签 字", "签 名")
_EN_TRIAD_TOKENS = (
    "author",
    "reviewer",
    "approver",
    "prepared by",
    "reviewed by",
    "approved by",
)
_ZH_TRIAD_TOKENS = ("编制人", "编制", "审核人", "审核", "批准人", "批准", "编写", "作者")
_SIGN_SLOT_PLACEHOLDER_RE = re.compile(
    r"(?:_{3,}|\.{3,}|·{3,}|□|■|[（(]\s*[年Y]\s*[）)]\s*[（(]\s*[月M]\s*[）)]\s*[（(]\s*[日D]\s*[）)])"
)
_SIGN_SLOT_LABEL_RE = re.compile(r"(?:编制|编写|作者|审核|复核|批准|签字|签名|日期|Date|Signed)", re.I)
_DOCX_EDGE_START_RATIO = 0.22
_DOCX_EDGE_END_RATIO = 0.78
_STRONG_SIGN_RULES = frozenset(
    {
        "strong_block_rule",
        "triad_rule",
        "triad_text_hint",
        "triad_text_hint_en",
        "role_date_rule",
        "role_placeholder_rule",
        "sign_slot_placeholder_rule",
        "table_header_rule",
        "docx_role_with_date_row",
        "docx_multi_role_row",
        "multi_role_rule",
        "org_seal_rule",
    }
)


def _get_detect_hint_weight() -> float:
    """可配置的 hint 权重（系统设置 SIGN_DETECT_HINT_WEIGHT）。"""
    v = 1.0
    try:
        from runtime_settings.resolve import get_setting

        v = float(get_setting("SIGN_DETECT_HINT_WEIGHT"))
    except Exception:
        v = 1.0
    if not (v == v):  # NaN
        v = 1.0
    return max(0.2, min(2.5, v))


def _scale_hint_conf(base: float, weight: float, *, low: float = 0.0, high: float = 0.99) -> float:
    x = float(base) * float(weight)
    if x < low:
        return low
    if x > high:
        return high
    return x


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())


def _contains_any(text: str, tokens: Sequence[str]) -> bool:
    t = (text or "").strip()
    if not t:
        return False
    low = t.lower()
    for x in tokens:
        if not x:
            continue
        if all(ord(c) < 128 for c in x):
            if x.lower() in low:
                return True
        else:
            if x in t:
                return True
    return False


def _detect_placeholder_forms(text: str) -> List[str]:
    """
    识别"留白形态"：返回命中的形态标签列表（不止枚举几种符号）。
    标签语义：
      - underscore_run / dot_run / box_symbol / ymd_brackets：明确符号占位
      - label_trailing_space：签批标签后留长空白
      - label_dangling_colon：签批标签后只剩冒号（典型"待填写"）
      - symbol_run：任意连续符号段（自适应）
      - bracket_empty：括号内空 / 仅含空白或符号
    """
    t = str(text or "").strip()
    if not t:
        return []
    forms: List[str] = []
    seen = set()

    def _add(tag: str) -> None:
        if tag and tag not in seen:
            seen.add(tag)
            forms.append(tag)

    if re.search(r"_{3,}", t):
        _add("underscore_run")
    if re.search(r"\.{3,}|·{3,}", t):
        _add("dot_run")
    if re.search(r"[□■]", t):
        _add("box_symbol")
    if re.search(
        r"[（(]\s*[年Y]\s*[）)]\s*[（(]\s*[月M]\s*[）)]\s*[（(]\s*[日D]\s*[）)]",
        t,
    ):
        _add("ymd_brackets")

    has_label = bool(_SIGN_SLOT_LABEL_RE.search(t))
    if has_label:
        if re.search(
            r"(编制|编写|作者|审核|复核|批准|签字|签名|日期)\s*[:：]?\s{3,}", t
        ):
            _add("label_trailing_space")
        if re.search(
            r"(?:编制|编写|作者|审核|复核|批准|签字|签名|日期|Date|Signed)\s*[:：]\s*$",
            t,
            re.I,
        ):
            _add("label_dangling_colon")
        for seg in re.findall(r"[^\w\u4e00-\u9fff]{3,}", t):
            if seg.strip():
                _add("symbol_run")
                break
        if re.search(r"[（(]\s*[^\w\u4e00-\u9fff]?\s*[）)]", t):
            _add("bracket_empty")
    return forms


def _has_sign_slot_placeholder(text: str) -> bool:
    return bool(_detect_placeholder_forms(text))


def _docx_block_position_score(
    source_hint: str, total_tables: int, total_paragraphs: int
) -> Optional[float]:
    s = str(source_hint or "").strip()
    m = re.match(r"^table(\d+)\.row\d+$", s, re.IGNORECASE)
    if m:
        idx = int(m.group(1))
        if total_tables <= 1:
            return 0.0
        return max(0.0, min(1.0, float(idx - 1) / float(max(1, total_tables - 1))))
    m = re.match(r"^paragraph(\d+)$", s, re.IGNORECASE)
    if m:
        idx = int(m.group(1))
        if total_paragraphs <= 1:
            return 0.0
        return max(0.0, min(1.0, float(idx - 1) / float(max(1, total_paragraphs - 1))))
    if s.lower() in ("header", "footer", "first_page_header", "first_page_footer"):
        # 页眉页脚视为边缘区域
        return 0.0 if "header" in s.lower() else 1.0
    return None


def _is_docx_edge_position(score: Optional[float]) -> bool:
    if score is None:
        return False
    return score <= _DOCX_EDGE_START_RATIO or score >= _DOCX_EDGE_END_RATIO


def _block_has_sign_context(block: Dict[str, Any]) -> bool:
    if not isinstance(block, dict):
        return False
    fields = block.get("fields") if isinstance(block.get("fields"), list) else []
    has_role = any(
        isinstance(f, dict) and f.get("type") == "role_id" and str(f.get("name") or "").strip()
        for f in fields
    )
    if not has_role:
        return False
    has_date = any(isinstance(f, dict) and f.get("type") == "date" for f in fields)
    has_action = any(isinstance(f, dict) and f.get("type") == "action" for f in fields)
    txt = (str(block.get("label_preview") or "") + " " + str(block.get("source_hint") or "")).strip()
    has_placeholder = _has_sign_slot_placeholder(txt)
    matched = {str(x or "").strip() for x in (block.get("matched_rules") or [])}
    return bool(has_date or has_action or has_placeholder or (matched & _STRONG_SIGN_RULES))


def _filter_docx_blocks_by_edge_priority(
    blocks: List[Dict[str, Any]], total_tables: int, total_paragraphs: int
) -> Tuple[List[Dict[str, Any]], bool]:
    strong_blocks: List[Tuple[int, Dict[str, Any], bool]] = []
    for b in blocks:
        if not _block_has_sign_context(b):
            continue
        score = _docx_block_position_score(
            str(b.get("source_hint") or ""), total_tables, total_paragraphs
        )
        strong_blocks.append((len(strong_blocks), b, _is_docx_edge_position(score)))
    if not strong_blocks:
        return blocks, False

    edge_cnt = sum(1 for _, _, is_edge in strong_blocks if is_edge)
    body_cnt = len(strong_blocks) - edge_cnt

    # 自适应策略：仅当“边缘签批证据明显占优”才过滤正文，避免硬性写死。
    edge_dominant = edge_cnt >= 2 and edge_cnt >= body_cnt + 1
    if not edge_dominant:
        return blocks, False

    kept = []
    for b in blocks:
        score = _docx_block_position_score(
            str(b.get("source_hint") or ""), total_tables, total_paragraphs
        )
        if _is_docx_edge_position(score) and _block_has_sign_context(b):
            kept.append(b)
    if kept:
        return kept, True
    return blocks, False


def _filter_role_evidence_to_edge(
    role_evidence: Dict[str, List[Dict[str, Any]]], total_tables: int, total_paragraphs: int
) -> Dict[str, List[Dict[str, Any]]]:
    out: Dict[str, List[Dict[str, Any]]] = {}
    for rid, arr in (role_evidence or {}).items():
        if rid not in ROLE_ID_TO_KEYWORD or not isinstance(arr, list):
            continue
        kept = []
        for ev in arr:
            if not isinstance(ev, dict):
                continue
            score = _docx_block_position_score(
                str(ev.get("source_hint") or ""), total_tables, total_paragraphs
            )
            if _is_docx_edge_position(score):
                kept.append(ev)
        if kept:
            out[rid] = kept[:6]
    return out


def _match_role_in_short_line(line: str) -> List[str]:
    """单行内角色识别：整行=标签 或 行首=标签+冒号/占位；长关键词优先，避免重复。"""
    out: List[str] = []
    seen: set = set()
    pairs: List[Tuple[str, str]] = []
    for rid in ROLE_ID_TO_KEYWORD:
        for kw in role_keywords(rid):
            pairs.append((rid, kw))
    pairs.sort(key=lambda p: len(p[1]), reverse=True)
    for rid, kw in pairs:
        if cell_text_matches_keyword(line, kw) or xlsx_cell_has_leading_role_keyword(line, kw):
            rid2 = canonical_sign_role_id(rid, kw)
            if rid2 in seen:
                continue
            seen.add(rid2)
            out.append(rid2)
    return out


def _paragraph_role_ids(text: str) -> List[str]:
    """段落/短文本中的角色标签（避免在长句中子串误匹配）。"""
    t = _norm(text or "")
    if not t:
        return []
    out: List[str] = []
    seen: set = set()
    if len(t) <= 96:
        for rid in _match_role_in_short_line(t):
            if rid not in seen:
                seen.add(rid)
                out.append(rid)
        if out:
            return out
    for line in re.split(r"[\r\n]+", t):
        line = _norm(line)
        if not line or len(line) > 120:
            continue
        for rid in _match_role_in_short_line(line):
            if rid not in seen:
                seen.add(rid)
                out.append(rid)
    return out


def _role_ids_for_text_label(label: str) -> List[str]:
    """把命中的标签映射成 role_id（按 ROLE_ID_TO_KEYWORD 反查；同时支持「标签:任意内容」形式）。"""
    s = _norm(label or "")
    if not s:
        return []
    return _match_role_in_short_line(s)


@dataclass
class DetectedBlock:
    block_id: str
    confidence: float
    matched_rules: List[str]
    fields: List[Dict[str, str]]
    source_hint: str
    label_preview: str = ""
    placeholder_forms: List[str] = None  # type: ignore[assignment]

    def to_dict(self) -> dict:
        d = {
            "block_id": self.block_id,
            "confidence": self.confidence,
            "matched_rules": self.matched_rules,
            "fields": self.fields,
            "source_hint": self.source_hint,
        }
        if self.label_preview:
            d["label_preview"] = self.label_preview
        if self.placeholder_forms:
            d["placeholder_forms"] = list(self.placeholder_forms)
        return d


def _append_role_evidence(
    role_evidence: Dict[str, List[Dict[str, Any]]],
    role_ids: Sequence[str],
    confidence: float,
    source_hint: str,
    matched_rules: Sequence[str],
    label_preview: str = "",
) -> None:
    for rid in sorted(set(role_ids or [])):
        if rid not in ROLE_ID_TO_KEYWORD:
            continue
        arr = role_evidence.setdefault(rid, [])
        arr.append(
            {
                "confidence": float(confidence),
                "source_hint": str(source_hint or ""),
                "matched_rules": list(matched_rules or []),
                "label_preview": str(label_preview or "")[:120],
            }
        )


def _compact_role_evidence(role_evidence: Dict[str, List[Dict[str, Any]]], top_n: int = 3) -> Dict[str, List[Dict[str, Any]]]:
    out: Dict[str, List[Dict[str, Any]]] = {}
    for rid, arr in (role_evidence or {}).items():
        if rid not in ROLE_ID_TO_KEYWORD or not isinstance(arr, list):
            continue
        seen = set()
        clean = []
        arr_sorted = sorted(arr, key=lambda x: float(x.get("confidence") or 0.0), reverse=True)
        for e in arr_sorted:
            if not isinstance(e, dict):
                continue
            key = (
                str(e.get("source_hint") or ""),
                tuple(str(x) for x in (e.get("matched_rules") or [])),
                str(e.get("label_preview") or ""),
            )
            if key in seen:
                continue
            seen.add(key)
            clean.append(e)
            if len(clean) >= max(1, int(top_n)):
                break
        if clean:
            out[rid] = clean
    return out


def _score_block(
    role_ids: List[str], joined_text: str, has_date: bool, has_action: bool, has_placeholder: bool = False
) -> Tuple[float, List[str]]:
    """role_ids 为英文 id（author/reviewer/…）；joined_text 为单元格拼接后的原文，用于中文关键词规则。"""
    matched: List[str] = []
    rset = set(role_ids)
    jt = joined_text or ""
    # 仅「执行+审核」且无日期/动作，通常是用例表列头，不应作为签字角色命中。
    if rset and rset.issubset({"executor", "reviewer"}) and not has_date and not has_action and not has_placeholder:
        return 0.0, ["weak_executor_reviewer_only"]
    if rset and (has_date or has_placeholder):
        matched.append("strong_block_rule")
    if rset and has_placeholder:
        matched.append("sign_slot_placeholder_rule")
    if {"author", "reviewer", "approver"}.issubset(rset):
        matched.append("triad_rule")
    elif _contains_any(jt, _ZH_TRIAD_TOKENS) and len(rset) >= 2:
        matched.append("triad_text_hint")
    elif _contains_any(jt.lower(), _EN_TRIAD_TOKENS) and len(rset) >= 2:
        matched.append("triad_text_hint_en")
    elif len(rset) >= 2 and rset.intersection({"author", "reviewer", "approver"}) and (
        has_date or has_action or "author" in rset or "approver" in rset
    ):
        matched.append("multi_role_rule")
    elif len(rset) >= 2 and rset.intersection({"author", "reviewer", "approver"}):
        matched.append("multi_role_weak_context")
    if _contains_any(jt, ("企业负责人", "法定代表人", "法人代表")) and (has_date or has_action):
        matched.append("org_seal_rule")
    if has_action and has_date and rset:
        matched.append("table_header_rule")
    if len(rset) >= 2 and has_action:
        matched.append("multi_role_action_rule")
    if len(rset) >= 1 and has_date and "strong_block_rule" not in matched:
        matched.append("role_date_rule")
    if len(rset) >= 1 and has_placeholder and "strong_block_rule" not in matched:
        matched.append("role_placeholder_rule")

    # 置信度：强规则优先
    if "triad_rule" in matched or "strong_block_rule" in matched:
        return 0.93, matched
    if "triad_text_hint" in matched or "triad_text_hint_en" in matched:
        return 0.88, matched
    if "multi_role_rule" in matched:
        return 0.8, matched
    if "multi_role_weak_context" in matched:
        return 0.62, matched
    if "org_seal_rule" in matched:
        return 0.86, matched
    if "table_header_rule" in matched:
        return 0.82, matched
    if "multi_role_action_rule" in matched:
        return 0.78, matched
    if "role_date_rule" in matched:
        return 0.74, matched
    if "role_placeholder_rule" in matched:
        return 0.72, matched
    if rset and (has_date or has_action or has_placeholder):
        return 0.72, matched or ["medium_block_rule"]
    if has_action and rset:
        return 0.55, matched or ["action_with_role"]
    if has_action:
        return 0.45, ["single_token_fallback"]
    return 0.0, []


def _cell_role_ids_multiline(cell_text: str) -> List[str]:
    """同一单元格多行签批（Author/Reviewer/Approver 纵向排列）逐行识别。"""
    raw = str(cell_text or "")
    if not raw.strip():
        return []
    parts = [raw]
    if re.search(r"[\r\n]", raw):
        parts = [p for p in re.split(r"[\r\n]+", raw) if p and p.strip()]
    out: List[str] = []
    seen: set = set()
    for part in parts:
        s = _norm(part)
        if not s:
            continue
        for rid in _table_cell_role_ids(s):
            if rid not in seen:
                seen.add(rid)
                out.append(rid)
        for rid in _role_ids_for_text_label(s):
            if rid not in seen:
                seen.add(rid)
                out.append(rid)
    return out


def _table_cell_role_ids(cell_text: str) -> List[str]:
    """Word/Excel 单元格内匹配到的 role_id（整格或格首同义词）；长关键词优先。"""
    s = _norm(cell_text or "")
    if not s or len(s) > 96:
        return []
    pairs: List[Tuple[str, str]] = []
    for rid in ROLE_ID_TO_KEYWORD:
        for kw in role_keywords(rid):
            pairs.append((rid, kw))
    pairs.sort(key=lambda p: len(p[1]), reverse=True)
    out: List[str] = []
    seen: set = set()
    for rid, kw in pairs:
        if xlsx_cell_has_leading_role_keyword(s, kw):
            rid2 = canonical_sign_role_id(rid, kw)
            if rid2 in seen:
                continue
            seen.add(rid2)
            out.append(rid2)
    return out


def _xlsx_effective_scan_dims(ws, cap_rows: int, cap_cols: int) -> Tuple[int, int]:
    """工作表真实占用范围（多页/长表时 max_row 可能偏小，用 calculate_dimension 补全）。"""
    eff_r = int(ws.max_row or 1)
    eff_c = int(ws.max_column or 1)
    try:
        from openpyxl.utils import range_boundaries

        dim = ws.calculate_dimension()
        if dim:
            _min_c, _min_r, max_c, max_r = range_boundaries(dim)
            eff_r = max(eff_r, int(max_r))
            eff_c = max(eff_c, int(max_c))
    except Exception:
        pass
    return min(max(eff_r, 1), cap_rows), min(max(eff_c, 1), cap_cols)


def _harvest_roles_xlsx_cell(ws, max_scan_rows: int, max_scan_cols: int) -> Dict[str, float]:
    """逐格扫描：整格或格首为角色标签即记入（补漏单行检测不到的表头）。"""
    found: Dict[str, float] = {}
    for r in range(1, max_scan_rows + 1):
        for c in range(1, max_scan_cols + 1):
            v = ws.cell(row=r, column=c).value
            if v is None:
                continue
            s = _norm(str(v))
            if not s or len(s) > 96:
                continue
            for rid in _cell_role_ids_multiline(s):
                found[rid] = max(found.get(rid, 0.0), 0.58)
    return found


def _iter_docx_tables(doc) -> List:
    """全量表（签字落位等场景）；识别角色请用 _iter_docx_tables_for_detect。"""
    return _iter_docx_tables_for_detect(doc, max_total=10000)


def _iter_docx_tables_for_detect(doc, max_body_tables: int = 32, max_total: int = 96) -> List:
    """
    识别专用：SRS 等长文档会在每节页眉/页脚重复签批表，全量遍历会极慢甚至像卡死。
    仅取正文前/后若干表 + 首尾节的页眉页脚表。
    """
    tables: List = []
    body = list(getattr(doc, "tables", None) or [])
    if len(body) > max_body_tables:
        head_n = min(8, max_body_tables // 2)
        tail_n = max_body_tables - head_n
        body = body[:head_n] + body[-tail_n:]
    tables.extend(body)
    try:
        sections = list(doc.sections)
        pick = []
        if sections:
            pick.append(sections[0])
            if len(sections) > 1:
                pick.append(sections[-1])
        for section in pick:
            for hf in (
                section.footer,
                section.header,
                section.first_page_footer,
                section.first_page_header,
            ):
                if hf is None:
                    continue
                tables.extend(list(getattr(hf, "tables", None) or []))
    except Exception:
        pass
    if len(tables) > max_total:
        return tables[:max_total]
    return tables


def _docx_has_core_triad(role_ids_found: Dict[str, float]) -> bool:
    return {"author", "reviewer", "approver"}.issubset(set(role_ids_found.keys()))


_ROLE_KW_FLAT_CACHE: Optional[Tuple[str, ...]] = None


def _all_role_keywords_flat() -> Tuple[str, ...]:
    global _ROLE_KW_FLAT_CACHE
    if _ROLE_KW_FLAT_CACHE is None:
        seen: set = set()
        flat: List[str] = []
        for rid in ROLE_ID_TO_KEYWORD:
            for kw in role_keywords(rid):
                k = str(kw or "").strip()
                if not k or k in seen:
                    continue
                seen.add(k)
                flat.append(k)
        flat.sort(key=len, reverse=True)
        _ROLE_KW_FLAT_CACHE = tuple(flat)
    return _ROLE_KW_FLAT_CACHE


def _paragraph_might_have_role(text: str) -> bool:
    t = text or ""
    if not t:
        return False
    if _contains_any(t, _DATE_TOKENS) or _contains_any(t, _ACTION_TOKENS):
        return True
    for kw in _all_role_keywords_flat():
        if all(ord(c) < 128 for c in kw):
            if kw.lower() in t.lower():
                return True
        elif kw in t:
            return True
    return False


def _apply_docx_table_triad_boost(doc, role_ids_found: Dict[str, float]) -> None:
    """同一表内凑齐 author+reviewer+approver（或其中两个）时抬高置信度。"""
    core = {"author", "reviewer", "approver"}
    for table in _iter_docx_tables_for_detect(doc):
        ids: set = set()
        for row in table.rows:
            for cell in row.cells:
                ids.update(_cell_role_ids_multiline(_norm(cell_effective_text(cell) or "")))
        hit = ids & core
        if not hit:
            continue
        conf = 0.92 if core.issubset(ids) else 0.8
        for rid in ids:
            if rid in ROLE_ID_TO_KEYWORD:
                role_ids_found[rid] = max(role_ids_found.get(rid, 0.0), conf)


def _xlsx_row_sign_signals(
    ws,
    r: int,
    max_c: int,
    scan_last_row: int,
    _dtoks: Sequence[str],
    _atoks: Sequence[str],
) -> Tuple[List[str], bool, bool, List[str]]:
    """扫描第 r 行（并合并 r+1 行常见日期/签字格）得到角色、日期/动作标记与拼接文本。"""
    role_ids: List[str] = []
    has_date = False
    has_action = False
    cell_texts: List[str] = []
    for c in range(1, max_c + 1):
        v = ws.cell(row=r, column=c).value
        if v is None:
            continue
        s = _norm(str(v))
        if not s or len(s) < 2:
            continue
        cell_texts.append(s)
        role_ids.extend(_cell_role_ids_multiline(s))
        if any(cell_text_matches_keyword(s, d) for d in _DATE_TOKENS) or any(
            xlsx_cell_has_leading_role_keyword(s, d) for d in _dtoks
        ):
            has_date = True
        if any(cell_text_matches_keyword(s, a) for a in _ACTION_TOKENS) or any(
            xlsx_cell_has_leading_role_keyword(s, a) for a in _atoks
        ):
            has_action = True

    if r < scan_last_row:
        for c in range(1, max_c + 1):
            v2 = ws.cell(row=r + 1, column=c).value
            if v2 is None:
                continue
            s2 = _norm(str(v2))
            if not s2:
                continue
            cell_texts.append(s2)
            if any(cell_text_matches_keyword(s2, d) for d in _DATE_TOKENS) or any(
                xlsx_cell_has_leading_role_keyword(s2, d) for d in _dtoks
            ):
                has_date = True
            if any(cell_text_matches_keyword(s2, a) for a in _ACTION_TOKENS) or any(
                xlsx_cell_has_leading_role_keyword(s2, a) for a in _atoks
            ):
                has_action = True

    return role_ids, has_date, has_action, cell_texts


def _append_xlsx_block(
    blocks: List[DetectedBlock],
    bi_holder: List[int],
    role_ids: List[str],
    joined: str,
    has_date: bool,
    has_action: bool,
    conf: float,
    matched: List[str],
    sheet_title: str,
    row_index: int,
    role_ids_found: Dict[str, float],
    role_evidence: Dict[str, List[Dict[str, Any]]],
) -> None:
    fields: List[Dict[str, str]] = []
    for rid in sorted(set(role_ids)):
        fields.append({"name": rid, "type": "role_id"})
    if has_date:
        fields.append({"name": "日期", "type": "date"})
    if has_action:
        fields.append({"name": "签字/盖章", "type": "action"})
    bi = bi_holder[0]
    bi_holder[0] = bi + 1
    b = DetectedBlock(
        block_id=f"xlsx_{bi:03d}",
        confidence=conf,
        matched_rules=matched,
        fields=fields,
        source_hint=f"{sheet_title}!row{row_index}",
        label_preview=(joined or "")[:200],
        placeholder_forms=_detect_placeholder_forms(joined),
    )
    blocks.append(b)
    for rid in set(role_ids):
        role_ids_found[rid] = max(role_ids_found.get(rid, 0.0), conf)
    _append_role_evidence(
        role_evidence,
        role_ids,
        conf,
        f"{sheet_title}!row{row_index}",
        matched,
        joined,
    )


def detect_xlsx(path: str, max_scan_rows: int = 8000, max_scan_cols: int = 128) -> dict:
    from openpyxl import load_workbook

    # data_only=True 读取公式格缓存结果，避免格内为「=」开头公式时识别不到显示文字
    wb = load_workbook(path, data_only=True)
    blocks: List[DetectedBlock] = []
    role_ids_found: Dict[str, float] = {}
    role_evidence: Dict[str, List[Dict[str, Any]]] = {}
    bi_holder = [1]

    for ws in wb.worksheets:
        eff_rows, eff_cols = _xlsx_effective_scan_dims(ws, max_scan_rows, max_scan_cols)

        # 逐格补漏（与有效扫描范围一致，避免只扫到「第一页」）
        for rid, conf in _harvest_roles_xlsx_cell(ws, eff_rows, eff_cols).items():
            role_ids_found[rid] = max(role_ids_found.get(rid, 0.0), conf)
            _append_role_evidence(
                role_evidence,
                [rid],
                conf,
                f"{ws.title or '?'}!cell_scan",
                ["xlsx_cell_scan_role_hint"],
                "",
            )

        _dtoks = sorted(_DATE_TOKENS, key=len, reverse=True)
        _atoks = sorted(_ACTION_TOKENS, key=len, reverse=True)
        rows_with_block = set()

        for r in range(1, eff_rows + 1):
            role_ids, has_date, has_action, cell_texts = _xlsx_row_sign_signals(
                ws, r, eff_cols, eff_rows, _dtoks, _atoks
            )
            joined = " | ".join(cell_texts)
            has_placeholder = _has_sign_slot_placeholder(joined)
            if not role_ids and not has_action:
                continue

            conf, matched = _score_block(role_ids, joined, has_date, has_action, has_placeholder)
            if conf <= 0:
                continue
            if conf < 0.5:
                continue

            _append_xlsx_block(
                blocks,
                bi_holder,
                role_ids,
                joined,
                has_date,
                has_action,
                conf,
                matched,
                ws.title or "?",
                r,
                role_ids_found,
                role_evidence,
            )
            rows_with_block.add(r)

        # 补漏：后面几页/另一块签字区常因无日期列等导致分数<0.5 被整行丢弃，这里按行再收一遍
        for r in range(1, eff_rows + 1):
            if r in rows_with_block:
                continue
            role_ids, has_date, has_action, cell_texts = _xlsx_row_sign_signals(
                ws, r, eff_cols, eff_rows, _dtoks, _atoks
            )
            if not role_ids:
                continue
            joined = " | ".join(cell_texts)
            has_placeholder = _has_sign_slot_placeholder(joined)
            conf, matched = _score_block(role_ids, joined, has_date, has_action, has_placeholder)
            if conf < 0.5:
                rset = set(role_ids)
                if (
                    (len(rset) >= 2 and not rset.issubset({"executor", "reviewer"}))
                    or has_date
                    or has_action
                    or has_placeholder
                ):
                    conf = 0.56
                    matched = list(matched or []) + ["xlsx_multi_page_supplement"]
                elif len(set(role_ids)) == 1 and has_date:
                    conf = 0.56
                    matched = list(matched or []) + ["xlsx_single_role_date_supplement"]
                else:
                    continue
            if conf <= 0:
                continue

            _append_xlsx_block(
                blocks,
                bi_holder,
                role_ids,
                joined,
                has_date,
                has_action,
                conf,
                matched,
                ws.title or "?",
                r,
                role_ids_found,
                role_evidence,
            )

    for b in blocks:
        for f in b.fields:
            if f.get("type") == "role_id" and f.get("name") in ROLE_ID_TO_KEYWORD:
                rid = f["name"]
                role_ids_found[rid] = max(role_ids_found.get(rid, 0.0), float(b.confidence) * 0.95)

    role_ev = _compact_role_evidence(role_evidence)
    placeholder_forms_seen = sorted({
        f
        for b in blocks
        for f in (b.placeholder_forms or [])
        if f
    })
    placeholder_blocks_count = sum(1 for b in blocks if b.placeholder_forms)
    debug_summary = {
        "kind": "xlsx",
        "total_blocks": len(blocks),
        "placeholder_forms_seen": placeholder_forms_seen,
        "placeholder_blocks_count": placeholder_blocks_count,
        "role_count": len(role_ids_found),
        "rules_matched": sorted(
            {
                str(m)
                for b in blocks
                for m in (b.matched_rules or [])
                if str(m).strip()
            }
        ),
    }
    return {
        "ok": True,
        "kind": "xlsx",
        "roles": [{"id": rid, "confidence": role_ids_found[rid]} for rid in sorted(role_ids_found)],
        "blocks": [b.to_dict() for b in blocks],
        "role_evidence": role_ev,
        "debug_summary": debug_summary,
    }


def _harvest_roles_docx_tables(doc) -> Dict[str, float]:
    found: Dict[str, float] = {}
    for table in _iter_docx_tables_for_detect(doc):
        for row in table.rows:
            for cell in row.cells:
                s = _norm(cell_effective_text(cell) or "")
                if not s or len(s) > 120:
                    continue
                for rid in _cell_role_ids_multiline(s):
                    found[rid] = max(found.get(rid, 0.0), 0.58)
    return found


def _harvest_roles_docx_headers_footers(doc) -> Dict[str, float]:
    """页眉/页脚中的签字表（需求规范等常把签批栏放在页脚）。"""
    found: Dict[str, float] = {}
    try:
        for section in doc.sections:
            for hf in (section.header, section.footer, section.first_page_header, section.first_page_footer):
                if hf is None:
                    continue
                for p in hf.paragraphs:
                    t = _norm(paragraph_effective_text(p) or "")
                    if not t:
                        continue
                    for rid in _paragraph_role_ids(t):
                        found[rid] = max(found.get(rid, 0.0), 0.55)
                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            s = _norm(cell_effective_text(cell) or "")
                            if not s or len(s) > 120:
                                continue
                            for rid in _cell_role_ids_multiline(s):
                                found[rid] = max(found.get(rid, 0.0), 0.58)
    except Exception:
        pass
    return found


def detect_docx(path: str, max_paragraphs: int = 1200, *, light: bool = False) -> dict:
    from docx import Document

    from sign_handlers.docx_revision_text import docx_has_track_changes

    has_rev = docx_has_track_changes(path)
    doc = Document(path)
    if light:
        max_paragraphs = min(max_paragraphs, 300)
    blocks: List[DetectedBlock] = []
    role_ids_found: Dict[str, float] = {}
    role_evidence: Dict[str, List[Dict[str, Any]]] = {}
    bi = 1

    for rid, conf in _harvest_roles_docx_tables(doc).items():
        role_ids_found[rid] = max(role_ids_found.get(rid, 0.0), conf)
        _append_role_evidence(
            role_evidence,
            [rid],
            conf,
            "docx_table_scan",
            ["docx_table_scan_role_hint"],
            "",
        )
    for rid, conf in _harvest_roles_docx_headers_footers(doc).items():
        role_ids_found[rid] = max(role_ids_found.get(rid, 0.0), conf)
        _append_role_evidence(
            role_evidence,
            [rid],
            conf,
            "docx_header_footer_scan",
            ["docx_header_footer_scan_role_hint"],
            "",
        )
    _apply_docx_table_triad_boost(doc, role_ids_found)

    tables_for_detect = _iter_docx_tables_for_detect(
        doc, max_body_tables=16 if light else 32, max_total=48 if light else 96
    )

    # 1) 表格行：强场景（含页眉页脚表；签批栏常「角色行 + 下一行 Date」）
    for ti, table in enumerate(tables_for_detect):
        rows_list = list(table.rows)
        for ri, row in enumerate(rows_list):
            texts = [_norm(cell_effective_text(c) or "") for c in row.cells]
            joined = " | ".join(t for t in texts if t)
            if ri + 1 < len(rows_list):
                texts_next = [_norm(cell_effective_text(c) or "") for c in rows_list[ri + 1].cells]
                joined_next = " | ".join(t for t in texts_next if t)
                if joined_next:
                    joined = (joined + " | " + joined_next) if joined else joined_next
            if not joined:
                continue
            role_ids: List[str] = []
            for t in texts:
                if not t:
                    continue
                role_ids.extend(_cell_role_ids_multiline(t))
            has_date = _contains_any(joined, _DATE_TOKENS)
            has_action = _contains_any(joined, _ACTION_TOKENS)
            has_placeholder = _has_sign_slot_placeholder(joined)

            conf, matched = _score_block(role_ids, joined, has_date, has_action, has_placeholder)
            if conf < 0.5:
                rset = set(role_ids)
                if len(rset) >= 1 and (has_date or has_placeholder):
                    conf = 0.72
                    matched = list(matched or []) + ["docx_role_with_date_row"]
                elif len(rset) >= 2 and not rset.issubset({"executor", "reviewer"}):
                    conf = 0.76
                    matched = list(matched or []) + ["docx_multi_role_row"]
                else:
                    continue
            fields: List[Dict[str, str]] = []
            for rid in sorted(set(role_ids)):
                fields.append({"name": rid, "type": "role_id"})
            if has_date:
                fields.append({"name": "日期", "type": "date"})
            if has_action:
                fields.append({"name": "签字/盖章", "type": "action"})
            b = DetectedBlock(
                block_id=f"docx_t{ti+1}_r{ri+1}_{bi:03d}",
                confidence=conf,
                matched_rules=matched,
                fields=fields,
                source_hint=f"table{ti+1}.row{ri+1}",
                label_preview=(joined or "")[:200],
                placeholder_forms=_detect_placeholder_forms(joined),
            )
            blocks.append(b)
            bi += 1
            for rid in set(role_ids):
                role_ids_found[rid] = max(role_ids_found.get(rid, 0.0), conf)
            _append_role_evidence(
                role_evidence,
                role_ids,
                conf,
                f"table{ti+1}.row{ri+1}",
                matched,
                joined,
            )
        if light and _docx_has_core_triad(role_ids_found) and ti >= 4:
            break
        if not light and _docx_has_core_triad(role_ids_found) and ti >= 12:
            break

    # 2) 段落：角色+日期同段（长文档减少扫描上限）
    para_cap = max_paragraphs
    try:
        if len(doc.paragraphs) > 2500:
            para_cap = min(para_cap, 400 if light else 800)
    except Exception:
        pass
    if light and _docx_has_core_triad(role_ids_found):
        para_cap = 0
    for pi, p in enumerate(doc.paragraphs[:para_cap]):
        t = _norm(paragraph_effective_text(p) or "")
        if not t:
            continue
        if not _paragraph_might_have_role(t):
            continue
        role_ids = _paragraph_role_ids(t)
        has_date = _contains_any(t, _DATE_TOKENS)
        has_action = _contains_any(t, _ACTION_TOKENS)
        has_placeholder = _has_sign_slot_placeholder(t)
        conf, matched = _score_block(role_ids, t, has_date, has_action, has_placeholder)
        if conf < 0.68:
            continue
        fields: List[Dict[str, str]] = [{"name": rid, "type": "role_id"} for rid in sorted(set(role_ids))]
        if has_date:
            fields.append({"name": "日期", "type": "date"})
        if has_action:
            fields.append({"name": "签字/盖章", "type": "action"})
        b = DetectedBlock(
            block_id=f"docx_p{pi+1}_{bi:03d}",
            confidence=conf,
            matched_rules=matched,
            fields=fields,
            source_hint=f"paragraph{pi+1}",
            label_preview=(t or "")[:200],
            placeholder_forms=_detect_placeholder_forms(t),
        )
        blocks.append(b)
        bi += 1
        for rid in set(role_ids):
            role_ids_found[rid] = max(role_ids_found.get(rid, 0.0), conf)
        _append_role_evidence(
            role_evidence,
            role_ids,
            conf,
            f"paragraph{pi+1}",
            matched,
            t,
        )

    edge_filter_applied = False
    blocks_dict = [b.to_dict() for b in blocks]
    blocks_dict, edge_filter_applied = _filter_docx_blocks_by_edge_priority(
        blocks_dict,
        total_tables=len(tables_for_detect),
        total_paragraphs=len(doc.paragraphs),
    )
    if edge_filter_applied:
        # 发现封面/文末已有强签批块时，按用户口径仅保留封面/文末签批，忽略正文（含正文表格）。
        role_ids_found = {}
        role_evidence = _filter_role_evidence_to_edge(
            role_evidence,
            total_tables=len(tables_for_detect),
            total_paragraphs=len(doc.paragraphs),
        )

    for b in blocks_dict:
        for f in (b.get("fields") or []):
            if f.get("type") == "role_id" and f.get("name") in ROLE_ID_TO_KEYWORD:
                rid = f["name"]
                role_ids_found[rid] = max(
                    role_ids_found.get(rid, 0.0), float(b.get("confidence") or 0.0) * 0.95
                )

    role_ev = _compact_role_evidence(role_evidence)
    placeholder_forms_seen = sorted({
        f
        for b in blocks_dict
        for f in (b.get("placeholder_forms") or [])
        if f
    })
    placeholder_blocks_count = sum(
        1 for b in blocks_dict if b.get("placeholder_forms")
    )
    debug_summary = {
        "kind": "docx",
        "track_changes_present": bool(has_rev),
        "text_mode": "revision_accepted_ooxml",
        "light_scan": bool(light),
        "tables_scanned": len(tables_for_detect),
        "total_blocks": len(blocks_dict),
        "edge_filter_applied": bool(edge_filter_applied),
        "placeholder_forms_seen": placeholder_forms_seen,
        "placeholder_blocks_count": placeholder_blocks_count,
        "role_count": len(role_ids_found),
        "rules_matched": sorted(
            {
                str(m)
                for b in blocks_dict
                for m in (b.get("matched_rules") or [])
                if str(m).strip()
            }
        ),
    }
    return {
        "ok": True,
        "kind": "docx",
        "roles": [{"id": rid, "confidence": role_ids_found[rid]} for rid in sorted(role_ids_found)],
        "blocks": blocks_dict,
        "role_evidence": role_ev,
        "debug_summary": debug_summary,
    }


def _normalize_detect_hint(detect_hint: Any) -> Dict[str, Any]:
    if not isinstance(detect_hint, dict):
        return {}
    expected = []
    for rid in detect_hint.get("expected_roles") or []:
        rs = str(rid or "").strip()
        if rs in ROLE_ID_TO_KEYWORD and rs not in expected:
            expected.append(rs)
    kws = []
    seen = set()
    for kw in detect_hint.get("label_keywords") or []:
        s = str(kw or "").strip()
        if not s or s in seen:
            continue
        seen.add(s)
        kws.append(s[:64])
        if len(kws) >= 24:
            break
    manual = []
    ocr_kws = []
    for kw in detect_hint.get("manual_keywords") or []:
        s = str(kw or "").strip()
        if s and s not in manual:
            manual.append(s[:64])
    for kw in detect_hint.get("ocr_keywords") or []:
        s = str(kw or "").strip()
        if s and s not in ocr_kws:
            ocr_kws.append(s[:64])
    ocr_stats = detect_hint.get("ocr_hint_stats")
    if not isinstance(ocr_stats, dict):
        ocr_stats = {}
    return {
        "expected_roles": expected,
        "label_keywords": kws,
        "manual_keywords": manual,
        "ocr_keywords": ocr_kws,
        "ocr_hint_stats": ocr_stats,
    }


def _apply_detect_hint_soft(result: dict, detect_hint: Any) -> dict:
    if not isinstance(result, dict) or not result.get("ok"):
        return result
    hint = _normalize_detect_hint(detect_hint)
    expected = hint.get("expected_roles") or []
    kws = hint.get("label_keywords") or []
    manual_kws = hint.get("manual_keywords") or []
    ocr_kws = hint.get("ocr_keywords") or []
    ocr_stats = hint.get("ocr_hint_stats") if isinstance(hint.get("ocr_hint_stats"), dict) else {}
    if not expected and not kws:
        return result
    hint_weight = _get_detect_hint_weight()

    result = dict(result)
    roles = []
    role_conf: Dict[str, float] = {}
    for rr in result.get("roles") or []:
        if not isinstance(rr, dict):
            continue
        rid = str(rr.get("id") or "").strip()
        if rid not in ROLE_ID_TO_KEYWORD:
            continue
        cf = float(rr.get("confidence") or 0.0)
        role_conf[rid] = max(role_conf.get(rid, 0.0), cf)
    blocks = [b for b in (result.get("blocks") or []) if isinstance(b, dict)]
    full_text = " ".join(
        str(b.get("label_preview") or "") + " " + str(b.get("source_hint") or "")
        for b in blocks
    )
    role_evidence = (
        dict(result.get("role_evidence"))
        if isinstance(result.get("role_evidence"), dict)
        else {}
    )
    kw_hits = 0
    for b in blocks:
        txt = (str(b.get("label_preview") or "") + " " + str(b.get("source_hint") or "")).strip()
        if not txt:
            continue
        if kws and any(kw in txt for kw in kws):
            kw_hits += 1
            for rid in expected:
                arr = role_evidence.get(rid) if isinstance(role_evidence.get(rid), list) else []
                arr.append(
                    {
                        "confidence": _scale_hint_conf(0.58, hint_weight),
                        "source_hint": "detect_hint_keyword_context",
                        "matched_rules": ["detect_hint_keyword"],
                        "label_preview": txt[:120],
                    }
                )
                role_evidence[rid] = arr[:6]

    for rid in expected:
        if rid in role_conf:
            role_conf[rid] = max(role_conf[rid], _scale_hint_conf(0.82, hint_weight))
            continue
        if any(kw and kw in full_text for kw in role_keywords(rid)):
            role_conf[rid] = _scale_hint_conf(0.56, hint_weight)
            arr = role_evidence.get(rid) if isinstance(role_evidence.get(rid), list) else []
            arr.append(
                {
                    "confidence": _scale_hint_conf(0.56, hint_weight),
                    "source_hint": "detect_hint_role_recover",
                    "matched_rules": ["detect_hint_expected_role"],
                    "label_preview": (full_text or rid)[:120],
                }
            )
            role_evidence[rid] = arr[:6]

    roles = [{"id": rid, "confidence": role_conf[rid]} for rid in sorted(role_conf)]
    result["roles"] = roles
    result["role_evidence"] = role_evidence
    ocr_doc_hits = [kw for kw in ocr_kws if kw and kw in full_text]
    ds = dict(result.get("debug_summary") or {})
    ds["detect_hint_used"] = True
    ds["detect_hint_weight"] = hint_weight
    ds["detect_hint_expected_roles"] = expected
    if manual_kws:
        ds["detect_hint_manual_keywords"] = manual_kws[:12]
    if ocr_kws:
        ds["detect_hint_ocr_keywords"] = ocr_kws[:12]
    if ocr_doc_hits:
        ds["detect_hint_ocr_keyword_hits"] = ocr_doc_hits[:12]
    if ocr_stats:
        ds["detect_hint_ocr_stats"] = {
            "enabled": bool(ocr_stats.get("enabled")),
            "images_configured": int(ocr_stats.get("images_configured") or 0),
            "images_tried": int(ocr_stats.get("images_tried") or 0),
            "images_read_ok": int(ocr_stats.get("images_read_ok") or 0),
            "skipped_reason": str(ocr_stats.get("skipped_reason") or "")[:64],
        }
    if kws:
        ds["detect_hint_keywords"] = kws[:12]
        ds["detect_hint_keyword_hits"] = kw_hits
    result["debug_summary"] = ds
    return result


_WORK_INSTRUCTION_NAME_RE = re.compile(
    r"作业指导书|过程检验|成品检验|SOP\s*0*\d",
    re.I,
)
_TAIL_INSPECTOR_LABEL_RE = re.compile(
    r"检验人|检验员|检查人|质检员|质检负责人",
    re.I,
)
_SIGNOFF_LABEL_RE = re.compile(r"编制|编写|作者|审核|复核|批准", re.I)


def _block_is_tail_inspector_only(block: dict) -> bool:
    """文末「检验人/检验员」等误匹配 reviewer_tail，非封面编审批签批栏。"""
    if not isinstance(block, dict):
        return False
    fields = block.get("fields") if isinstance(block.get("fields"), list) else []
    role_names = [
        str(f.get("name") or "").strip()
        for f in fields
        if isinstance(f, dict) and f.get("type") == "role_id"
    ]
    if not role_names:
        return False
    preview = (str(block.get("label_preview") or "") + str(block.get("source_hint") or "")).strip()
    if not _TAIL_INSPECTOR_LABEL_RE.search(preview):
        return False
    if _SIGNOFF_LABEL_RE.search(preview):
        return False
    canon = {canonical_sign_role_id(r) for r in role_names}
    return canon <= {"reviewer", "reviewer_tail"}


def _harmonize_role_evidence(role_evidence: Any) -> Dict[str, List[Dict[str, Any]]]:
    if not isinstance(role_evidence, dict):
        return {}
    out: Dict[str, List[Dict[str, Any]]] = {}
    for rid, arr in role_evidence.items():
        rid2 = canonical_sign_role_id(str(rid))
        if rid2 not in ROLE_ID_TO_KEYWORD:
            continue
        bucket = out.setdefault(rid2, [])
        if not isinstance(arr, list):
            continue
        for ev in arr:
            if isinstance(ev, dict) and ev not in bucket:
                bucket.append(ev)
        out[rid2] = bucket[:6]
    return out


def _harmonize_detect_result(result: dict, source_name: str = "") -> dict:
    """
    统一需签角色集：reviewer_tail 并入 reviewer；作业指导书等已有编审批三角色时，
    去掉文末「检验人」误识别的 reviewer_tail 块，避免签字位与需签角色列不一致。
    """
    if not isinstance(result, dict) or not result.get("ok"):
        return result
    out = dict(result)
    role_conf: Dict[str, float] = {}
    for rr in out.get("roles") or []:
        if not isinstance(rr, dict):
            continue
        rid = canonical_sign_role_id(str(rr.get("id") or "").strip())
        if rid not in ROLE_ID_TO_KEYWORD:
            continue
        role_conf[rid] = max(role_conf.get(rid, 0.0), float(rr.get("confidence") or 0.0))

    if "reviewer_tail" in role_conf:
        rt = role_conf.pop("reviewer_tail")
        role_conf["reviewer"] = max(role_conf.get("reviewer", 0.0), rt)

    nm = str(source_name or "").strip()
    is_work_instruction = bool(_WORK_INSTRUCTION_NAME_RE.search(nm))
    has_triad = {"author", "reviewer", "approver"}.issubset(role_conf.keys())

    blocks_in = [b for b in (out.get("blocks") or []) if isinstance(b, dict)]
    blocks_out: List[dict] = []
    for b in blocks_in:
        if is_work_instruction and has_triad and _block_is_tail_inspector_only(b):
            continue
        b2 = dict(b)
        fields = []
        for f in b.get("fields") or []:
            if not isinstance(f, dict):
                continue
            f2 = dict(f)
            if f2.get("type") == "role_id" and f2.get("name"):
                f2["name"] = canonical_sign_role_id(str(f2["name"]))
            fields.append(f2)
        b2["fields"] = fields
        blocks_out.append(b2)

    if is_work_instruction and has_triad:
        role_conf.pop("reviewer_tail", None)

    # 块内角色与顶层 roles 对齐
    for b in blocks_out:
        for f in b.get("fields") or []:
            if not isinstance(f, dict) or f.get("type") != "role_id":
                continue
            rid = str(f.get("name") or "").strip()
            if rid not in ROLE_ID_TO_KEYWORD:
                continue
            role_conf[rid] = max(role_conf.get(rid, 0.0), float(b.get("confidence") or 0.0) * 0.95)

    out["roles"] = [
        {"id": rid, "confidence": role_conf[rid]} for rid in sorted(role_conf)
    ]
    out["blocks"] = blocks_out
    out["role_evidence"] = _harmonize_role_evidence(out.get("role_evidence"))
    ds = dict(out.get("debug_summary") or {})
    if is_work_instruction and has_triad:
        ds["roles_harmonized"] = "work_instruction_triad"
    out["debug_summary"] = ds
    return out


def detect_file(path: str, source_name: str = "", mode: str = "auto", detect_hint: Any = None) -> dict:
    ext = os.path.splitext(path)[1].lower()
    light = mode == "light"
    # 规则/标误只作为提示，不以其结论替代真实识别；auto 模式保持正常扫描深度。
    src = source_name or os.path.basename(path)
    if ext == ".xlsx":
        result = _apply_detect_hint_soft(detect_xlsx(path), detect_hint)
    elif ext == ".docx":
        result = _apply_detect_hint_soft(detect_docx(path, light=light), detect_hint)
    else:
        return {"ok": False, "error": f"不支持的格式: {ext}"}
    if isinstance(result, dict) and result.get("ok"):
        result = _harmonize_detect_result(result, source_name=src)
    return result

