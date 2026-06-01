# -*- coding: utf-8 -*-
"""
Word .docx：在模板已预留的空白处插入签名/日期 PNG（不新增签名字段）。
定位顺序：
1) 表格内「角色/姓名格 + 邻列 Date/日期 表头」→ 签名只插角色侧（姓名后或尾部空白），日期只插日期列；
2) 三列表「角色+姓名 | 空签字列 | 日期列」→ 签名插中间格，日期插日期列；
3) 原逻辑：关键词单元格 → 右侧空单元格；段落内关键词后的占位符/下划线 run。
"""
from __future__ import annotations

import io
import os
import re
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.document import Document as DocumentObject
from docx.shared import Cm
from docx.table import Table
from docx.text.paragraph import Paragraph

from sign_handlers.config import (
    ROLE_ID_TO_KEYWORD,
    is_replaceable_prefilled_slot_text,
    role_keywords_for_apply,
)
from sign_handlers.label_match import (
    cell_has_label_inline_reservation,
    cell_has_role_keyword,
    cell_has_signoff_inline_reservation,
    cell_is_bare_role_column_header,
    cell_is_role_signoff_label_slot,
    cell_looks_like_signoff_date_label,
    cell_text_matches_keyword,
    paragraph_text_keyword_end_offset,
    xlsx_cell_has_leading_role_keyword,
)
from sign_handlers.docx_revision_text import cell_effective_text, paragraph_effective_text
from sign_handlers.png_word_compat import (
    prepare_png_for_word,
    prepare_signature_date_pair_for_word,
)

# 可见占位字符（作者：____ 等）
_PLACEHOLDER_CHARS = re.compile(r"^[\s_\-—–\u2014\u2015\u2500\u3000\.·]+$")
# 右侧「日期」列表头（整格以 Date/日期 开头）
_DATE_HEADER_CELL = re.compile(r"^\s*(日期|Date)\s*[:：]?\s*", re.IGNORECASE)
_REVISION_ROW_NOISE_TERMS = (
    "更改历史",
    "修订记录",
    "变更记录",
    "版本",
    "更改日期",
    "变更日期",
    "修订日期",
    "变更内容",
    "更改内容",
    "修订内容",
    "变更号",
    "修订号",
    "版本号",
    "版次",
    "修订内容",
    "修订说明",
)
_REVISION_VERSION_TOKEN_RE = re.compile(
    r"(?:\b[A-Z]/\d+\b|\bV?\d+(?:\.\d+){1,4}\b|版本\s*[A-Z0-9./-]+)",
    re.IGNORECASE,
)
_REVISION_DATE_TOKEN_RE = re.compile(
    r"(?:20\d{2}[./-]\d{1,2}[./-]\d{1,2}|日期|Date)",
    re.IGNORECASE,
)
_TABLE_HEADER_NOISE_TERMS = (
    "用例",
    "步骤",
    "预期",
    "结果",
    "测试结果",
    "测试人员",
    "测试人",
    "执行结果",
    "执行人员",
    "执行人",
    "测试项",
    "测试内容",
    "traceability",
    "requirement",
)
_TABLE_SIGNOFF_HINT_TERMS = (
    "/日期",
    "日期：",
    "Date:",
    "Date：",
    "编制",
    "审核",
    "批准",
    "复核",
    "签字",
    "签名",
)

_PIC_WIDTH = Cm(2.8)
# 日期与签名使用同一显示宽度，避免缩放后笔画变细、观感不一致
_PIC_WIDTH_DATE = _PIC_WIDTH


def _is_emptyish_text(s: str) -> bool:
    if s is None:
        return True
    t = str(s).strip()
    if not t:
        return True
    if _PLACEHOLDER_CHARS.match(t):
        return True
    return len(t) < 2


def _is_slot_target_text(s: str) -> bool:
    """签名/日期可落位目标：空白占位，或可替换的电脑输入值。"""
    return _is_emptyish_text(s) or is_replaceable_prefilled_slot_text(s)


def _cell_text(cell) -> str:
    return (cell_effective_text(cell) or "").strip()


def _paragraph_full_text(p: Paragraph) -> str:
    return paragraph_effective_text(p) or ""


def _run_visible_text(r) -> str:
    try:
        from sign_handlers.docx_revision_text import ooxml_effective_text

        return ooxml_effective_text(r._element)
    except Exception:
        return r.text or ""


def _run_has_underline_or_border(r) -> bool:
    try:
        f = r.font
        if f is not None and f.underline:
            return True
    except Exception:
        pass
    try:
        if r._element.rPr is not None:
            ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            rpr = r._element.rPr
            if rpr.find(f"{ns}u") is not None:
                return True
    except Exception:
        pass
    return False


def _find_keyword_in_paragraph(p: Paragraph, keyword: str) -> int:
    return paragraph_text_keyword_end_offset(_paragraph_full_text(p), keyword)


def _clear_runs_after_offset(p: Paragraph, start_char_offset: int) -> None:
    """从段落文本的 start_char_offset 起删除文本（用于去掉 ____ 占位）。"""
    full = _paragraph_full_text(p)
    if start_char_offset <= 0 or start_char_offset >= len(full):
        return
    # 勿用 strip() 判断「无内容」：关键词后常见仅 Tab/全角空格/制表符前导点线，
    # strip 后为空会误判并跳过删除，导致后续 add_run 把图片追加到段末（离标签很远、易换行）。
    if not full[start_char_offset:]:
        return
    acc = 0
    truncate_after = None
    for i, r in enumerate(list(p.runs)):
        rt = _run_visible_text(r)
        ln = len(rt)
        if ln == 0 and (r.text or ""):
            acc += len(r.text or "")
            continue
        if acc + ln <= start_char_offset:
            acc += ln
            continue
        if acc >= start_char_offset:
            if _run_visible_text(r) or not (r.text or ""):
                r.text = ""
            acc += max(ln, len(r.text or ""))
            continue
        cut = start_char_offset - acc
        if _run_visible_text(r):
            raw = r.text or ""
            r.text = raw[:cut] if raw else ""
        truncate_after = i
        acc += ln
        break
    if truncate_after is not None:
        for r in list(p.runs)[truncate_after + 1 :]:
            if _run_visible_text(r) or not (r.text or ""):
                r.text = ""


def _truncate_after_offset_and_get_anchor_run_idx(
    p: Paragraph, start_char_offset: int
) -> int:
    """
    从 start_char_offset 起清空文本，并返回“插入锚点 run”的索引：
    - 若 offset 落在某个 run 内：该 run 被截断到 offset，返回该 run idx
    - 若 offset 正好在 run 边界：返回 offset 前的 run idx（最靠近标题的那个）
    - 若无法定位：返回最后一个 run idx（或 -1 表示无 run）
    """
    full = _paragraph_full_text(p)
    if not p.runs:
        return -1
    if start_char_offset <= 0:
        # 标题在段首：锚点设为第一个 run（图片插在它后）
        return 0
    if start_char_offset >= len(full):
        return len(p.runs) - 1

    def _run_is_placeholder(r) -> bool:
        rt = _run_visible_text(r) or (r.text or "")
        # 下划线/边框 run 通常是“__ / 点线”占位；直接当作占位清空
        if _run_has_underline_or_border(r):
            return True
        # 仅包含空白/下划线/横线/点线等字符：当作占位清空
        try:
            if _PLACEHOLDER_CHARS.match(rt):
                return True
        except Exception:
            pass
        # 兜底：全空白也当作占位
        if not rt or not str(rt).strip():
            return True
        return False

    acc = 0
    last_nonempty_idx = 0
    for i, r in enumerate(list(p.runs)):
        rt = _run_visible_text(r)
        ln = len(rt)
        if not ln and (r.text or ""):
            acc += len(r.text or "")
            continue
        if rt:
            last_nonempty_idx = i
        seg = acc + ln
        if seg < start_char_offset:
            acc = seg
            continue
        if seg == start_char_offset:
            for rr in list(p.runs)[i + 1 :]:
                if _run_is_placeholder(rr):
                    rr.text = ""
            return i
        cut = start_char_offset - acc
        raw = r.text or ""
        if rt:
            r.text = raw[:cut] if raw else ""
        for rr in list(p.runs)[i + 1 :]:
            if _run_is_placeholder(rr):
                rr.text = ""
        return i
    # fallback：按原逻辑清空，锚点取最后一个含文本的 run
    _clear_runs_after_offset(p, start_char_offset)
    return last_nonempty_idx


def _move_run_after_run_idx(p: Paragraph, new_run, anchor_run_idx: int) -> None:
    """
    python-docx 只能在段末 add_run；此处把 new_run 的 XML 节点移动到 anchor_run_idx 后面，
    以保证图片紧接字段标题后（不会跑到点线/空格末尾导致换行）。
    """
    try:
        if anchor_run_idx < 0:
            return
        runs = list(p.runs)
        if not runs:
            return
        if anchor_run_idx >= len(runs):
            anchor_run_idx = len(runs) - 1
        anchor_r = runs[anchor_run_idx]._r
        new_r = new_run._r
        # 先从末尾移除，再插入到 anchor 后
        p._p.remove(new_r)
        insert_pos = p._p.index(anchor_r) + 1
        p._p.insert(insert_pos, new_r)
    except Exception:
        # 若低层 XML 操作失败，退化为段末追加（不致使签署失败）
        return


def _insert_pictures_in_paragraph(
    p: Paragraph, sig_png: Optional[bytes], date_png: Optional[bytes]
) -> None:
    """在同一段落中按需插入签名/日期（可只插入其一）。默认追加到段末。"""
    if sig_png:
        p.add_run().add_picture(io.BytesIO(sig_png), width=_PIC_WIDTH)
    if sig_png and date_png:
        p.add_run(" ")
    if date_png:
        p.add_run().add_picture(io.BytesIO(date_png), width=_PIC_WIDTH_DATE)


def _insert_pictures_after_anchor_run_idx(
    p: Paragraph,
    anchor_run_idx: int,
    sig_png: Optional[bytes],
    date_png: Optional[bytes],
) -> None:
    """
    在 anchor_run_idx 之后插入签名/日期图片，并把对应 run 移动到锚点后，
    以保证紧贴字段标题，不会跑到点线/空格末尾导致换行。
    """
    runs_to_place = []
    if sig_png:
        rs = p.add_run()
        rs.add_picture(io.BytesIO(sig_png), width=_PIC_WIDTH)
        runs_to_place.append(rs)
    if sig_png and date_png:
        runs_to_place.append(p.add_run(" "))
    if date_png:
        rd = p.add_run()
        rd.add_picture(io.BytesIO(date_png), width=_PIC_WIDTH_DATE)
        runs_to_place.append(rd)

    # 逐个移动：每次都插到“当前锚点后”，并推进锚点以保持顺序
    cur_anchor = anchor_run_idx
    for r in runs_to_place:
        _move_run_after_run_idx(p, r, cur_anchor)
        cur_anchor += 1


def _cell_looks_like_date_header_cell(text: str) -> bool:
    t = (text or "").strip()
    return bool(_DATE_HEADER_CELL.match(t))


def _find_date_label_end_in_paragraph(p: Paragraph) -> int:
    full = _paragraph_full_text(p)
    if not full:
        return -1
    m = re.search(r"(?i)(?:\bDate\b|日期)\s*[:：]?", full)
    return m.end() if m else -1


def _insert_sig_only_at_char_offset(p: Paragraph, char_offset: int, sig_png: bytes) -> None:
    # 紧接「Author:/Approver:」等标题后插入：从插入点起清空其后文本/点线占位，
    # 再把图片 run 移到标题 run 后（避免落在点线/空格末尾导致换行）。
    anchor_idx = _truncate_after_offset_and_get_anchor_run_idx(p, char_offset)
    rn = p.add_run()
    rn.add_picture(io.BytesIO(sig_png), width=_PIC_WIDTH)
    _move_run_after_run_idx(p, rn, anchor_idx)


def _insert_date_only_in_date_cell(date_cell, date_png: bytes) -> bool:
    """在「日期/Date」标签后的空白处只插入日期手写图。"""
    for p in date_cell.paragraphs:
        offd = _find_date_label_end_in_paragraph(p)
        if offd < 0:
            continue
        anchor_idx = _truncate_after_offset_and_get_anchor_run_idx(p, offd)
        rn = p.add_run()
        rn.add_picture(io.BytesIO(date_png), width=_PIC_WIDTH_DATE)
        _move_run_after_run_idx(p, rn, anchor_idx)
        return True
    if date_cell.paragraphs:
        p0 = date_cell.paragraphs[0]
        for r in list(p0.runs):
            r.text = ""
        rn = p0.add_run()
        rn.add_picture(io.BytesIO(date_png), width=_PIC_WIDTH_DATE)
        return True
    p = date_cell.add_paragraph()
    rn = p.add_run()
    rn.add_picture(io.BytesIO(date_png), width=_PIC_WIDTH_DATE)
    return True


def _insert_sig_in_role_cell_for_adjacent_date_column(
    role_cell, keyword: str, sig_png: bytes
) -> bool:
    for p in role_cell.paragraphs:
        off = _find_keyword_in_paragraph(p, keyword)
        if off < 0:
            continue
        _insert_sig_only_at_char_offset(p, off, sig_png)
        return True
    return False


def _insert_sig_only_in_empty_cell(cell, sig_png: bytes) -> None:
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    for r in list(p.runs):
        r.text = ""
    p.add_run().add_picture(io.BytesIO(sig_png), width=_PIC_WIDTH)


def _insert_sig_and_date_in_empty_cell(cell, sig_png: Optional[bytes], date_png: Optional[bytes]) -> None:
    """把签名/日期插入到“预留空白格”中（清空原占位）。"""
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    for r in list(p.runs):
        r.text = ""
    if sig_png:
        p.add_run().add_picture(io.BytesIO(sig_png), width=_PIC_WIDTH)
    if sig_png and date_png:
        p.add_run(" ")
    if date_png:
        p.add_run().add_picture(io.BytesIO(date_png), width=_PIC_WIDTH_DATE)


def _append_date_to_cell(cell, date_png: bytes) -> None:
    """在已有签名图的空白格后追加日期图（签批栏签名/日期分列）。"""
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if p.runs:
        p.add_run(" ")
    p.add_run().add_picture(io.BytesIO(date_png), width=_PIC_WIDTH_DATE)


def _parse_docx_cell_loc(loc: str) -> Optional[Tuple[int, int, int]]:
    """table#2.r10.c3 → (2, 10, 3) 1-based。"""
    s = str(loc or "").strip()
    m = re.match(r"^\s*table#?(\d+)\.r(\d+)\.c(\d+)\s*$", s, re.IGNORECASE)
    if not m:
        return None
    return int(m.group(1)), int(m.group(2)), int(m.group(3))


def _docx_detect_ordered_tables(doc: DocumentObject) -> List[Table]:
    try:
        from sign_handlers.detect_fields import _iter_docx_tables_for_detect

        return list(
            _iter_docx_tables_for_detect(doc, max_body_tables=32, max_total=64)
        )
    except Exception:
        return []


def _docx_row_cells(table: Table, row_1based: int) -> Tuple[int, List]:
    """返回 (ri, cells) 或 (-1, [])。"""
    ri = int(row_1based) - 1
    rows_list = list(table.rows)
    if ri < 0 or ri >= len(rows_list):
        return -1, []
    return ri, list(rows_list[ri].cells)


def _place_four_column_row(
    cells,
    label_ci: int,
    keyword: str,
    sig_png: Optional[bytes],
    date_png: Optional[bytes],
) -> bool:
    """角色 | 姓名空白 | 日期标签 | 日期空白。"""
    if not _is_label_blank_date_blank_row(cells, label_ci, keyword):
        return False
    placed = False
    sig_cell = cells[label_ci + 1]
    date_hdr = cells[label_ci + 2]
    date_val = cells[label_ci + 3]
    if sig_png and _is_slot_target_text(_cell_text(sig_cell)):
        _insert_sig_only_in_empty_cell(sig_cell, sig_png)
        placed = True
    if date_png:
        if _is_slot_target_text(_cell_text(date_val)):
            _insert_sig_and_date_in_empty_cell(date_val, None, date_png)
            placed = True
        elif _insert_date_only_in_date_cell(date_hdr, date_png):
            placed = True
        else:
            dt_target = _next_distinct_cell(cells, label_ci + 2)
            if dt_target is not None and _is_slot_target_text(_cell_text(dt_target)):
                _insert_sig_and_date_in_empty_cell(dt_target, None, date_png)
                placed = True
    if sig_png and not date_png:
        return placed
    if date_png:
        return placed
    return placed


def _try_docx_role_layout_cells(
    doc: DocumentObject,
    role_id: str,
    placement_plan: Optional[dict],
    keywords: List[str],
    sig_png: Optional[bytes],
    date_png: Optional[bytes],
) -> bool:
    """
    按识别阶段 signature_layout 给出的单元格坐标直接落位（与 detect 表序号一致）。
    仅在格可写时成功，不伪造结果。
    """
    if not isinstance(placement_plan, dict):
        return False
    rp = placement_plan.get(str(role_id))
    if not isinstance(rp, dict):
        return False
    name_cell = rp.get("name_cell")
    if not isinstance(name_cell, dict):
        return False
    want_date = bool(rp.get("date_slot", True))
    dt = date_png if want_date else None
    tno = int(name_cell.get("table") or 0)
    rno = int(name_cell.get("row") or 0)
    if tno < 1 or rno < 1:
        return False
    tables = _docx_detect_ordered_tables(doc)
    if tno > len(tables):
        return False
    table = tables[tno - 1]
    ri, cells = _docx_row_cells(table, rno)
    if ri < 0 or not cells:
        return False
    rows_list = list(table.rows)
    for kw in keywords:
        for ci, cell in enumerate(cells):
            ct = _cell_text(cell)
            if not cell_has_role_keyword(ct, kw):
                continue
            if _place_four_column_row(cells, ci, sig_png, dt, kw):
                return True
            if _place_sig_date_at_signoff_anchor(
                cell, kw, cells, ci, rows_list, ri, sig_png, dt
            ):
                return True
    return False


def _is_label_blank_date_blank_row(cells, label_ci: int, keyword: str) -> bool:
    """
    四列签批：角色标签 | 姓名空白 | 日期标签 | 日期空白。
    软件发布说明、部分作业指导书常见；标签格仅为「作者/审核/批准」时仍应可落位。
    """
    if label_ci + 3 >= len(cells):
        return False
    lt = _cell_text(cells[label_ci])
    if not cell_has_role_keyword(lt, keyword):
        return False
    if not _is_slot_target_text(_cell_text(cells[label_ci + 1])):
        return False
    if not _cell_looks_like_date_header_cell(_cell_text(cells[label_ci + 2])):
        return False
    if not _is_slot_target_text(_cell_text(cells[label_ci + 3])):
        return False
    return True


def _next_distinct_cell(cells, idx: int):
    """返回同一行里 idx 之后第一个“不同的 cell”（按合并单元格去重）。"""
    try:
        base = cells[idx]
        base_tc = getattr(base, "_tc", None)
        for j in range(idx + 1, len(cells)):
            c2 = cells[j]
            if getattr(c2, "_tc", None) is not None and getattr(c2, "_tc", None) == base_tc:
                continue
            return c2
    except Exception:
        pass
    return None


def _try_table_adjacent_date_column(
    table: Table,
    keyword: str,
    sig_png: Optional[bytes],
    date_png: Optional[bytes],
) -> bool:
    """
    常见审批表：
    - 两列：左格「角色/姓名 + 签字空白」+ 右邻格「Date:/日期: + 空白」→ 签名只插左格，日期只插右格。
    - 三列：角色+姓名 | 空签字列 | 日期列 → 签名优先插左格关键词/冒号后（更贴近字段标记），日期插右格。
    避免两段图都落在左侧姓名后。
    """
    rows_list = list(table.rows)
    for ri in _table_row_scan_order(table):
        cells = rows_list[ri].cells
        for ci in range(len(cells) - 1):
            left = cells[ci]
            lt = _cell_text(left)
            if not xlsx_cell_has_leading_role_keyword(lt, keyword):
                continue
            if cell_is_bare_role_column_header(lt, keyword):
                if not _is_label_blank_date_blank_row(cells, ci, keyword):
                    continue
            date_idx = None
            if ci + 1 < len(cells) and _cell_looks_like_date_header_cell(_cell_text(cells[ci + 1])):
                date_idx = ci + 1
            elif ci + 2 < len(cells) and _cell_looks_like_date_header_cell(_cell_text(cells[ci + 2])):
                date_idx = ci + 2
            if date_idx is None:
                continue
            date_cell = cells[date_idx]
            placed_any = False
            if sig_png:
                if date_idx == ci + 1:
                    # 只有两列且右列就是 Date 标签：签名只能写在左格关键词后
                    if not _insert_sig_in_role_cell_for_adjacent_date_column(left, keyword, sig_png):
                        continue
                else:
                    # 三列（或更多）：字段格后面通常有“预留空白签字格”，优先插到该空白格里
                    mid = _next_distinct_cell(cells, ci)
                    if mid is None or mid._tc == date_cell._tc:
                        # 没有独立签字格：退回写在字段格关键词后
                        if not _insert_sig_in_role_cell_for_adjacent_date_column(left, keyword, sig_png):
                            continue
                    else:
                        if not _is_slot_target_text(_cell_text(mid)):
                            continue
                        _insert_sig_only_in_empty_cell(mid, sig_png)
                placed_any = True
            if date_png:
                # 日期：优先插到 Date 标签格的右侧空白格；否则退回 Date 标签格内
                date_target = None
                try:
                    date_target = _next_distinct_cell(cells, date_idx)
                except Exception:
                    date_target = None
                if date_target is not None and _is_slot_target_text(_cell_text(date_target)):
                    _insert_sig_and_date_in_empty_cell(date_target, None, date_png)
                    placed_any = True
                elif not _insert_date_only_in_date_cell(date_cell, date_png):
                    # 仅有日期时若没能插入（极少），继续找其它匹配
                    if not placed_any:
                        continue
                placed_any = True
            if placed_any:
                return True
    return False


def _try_paragraph_inline(
    p: Paragraph,
    keyword: str,
    sig_png: Optional[bytes],
    date_png: Optional[bytes],
) -> bool:
    """段落内：关键词后占位符清除并插图，或下划线 run 后插图，否则关键词后追加。"""
    off = _find_keyword_in_paragraph(p, keyword)
    if off < 0:
        return False
    full = _paragraph_full_text(p)
    rest = full[off:]
    rest_stripped = rest.lstrip(" ：:\t")
    # 同段有「日期」且要先签后日期：用户按角色提供 sig+date，都插在同一角色关联区域
    head = (rest_stripped[:20] if len(rest_stripped) > 20 else rest_stripped) if rest_stripped else ""
    if _is_emptyish_text(rest_stripped) or (head and _PLACEHOLDER_CHARS.match(head)):
        anchor_idx = _truncate_after_offset_and_get_anchor_run_idx(p, off)
        _insert_pictures_after_anchor_run_idx(p, anchor_idx, sig_png, date_png)
        return True
    # 尝试：从第一个 run 开始找关键词后的下划线 run
    acc = 0
    found_kw_end = False
    insert_after_run_idx = None
    for i, r in enumerate(p.runs):
        rt = r.text or ""
        seg = acc + len(rt)
        if not found_kw_end:
            # off 为关键词匹配结束位置（见 paragraph_text_keyword_end_offset）
            if acc < off <= seg:
                found_kw_end = True
        if found_kw_end and acc >= off and _run_has_underline_or_border(r):
            insert_after_run_idx = i
            break
        acc = seg
    if insert_after_run_idx is not None:
        # 在关键词后（通常为冒号后）清空占位，并把图片 run 移动到标题 run 后
        anchor_idx = _truncate_after_offset_and_get_anchor_run_idx(p, off)
        _insert_pictures_after_anchor_run_idx(p, anchor_idx, sig_png, date_png)
        return True
    # 关键词后已有非占位正文：避免误删，交其它段落/表格再匹配
    if rest_stripped and not _PLACEHOLDER_CHARS.match(rest_stripped):
        return False
    anchor_idx = _truncate_after_offset_and_get_anchor_run_idx(p, off)
    _insert_pictures_after_anchor_run_idx(p, anchor_idx, sig_png, date_png)
    return True


def _table_row_scan_order(table: Table) -> range:
    """宽表（如用例表）从下往上找签批行，避免先命中列头「测试人」。"""
    n = len(table.rows)
    if n > 10:
        return range(n - 1, -1, -1)
    return range(n)


def _find_date_label_cell_in_row(cells, label_ci: int, rows_list, ri: int):
    """同行或下一行找「日期/Date」列表头格。"""
    date_cell = None
    for j in range(label_ci + 1, len(cells)):
        if _cell_looks_like_date_header_cell(_cell_text(cells[j])):
            date_cell = cells[j]
            break
    if date_cell is None and ri + 1 < len(rows_list):
        nrow = rows_list[ri + 1]
        for cand_col in (label_ci, label_ci + 1, label_ci + 2):
            if cand_col < len(nrow.cells) and _cell_looks_like_date_header_cell(
                _cell_text(nrow.cells[cand_col])
            ):
                date_cell = nrow.cells[cand_col]
                break
    return date_cell


def _iter_reserved_blank_cells_same_row(cells, label_ci: int, keyword: str):
    """标签格右侧连续空白格（签批栏：第一格签名、第二格日期，遇下一角色标签则停止）。"""
    try:
        base_tc = getattr(cells[label_ci], "_tc", None)
    except Exception:
        base_tc = None
    for j in range(label_ci + 1, len(cells)):
        c2 = cells[j]
        try:
            if base_tc is not None and getattr(c2, "_tc", None) == base_tc:
                continue
        except Exception:
            pass
        t = _cell_text(c2)
        if _cell_looks_like_date_header_cell(t):
            continue
        if (
            cell_looks_like_signoff_date_label(t)
            and not cell_has_role_keyword(t, keyword)
        ):
            break
        if _is_emptyish_text(t):
            yield c2
            continue
        if is_replaceable_prefilled_slot_text(t):
            yield c2
            continue
        if cell_has_role_keyword(t, keyword):
            continue
        break


def _first_reserved_blank_cell_same_row(cells, label_ci: int, keyword: str):
    """标签格右侧：第一个空白 distinct 格（签字区，跳过日期列标题格）。"""
    for c2 in _iter_reserved_blank_cells_same_row(cells, label_ci, keyword):
        return c2
    return None


def _first_reserved_blank_cell_below(rows_list, ri: int, label_ci: int):
    """标签格正下方（或略偏右）的空白格。"""
    if ri + 1 >= len(rows_list):
        return None
    nrow = rows_list[ri + 1]
    for cand_col in (label_ci, label_ci + 1, label_ci + 2):
        if cand_col < len(nrow.cells):
            t = _cell_text(nrow.cells[cand_col])
            if _is_slot_target_text(t) and not _cell_looks_like_date_header_cell(t):
                return nrow.cells[cand_col]
    return None


def _cells_joined_text(cells) -> str:
    parts = []
    for c in cells or []:
        t = _cell_text(c)
        if t:
            parts.append(t)
    return " | ".join(parts)


def _is_revision_history_row(cells) -> bool:
    txt = _cells_joined_text(cells)
    if not txt:
        return False
    txt_l = txt.lower()
    hit = 0
    for term in _REVISION_ROW_NOISE_TERMS:
        if term in txt:
            hit += 1
            if hit >= 2:
                return True
    # 仅当出现「变更/修订/更改」这类明确修订动词且带版本/日期特征时才判为修订记录行；
    # 签批行不会出现这些动词，避免误伤「编制/审核/批准 + 日期」的签批行。
    if ("变更" in txt or "修订" in txt or "更改" in txt) and (
        _REVISION_VERSION_TOKEN_RE.search(txt) or _REVISION_DATE_TOKEN_RE.search(txt)
    ):
        return True
    # 英文版修订记录
    if ("revision" in txt_l or "change history" in txt_l) and _REVISION_DATE_TOKEN_RE.search(txt):
        return True
    return False


def _looks_like_data_table_header_row(cells) -> bool:
    txt = _cells_joined_text(cells)
    if not txt:
        return False
    txt_l = txt.lower()
    # 有明确签批提示词时，不按表头拦截
    for term in _TABLE_SIGNOFF_HINT_TERMS:
        if term.lower() in txt_l:
            return False
    nonempty = 0
    for c in cells or []:
        if _cell_text(c):
            nonempty += 1
    if nonempty < 5:
        return False
    hit = 0
    for term in _TABLE_HEADER_NOISE_TERMS:
        if term in txt_l:
            hit += 1
    return hit >= 2


def _cell_has_table_signoff_reservation(
    cell_text: str,
    keyword: str,
    cells,
    label_ci: int,
    rows_list,
    ri: int,
) -> bool:
    """
    表格签字锚点：有角色标签，且存在留位证据（同格下划线/空格、右侧空白格、下方空白格、角色/日期）。
    作者/审核/批准/测试人等统一规则。
    """
    if not cell_has_role_keyword(cell_text, keyword):
        return False
    # 变更历史/修订记录中的「作者」不是签批栏，不应落签。
    if _is_revision_history_row(cells):
        return False
    if cell_is_bare_role_column_header(cell_text, keyword):
        return _is_label_blank_date_blank_row(cells, label_ci, keyword)
    if cell_is_role_signoff_label_slot(cell_text, keyword):
        return True
    if cell_has_signoff_inline_reservation(cell_text, keyword):
        return True
    if cell_has_label_inline_reservation(cell_text, keyword):
        return True
    date_nearby = _find_date_label_cell_in_row(cells, label_ci, rows_list, ri) is not None
    same_row_slot = _first_reserved_blank_cell_same_row(cells, label_ci, keyword)
    if same_row_slot is not None:
        slot_t = _cell_text(same_row_slot)
        if _is_emptyish_text(slot_t):
            return True
        if is_replaceable_prefilled_slot_text(slot_t) and date_nearby:
            return True
    below_slot = _first_reserved_blank_cell_below(rows_list, ri, label_ci)
    if below_slot is not None:
        slot_t = _cell_text(below_slot)
        if _is_emptyish_text(slot_t):
            return True
        if is_replaceable_prefilled_slot_text(slot_t) and date_nearby:
            return True
    return False


def _place_sig_date_at_signoff_anchor(
    label_cell,
    keyword: str,
    cells,
    label_ci: int,
    rows_list,
    ri: int,
    sig_png: Optional[bytes],
    date_png: Optional[bytes],
) -> bool:
    """
    在签批锚点落位（顺序固定）：
    1) 签批栏「角色/日期」：同行第 1 个空白格签名、第 2 个空白格日期（与复核人一致）；
    2) 其它：同行/下方空白格；同格下划线留白；独立 Date/日期 列。
    """
    if _place_four_column_row(cells, label_ci, keyword, sig_png, date_png):
        return True

    placed_any = False
    ct = _cell_text(label_cell)
    date_cell = _find_date_label_cell_in_row(cells, label_ci, rows_list, ri)
    signoff_slot = cell_is_role_signoff_label_slot(ct, keyword)
    # 同格「角色/日期：____」优先一次性写入签名+日期，避免先落日期导致签名失败。
    if (
        signoff_slot
        and sig_png
        and date_png
        and cell_has_signoff_inline_reservation(ct, keyword)
    ):
        for p0 in label_cell.paragraphs:
            if _try_paragraph_inline(p0, keyword, sig_png, date_png):
                return True
    row_blanks = (
        list(_iter_reserved_blank_cells_same_row(cells, label_ci, keyword))
        if signoff_slot
        else []
    )

    if sig_png:
        if row_blanks:
            _insert_sig_and_date_in_empty_cell(row_blanks[0], sig_png, None)
            placed_any = True
        else:
            sig_blank = _first_reserved_blank_cell_same_row(cells, label_ci, keyword)
            if sig_blank is not None:
                _insert_sig_and_date_in_empty_cell(sig_blank, sig_png, None)
                placed_any = True
        if not placed_any:
            below = _first_reserved_blank_cell_below(rows_list, ri, label_ci)
            if below is not None:
                _insert_sig_and_date_in_empty_cell(below, sig_png, None)
                placed_any = True
        if not placed_any and cell_has_label_inline_reservation(ct, keyword):
            if _insert_sig_in_role_cell_for_adjacent_date_column(label_cell, keyword, sig_png):
                placed_any = True
        if not placed_any and signoff_slot:
            sig_blank = _next_distinct_cell(cells, label_ci)
            if sig_blank is not None and _is_slot_target_text(_cell_text(sig_blank)):
                if date_cell is None or sig_blank._tc != date_cell._tc:
                    _insert_sig_and_date_in_empty_cell(sig_blank, sig_png, None)
                    placed_any = True
        if not placed_any and signoff_slot and cell_has_signoff_inline_reservation(
            ct, keyword
        ):
            for p0 in label_cell.paragraphs:
                if _try_paragraph_inline(p0, keyword, sig_png, None):
                    placed_any = True
                    break

    if date_png:
        date_placed = False
        if date_cell is not None and not signoff_slot:
            dt_target = None
            try:
                dt_target = _next_distinct_cell(cells, cells.index(date_cell))
            except Exception:
                dt_target = None
            if dt_target is not None and _is_slot_target_text(_cell_text(dt_target)):
                _insert_sig_and_date_in_empty_cell(dt_target, None, date_png)
                date_placed = True
            elif _insert_date_only_in_date_cell(date_cell, date_png):
                date_placed = True
        elif signoff_slot and row_blanks:
            # 签批栏里若存在显式「日期」列，优先把日期落到其右侧/同格，避免错位到签名格。
            dt_target = None
            if date_cell is not None:
                try:
                    dt_target = _next_distinct_cell(cells, cells.index(date_cell))
                except Exception:
                    dt_target = None
            if dt_target is not None and _is_slot_target_text(_cell_text(dt_target)):
                _insert_sig_and_date_in_empty_cell(dt_target, None, date_png)
                date_placed = True
            elif date_cell is not None and _insert_date_only_in_date_cell(date_cell, date_png):
                date_placed = True
            elif sig_png and len(row_blanks) > 1:
                _insert_sig_and_date_in_empty_cell(row_blanks[1], None, date_png)
                date_placed = True
            elif sig_png and row_blanks:
                _append_date_to_cell(row_blanks[0], date_png)
                date_placed = True
            elif row_blanks:
                _insert_sig_and_date_in_empty_cell(row_blanks[0], None, date_png)
                date_placed = True
        elif signoff_slot and cell_has_signoff_inline_reservation(ct, keyword):
            for p0 in label_cell.paragraphs:
                if _try_paragraph_inline(p0, keyword, sig_png, date_png):
                    return True
        elif not sig_png:
            for p0 in label_cell.paragraphs:
                off = _find_keyword_in_paragraph(p0, keyword)
                if off >= 0:
                    _clear_runs_after_offset(p0, off)
                    _insert_pictures_in_paragraph(p0, None, date_png)
                    date_placed = True
                    break
        if date_placed:
            placed_any = True

    if (sig_png or date_png) and not placed_any:
        sig_blank = _first_reserved_blank_cell_same_row(cells, label_ci, keyword)
        if sig_blank is not None:
            _insert_sig_and_date_in_empty_cell(sig_blank, sig_png, date_png)
            placed_any = True

    return placed_any


def _try_table_four_column_scan(
    table: Table,
    keyword: str,
    sig_png: Optional[bytes],
    date_png: Optional[bytes],
) -> bool:
    rows_list = list(table.rows)
    for ri in _table_row_scan_order(table):
        cells = rows_list[ri].cells
        for ci in range(max(0, len(cells) - 3)):
            if _place_four_column_row(cells, ci, keyword, sig_png, date_png):
                return True
    return False


def _try_table_role(
    table: Table,
    keyword: str,
    sig_png: Optional[bytes],
    date_png: Optional[bytes],
    role_id: str = "",
) -> bool:
    if _try_table_four_column_scan(table, keyword, sig_png, date_png):
        return True
    if _try_table_adjacent_date_column(table, keyword, sig_png, date_png):
        return True
    rows_list = list(table.rows)
    for ri in _table_row_scan_order(table):
        row = rows_list[ri]
        cells = row.cells
        for ci, cell in enumerate(cells):
            ct = _cell_text(cell)
            if not _cell_has_table_signoff_reservation(ct, keyword, cells, ci, rows_list, ri):
                continue
            if _place_sig_date_at_signoff_anchor(
                cell, keyword, cells, ci, rows_list, ri, sig_png, date_png
            ):
                return True
    return False


def _iter_tables(doc: DocumentObject) -> list[Table]:
    out = []
    try:
        for t in doc.tables:
            out.append(t)
    except Exception:
        pass
    return out


def _iter_all_tables(doc: DocumentObject, *, footers_first: bool = False) -> list[Table]:
    """正文表 + 页眉页脚表（用例表签批栏常在每页页脚）。"""
    body: list[Table] = []
    hf: list[Table] = []
    try:
        body = list(doc.tables)
    except Exception:
        pass
    try:
        for section in doc.sections:
            for part in (
                section.footer,
                section.first_page_footer,
                section.header,
                section.first_page_header,
            ):
                if part is None:
                    continue
                hf.extend(list(getattr(part, "tables", None) or []))
    except Exception:
        pass
    if footers_first:
        return hf + body
    return body + hf


def _iter_body_paragraphs(doc: DocumentObject) -> list[Paragraph]:
    return list(doc.paragraphs)


def _planned_keywords_for_role(
    role_id: str,
    placement_plan: dict | None,
) -> list[str]:
    if not isinstance(placement_plan, dict):
        return []
    role_plan = placement_plan.get(str(role_id))
    if not isinstance(role_plan, dict):
        return []
    kws = role_plan.get("keywords")
    if not isinstance(kws, list):
        return []
    out: list[str] = []
    seen: set[str] = set()
    for item in kws:
        s = str(item or "").strip()
        if not s or s in seen:
            continue
        seen.add(s)
        out.append(s)
    out.sort(key=len, reverse=True)
    return out


def _planned_source_hints_for_role(
    role_id: str,
    placement_plan: dict | None,
) -> list[str]:
    if not isinstance(placement_plan, dict):
        return []
    role_plan = placement_plan.get(str(role_id))
    if not isinstance(role_plan, dict):
        return []
    arr = role_plan.get("source_hints")
    if not isinstance(arr, list):
        return []
    out: list[str] = []
    seen: set[str] = set()
    for item in arr:
        s = str(item or "").strip()
        if not s or s in seen:
            continue
        seen.add(s)
        out.append(s)
    return out


def _planned_layout_types_for_role(
    role_id: str,
    placement_plan: dict | None,
) -> list[str]:
    if not isinstance(placement_plan, dict):
        return []
    role_plan = placement_plan.get(str(role_id))
    if not isinstance(role_plan, dict):
        return []
    arr = role_plan.get("layout_types")
    if not isinstance(arr, list):
        return []
    out: list[str] = []
    seen: set[str] = set()
    for item in arr:
        s = str(item or "").strip()
        if not s or s in seen:
            continue
        seen.add(s)
        out.append(s)
    return out


def _parse_docx_table_row_hint(source_hint: str) -> tuple[int, int] | None:
    m = re.match(r"^\s*table(\d+)\.row(\d+)\s*$", str(source_hint or ""), re.IGNORECASE)
    if not m:
        return None
    return int(m.group(1)), int(m.group(2))


def _parse_docx_paragraph_hint(source_hint: str) -> int | None:
    m = re.match(r"^\s*paragraph(\d+)\s*$", str(source_hint or ""), re.IGNORECASE)
    if not m:
        return None
    return int(m.group(1))


def _try_table_role_at_row(
    table: Table,
    keyword: str,
    sig_png: Optional[bytes],
    date_png: Optional[bytes],
    row_no_1based: int,
) -> bool:
    rows_list = list(table.rows)
    ri = int(row_no_1based) - 1
    if ri < 0 or ri >= len(rows_list):
        return False
    row = rows_list[ri]
    cells = row.cells
    for ci, cell in enumerate(cells):
        ct = _cell_text(cell)
        if not _cell_has_table_signoff_reservation(ct, keyword, cells, ci, rows_list, ri):
            continue
        if _place_sig_date_at_signoff_anchor(
            cell, keyword, cells, ci, rows_list, ri, sig_png, date_png
        ):
            return True
    return False


def _try_docx_plan_placement(
    doc: DocumentObject,
    keywords: list[str],
    source_hints: list[str],
    sig_png: Optional[bytes],
    date_png: Optional[bytes],
) -> bool:
    if not keywords or not source_hints:
        return False
    try:
        from sign_handlers.detect_fields import _iter_docx_tables_for_detect

        detect_tables = _iter_docx_tables_for_detect(doc, max_body_tables=32, max_total=64)
    except Exception:
        detect_tables = []
    body_paras = _iter_body_paragraphs(doc)

    for hint in source_hints:
        tr = _parse_docx_table_row_hint(hint)
        if tr:
            tno, rno = tr
            candidates: list[Table] = []
            if detect_tables and 1 <= tno <= len(detect_tables):
                candidates.append(detect_tables[tno - 1])
            seen_tables: set[int] = set()
            uniq: list[Table] = []
            for tb in candidates:
                ident = id(tb)
                if ident in seen_tables:
                    continue
                seen_tables.add(ident)
                uniq.append(tb)
            for kw in keywords:
                for tb in uniq:
                    if _try_table_role_at_row(tb, kw, sig_png, date_png, rno):
                        return True
        pno = _parse_docx_paragraph_hint(hint)
        if pno is not None and 1 <= pno <= len(body_paras):
            p = body_paras[pno - 1]
            for kw in keywords:
                if _find_keyword_in_paragraph(p, kw) < 0:
                    continue
                if _try_paragraph_inline(p, kw, sig_png, date_png):
                    return True
    return False


def sign_docx(
    path: str,
    role_to_signature_png: dict,
    role_to_date_png: dict,
    out_path: Optional[str] = None,
    placement_plan: dict | None = None,
    placement_result: dict | None = None,
) -> str:
    path = os.path.abspath(path)
    if out_path is None:
        base, ext = os.path.splitext(path)
        out_path = f"{base}_signed{ext}"
    doc = Document(path)

    sig_map: dict = {}
    date_map: dict = {}
    for rid in ROLE_ID_TO_KEYWORD:
        sb = role_to_signature_png.get(rid)
        db = role_to_date_png.get(rid)
        if sb and db:
            s2, d2 = prepare_signature_date_pair_for_word(sb, db)
            if s2:
                sig_map[rid] = s2
            if d2:
                date_map[rid] = d2
        else:
            if sb:
                sig_map[rid] = prepare_png_for_word(sb) or sb
            if db:
                date_map[rid] = prepare_png_for_word(db) or db

    for role_id in ROLE_ID_TO_KEYWORD:
        sig = sig_map.get(role_id)
        dt = date_map.get(role_id)
        want_date = True
        if isinstance(placement_plan, dict):
            rp0 = placement_plan.get(str(role_id))
            if isinstance(rp0, dict) and "date_slot" in rp0:
                want_date = bool(rp0.get("date_slot"))
        dt_use = dt if want_date else None
        role_result = None
        if isinstance(placement_result, dict):
            role_result = placement_result.setdefault(
                str(role_id),
                {
                    "applied_sig": bool(sig),
                    "applied_date": bool(dt),
                    "placed": False,
                    "placed_by": "",
                    "keywords": [],
                },
            )
        if not sig and not dt_use:
            if isinstance(role_result, dict):
                role_result["placed"] = False
                role_result["placed_by"] = "skip_no_material"
            continue
        done = False
        fallback_kws = sorted(role_keywords_for_apply(role_id), key=lambda x: len(x), reverse=True)
        planned_kws = _planned_keywords_for_role(role_id, placement_plan)
        planned_hints = _planned_source_hints_for_role(role_id, placement_plan)
        planned_layout_types = _planned_layout_types_for_role(role_id, placement_plan)
        kws = []
        seen_kw: set[str] = set()
        for kw in planned_kws + fallback_kws:
            if kw in seen_kw:
                continue
            seen_kw.add(kw)
            kws.append(kw)
        if isinstance(role_result, dict):
            role_result["keywords"] = kws[:]
            role_result["source_hints"] = planned_hints[:]
            role_result["layout_types"] = planned_layout_types[:]
        placed_by = ""
        tables = _iter_all_tables(doc, footers_first=True)
        if (not done) and placement_plan and isinstance(placement_plan, dict):
            done = _try_docx_role_layout_cells(
                doc, role_id, placement_plan, kws, sig, dt_use
            )
            if done:
                placed_by = "layout_cells"
        if ("two_row_signoff_table" in planned_layout_types) and kws:
            for kw in kws:
                if done:
                    break
                for table in tables:
                    if _try_table_role(table, kw, sig, dt_use, role_id=role_id):
                        done = True
                        placed_by = "planned_layout_type"
                        break
        if (not done) and planned_hints and kws:
            done = _try_docx_plan_placement(doc, kws, planned_hints, sig, dt_use)
            if done:
                placed_by = "planned_source_hint"
        if not done:
            for kw in kws:
                if done:
                    break
                for table in tables:
                    if _try_table_role(table, kw, sig, dt_use, role_id=role_id):
                        done = True
                        if not placed_by:
                            placed_by = (
                                "planned_keywords" if planned_kws else "fallback_keywords"
                            )
                        break
                if done:
                    break
                for p in _iter_body_paragraphs(doc):
                    if _find_keyword_in_paragraph(p, kw) < 0:
                        continue
                    if _try_paragraph_inline(p, kw, sig, dt_use):
                        done = True
                        if not placed_by:
                            placed_by = (
                                "planned_keywords" if planned_kws else "fallback_keywords"
                            )
                        break
                if done:
                    break
                for table in tables:
                    rows_list = list(table.rows)
                    for ri, row in enumerate(rows_list):
                        row_cells = list(row.cells)
                        for ci, cell in enumerate(row_cells):
                            ct = _cell_text(cell)
                            if not _cell_has_table_signoff_reservation(
                                ct, kw, row_cells, ci, rows_list, ri
                            ):
                                continue
                            if _place_sig_date_at_signoff_anchor(
                                cell, kw, row_cells, ci, rows_list, ri, sig, dt_use
                            ):
                                done = True
                                if not placed_by:
                                    placed_by = (
                                        "planned_keywords" if planned_kws else "fallback_keywords"
                                    )
                                break
                        if done:
                            break
                    if done:
                        break

        if isinstance(role_result, dict):
            role_result["placed"] = bool(done)
            if done:
                role_result["placed_by"] = placed_by or "fallback_keywords"
            else:
                role_result["placed_by"] = "not_found"

    doc.save(out_path)
    return out_path
