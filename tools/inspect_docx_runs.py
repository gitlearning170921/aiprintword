# -*- coding: utf-8 -*-
from __future__ import annotations

import sys
import re
from docx import Document


def main() -> int:
    if len(sys.argv) < 2:
        print("Usage: python tools/inspect_docx_runs.py <path.docx>")
        return 2
    path = sys.argv[1]
    doc = Document(path)
    kw_re = re.compile(
        r"\b(Author|Reviewer|Approver)\b\s*[:：]?|\bDate\b\s*[:：]?|日期\s*[:：]?",
        re.I,
    )

    def dump_par(p, where: str) -> None:
        txt = p.text or ""
        if not kw_re.search(txt):
            return
        print("\n===", where, "===")
        print("PARA:", repr(txt))
        for i, r in enumerate(p.runs):
            rt = r.text or ""
            try:
                u = bool(r.font.underline)
            except Exception:
                u = None
            # python-docx 的底层 xpath 支持有限：用 local-name() 判断是否包含图片节点
            try:
                has_pic = bool(r._element.xpath(".//*[local-name()='pic']"))
            except Exception:
                has_pic = False
            pic_flag = " PIC" if has_pic else ""
            print(f"  run[{i}] u={u} text={rt!r}{pic_flag}")

    for i, p in enumerate(doc.paragraphs):
        dump_par(p, f"body p#{i}")

    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    dump_par(p, f"table{ti} r{ri}c{ci} p{pi}")

    print("\nDone")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

