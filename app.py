# -*- coding: utf-8 -*-
"""
Batch print web service for Word/Excel/PDF.
"""
import base64
import io
import json
import logging
import os
import socket
import re
import shutil
import subprocess
import sys
import tempfile
import threading
import time
import uuid
import zipfile
import functools
from datetime import datetime
from contextlib import contextmanager
from pathlib import Path
from queue import Queue, Empty
from dataclasses import dataclass, field
from typing import Optional

from flask import (
    Flask,
    request,
    jsonify,
    send_from_directory,
    send_file,
    Response,
    session,
)
from werkzeug.utils import secure_filename

LOG_LEVEL = (os.environ.get("AIPRINTWORD_LOG_LEVEL") or "INFO").upper()
if not logging.getLogger().handlers:
    logging.basicConfig(
        level=getattr(logging, LOG_LEVEL, logging.INFO),
        format="%(asctime)s %(levelname)s %(name)s: %(message)s",
    )
logger = logging.getLogger("aiprintword.web")


def _is_path_under_parent(parent_abs: str, path_abs: str) -> bool:
    """path_abs 是否为 parent_abs 本身或其子路径（大小写按平台归一）。"""
    parent = os.path.normcase(os.path.abspath(parent_abs))
    child = os.path.normcase(os.path.abspath(path_abs))
    if parent == child:
        return True
    if not parent.endswith(os.sep):
        parent = parent + os.sep
    return child.startswith(parent)


def _parse_incremental_output_dir(form):
    """
    处理完成后立即落盘的根目录（运行 Flask 的机器上的绝对路径）。
    若设置环境变量 AIPRINTWORD_ALLOWED_OUTPUT_PARENT，则输出目录必须位于该路径之下。
    """
    raw = (form.get("incremental_output_dir") or "").strip()
    if not raw:
        return None
    path = os.path.abspath(os.path.normpath(raw))
    try:
        from runtime_settings.resolve import get_setting

        allowed_raw = (str(get_setting("AIPRINTWORD_ALLOWED_OUTPUT_PARENT") or "")).strip()
    except Exception:
        allowed_raw = (os.environ.get("AIPRINTWORD_ALLOWED_OUTPUT_PARENT") or "").strip()
    if allowed_raw:
        allowed_abs = os.path.abspath(os.path.normpath(allowed_raw))
        if not _is_path_under_parent(allowed_abs, path):
            raise ValueError(
                "输出目录必须位于环境变量 AIPRINTWORD_ALLOWED_OUTPUT_PARENT 所允许的目录之下。"
                f" 当前允许根：{allowed_abs}"
            )
    try:
        os.makedirs(path, exist_ok=True)
    except OSError as e:
        raise ValueError(f"无法创建或使用输出目录：{path}（{e}）")
    if not os.path.isdir(path):
        raise ValueError(f"输出路径不是目录：{path}")
    return path


@dataclass
class _BatchJob:
    job_id: str
    cancel_event: threading.Event
    pause_event: threading.Event
    skip_current_event: threading.Event = field(default_factory=threading.Event)


_BATCH_JOB_REGISTRY = {}

# 本轮终版下载映射（local）：key -> {created_at, items:[{ftp_path, filename}]}
_FINAL_DOWNLOAD_CACHE = {}

# Allowed file extensions（批处理实际处理的文档类型）
ALLOWED_EXT = {
    ".doc", ".docx", ".docm",
    ".xls", ".xlsx", ".xlsm",
    ".pdf",
}
# 可上传的压缩包：服务端解压后仅将 ALLOWED_EXT 内的成员加入批处理
ARCHIVE_UPLOAD_EXT = {".zip", ".7z", ".rar"}
# 解压防护（单个压缩包）
_ARCHIVE_EXTRACT_MAX_FILES = 500
_ARCHIVE_EXTRACT_MAX_TOTAL_UNCOMPRESSED = 250 * 1024 * 1024
_ARCHIVE_EXTRACT_SINGLE_MAX = 80 * 1024 * 1024

app = Flask(__name__, static_folder="static", static_url_path="")
app.config["MAX_CONTENT_LENGTH"] = 256 * 1024 * 1024  # 256MB total upload limit
app.secret_key = os.environ.get("FLASK_SECRET_KEY") or "aiprintword-dev-secret-change-with-env"

# Project root
ROOT = os.path.dirname(os.path.abspath(__file__))
# 与 /api/aiprintword-build 中 build 字段一致；用于确认 5050 是否加载了当前这份 app.py
AIPRINTWORD_WEB_BUILD = 8
BATCH_EXPORT_ROOT = os.path.join(ROOT, "data", "batch_exports")
HANDOFF_DIR = os.path.join(ROOT, "data", "aiword_handoff")
_HANDOFF_LOCK = threading.Lock()
_HEAVY_SEM: Optional[threading.Semaphore] = None
_HEAVY_SEM_SIZE = 0
_DETECT_SEM: Optional[threading.Semaphore] = None
_DETECT_SEM_SIZE = 0


class _SignHeavyBusyError(Exception):
    """上传/签名等重任务并发槽位已满。"""


class _SignDetectBusyError(Exception):
    """文档识别并发槽位已满。"""


def _sign_heavy_semaphore() -> threading.Semaphore:
    global _HEAVY_SEM, _HEAVY_SEM_SIZE
    if _HEAVY_SEM is None:
        try:
            n = int((os.environ.get("SIGN_HEAVY_CONCURRENCY") or "3").strip() or "3")
        except ValueError:
            n = 3
        _HEAVY_SEM_SIZE = max(1, min(n, 16))
        _HEAVY_SEM = threading.Semaphore(_HEAVY_SEM_SIZE)
    return _HEAVY_SEM


@contextmanager
def _sign_heavy_op_slot():
    try:
        wait = int((os.environ.get("SIGN_HEAVY_WAIT_SEC") or "90").strip() or "90")
    except ValueError:
        wait = 90
    wait = max(5, min(wait, 600))
    sem = _sign_heavy_semaphore()
    if not sem.acquire(timeout=wait):
        raise _SignHeavyBusyError(
            "服务器正忙于其它上传/批量签名任务（同时最多 "
            + str(_HEAVY_SEM_SIZE)
            + " 个重任务）。请稍后再试，或让另一客户端先完成当前批次。"
        )
    try:
        yield
    finally:
        sem.release()


def sign_heavy_route(fn):
    """限制并发的重路由：素材上传、保存笔迹、生成/批量签名文档。"""

    @functools.wraps(fn)
    def _wrapped(*args, **kwargs):
        try:
            with _sign_heavy_op_slot():
                return fn(*args, **kwargs)
        except _SignHeavyBusyError as e:
            return jsonify({"ok": False, "error": str(e), "error_code": "server_busy"}), 503

    return _wrapped


def _sign_detect_semaphore() -> threading.Semaphore:
    global _DETECT_SEM, _DETECT_SEM_SIZE
    if _DETECT_SEM is None:
        try:
            n = int((os.environ.get("SIGN_DETECT_CONCURRENCY") or "2").strip() or "2")
        except ValueError:
            n = 2
        _DETECT_SEM_SIZE = max(1, min(n, 8))
        _DETECT_SEM = threading.Semaphore(_DETECT_SEM_SIZE)
    return _DETECT_SEM


@contextmanager
def _sign_detect_op_slot():
    try:
        wait = int((os.environ.get("SIGN_DETECT_WAIT_SEC") or "45").strip() or "45")
    except ValueError:
        wait = 45
    wait = max(3, min(wait, 300))
    sem = _sign_detect_semaphore()
    if not sem.acquire(timeout=wait):
        raise _SignDetectBusyError(
            "服务器识别任务已满（同时最多 "
            + str(_DETECT_SEM_SIZE)
            + " 个文档在识别）。请稍候自动重试，或等当前批次完成后再操作。"
        )
    try:
        yield
    finally:
        sem.release()


def sign_detect_route(fn):
    """限制并发的识别路由，避免多轮上传/识别占满 Flask 工作线程。"""

    @functools.wraps(fn)
    def _wrapped(*args, **kwargs):
        try:
            with _sign_detect_op_slot():
                return fn(*args, **kwargs)
        except _SignDetectBusyError as e:
            return jsonify({"ok": False, "error": str(e), "error_code": "detect_busy"}), 503

    return _wrapped


_HANDOFF_TTL_SEC = 30 * 60
_HANDOFF_MAX_BYTES = 80 * 1024 * 1024
BATCH_HISTORY_ROOT = os.path.join(ROOT, "data", "batch_history")
_BATCH_EXPORT_TOKEN_RE = re.compile(r"^[0-9a-f]{32}$")
# 可选：单行口令文件（与 .env 二选一即可，适合服务账号无 .env 拷贝权限时）
_ADMIN_TOKEN_FILE_DEFAULT = os.path.join(ROOT, "data", "admin_token.txt")
_ENV_FILE_ENCODINGS = ("utf-8-sig", "utf-8", "gbk")


def _resolved_dotenv_path() -> str:
    """项目 .env 绝对路径；可通过环境变量 AIPRINTWORD_DOTENV_PATH 指定（服务的工作目录与代码目录不一致时）。"""
    o = (os.environ.get("AIPRINTWORD_DOTENV_PATH") or "").strip().strip('"').strip("'")
    if o:
        return os.path.abspath(os.path.expanduser(o))
    return os.path.join(ROOT, ".env")


def _read_key_from_env_file(path: str, key: str) -> str:
    """仅从文件读取键值（不写 os.environ）；支持 KEY= 与 KEY =；多编码尝试。"""
    if not os.path.isfile(path):
        return ""
    key_norm = (key or "").strip()
    if not key_norm:
        return ""
    for enc in _ENV_FILE_ENCODINGS:
        try:
            with open(path, "r", encoding=enc) as f:
                for raw in f:
                    line = raw.strip()
                    if not line or line.startswith("#"):
                        continue
                    if "=" not in line:
                        continue
                    k, _, val = line.partition("=")
                    if k.strip() != key_norm:
                        continue
                    val = val.strip()
                    if len(val) >= 2 and val[0] == val[-1] and val[0] in "\"'":
                        val = val[1:-1]
                    return (val or "").strip()
            return ""
        except UnicodeDecodeError:
            continue
        except OSError:
            return ""
    return ""


def _parse_env_file_value(path: str, key: str) -> None:
    """从 .env 逐行解析键值并写入 os.environ（支持 KEY=val 与 KEY = val；兜底 dotenv 未生效）。"""
    val = _read_key_from_env_file(path, key)
    if val:
        os.environ[(key or "").strip()] = val


def _load_project_dotenv() -> None:
    """加载项目根目录 .env；保证管理口令等关键变量进入 os.environ。"""
    path = _resolved_dotenv_path()
    try:
        from dotenv import load_dotenv

        try:
            load_dotenv(path, override=True, encoding="utf-8-sig")
        except TypeError:
            load_dotenv(path, override=True)
    except ImportError:
        pass
    if not (os.environ.get("AIPRINTWORD_ADMIN_TOKEN") or "").strip():
        _parse_env_file_value(path, "AIPRINTWORD_ADMIN_TOKEN")


_load_project_dotenv()
logger.info(
    ".env 路径=%s 存在=%s AIPRINTWORD_ADMIN_TOKEN=%s",
    _resolved_dotenv_path(),
    os.path.isfile(_resolved_dotenv_path()),
    "已配置" if (os.environ.get("AIPRINTWORD_ADMIN_TOKEN") or "").strip() else "未配置",
)


def _sync_admin_token_from_disk() -> None:
    """
    每次管理 API 调用前从磁盘同步口令（不依赖仅启动时 load 一次）。
    解决：Flask 重载子进程、多实例、以及 .env 写成「KEY = 值」时旧版手动解析失败等。
    """
    path = _resolved_dotenv_path()
    try:
        from dotenv import load_dotenv

        try:
            load_dotenv(path, override=True, encoding="utf-8-sig")
        except TypeError:
            load_dotenv(path, override=True)
    except Exception:
        pass
    _parse_env_file_value(path, "AIPRINTWORD_ADMIN_TOKEN")
    if (os.environ.get("AIPRINTWORD_ADMIN_TOKEN") or "").strip():
        return
    token_paths = []
    ext = (os.environ.get("AIPRINTWORD_ADMIN_TOKEN_FILE") or "").strip()
    if ext:
        token_paths.append(ext)
    token_paths.append(_ADMIN_TOKEN_FILE_DEFAULT)
    for fp in token_paths:
        if not fp or not os.path.isfile(fp):
            continue
        try:
            with open(fp, encoding="utf-8-sig") as f:
                for raw in f:
                    line = raw.strip()
                    if not line or line.startswith("#"):
                        continue
                    os.environ["AIPRINTWORD_ADMIN_TOKEN"] = line
                    break
        except OSError:
            continue
        if (os.environ.get("AIPRINTWORD_ADMIN_TOKEN") or "").strip():
            break


def _batch_history_max() -> int:
    try:
        from runtime_settings.resolve import get_setting

        v = int(get_setting("AIPRINTWORD_HISTORY_MAX"))
        return max(5, min(v, 500))
    except Exception:
        return 50


def _sign_mysql_max_files() -> int:
    try:
        from runtime_settings.resolve import get_setting

        v = int(get_setting("SIGN_MYSQL_MAX_FILES"))
        return max(1, min(v, 10000))
    except Exception:
        return 500


def _sign_mysql_max_signed() -> int:
    try:
        from runtime_settings.resolve import get_setting

        v = int(get_setting("SIGN_MYSQL_MAX_SIGNED"))
        return max(1, min(v, 50000))
    except Exception:
        return 2000


# Common COM error codes
_COM_ERROR_CODES = (-2147352567, -2147467259, -2146827864)


def _format_com_error(e):
    """Format COM exceptions to user-friendly text."""
    try:
        s = str(e)
        if "?" in s and "?" in s:
            return s
    except Exception:
        pass
    try:
        if getattr(e, "args", None) and len(e.args) > 0 and e.args[0] in _COM_ERROR_CODES:
            return "????????????? WPS/Excel ?????????????????"
    except Exception:
        pass
    return str(e)


def _allowed_file(filename):
    return os.path.splitext(filename)[1].lower() in ALLOWED_EXT


def _is_archive_upload(filename: str) -> bool:
    return os.path.splitext(filename)[1].lower() in ARCHIVE_UPLOAD_EXT


def _zip_strip_leading_windows_drive(norm: str) -> str:
    """去掉 ZIP 内常见的「盘符绝对路径」前缀，如 C:/xxx、D:\\xxx（否则首段 C: 含冒号会被安全校验整包拒绝）。"""
    s = norm.replace("\\", "/").strip()
    while len(s) >= 3 and s[0].isalpha() and s[1] == ":" and s[2] == "/":
        s = s[3:].lstrip("/")
    return s


def _zip_inner_path_safe(name: str) -> Optional[str]:
    """ZIP 内相对路径规范化；拒绝路径穿越与非法段。"""
    if not name or str(name).endswith("/"):
        return None
    norm = _zip_strip_leading_windows_drive(str(name))
    if not norm:
        return None
    parts = [p for p in norm.split("/") if p]
    if not parts or ".." in parts:
        return None
    if parts[0] in ("/", "\\"):
        return None
    for p in parts:
        for ch in '<>:"|?*':
            if ch in p:
                return None
    return "/".join(parts)


def _expand_zipfile_members(zf: zipfile.ZipFile, extract_root: str, zip_rel: str):
    """从已打开的 ZipFile 中解压 ALLOWED_EXT 成员到 extract_root；返回 [(绝对路径, 展示相对名), ...]。"""
    zip_rel = zip_rel.replace("\\", "/").strip()
    out = []
    abs_root = os.path.abspath(extract_root)
    if not abs_root.endswith(os.sep):
        abs_root_prefix = abs_root + os.sep
    else:
        abs_root_prefix = abs_root
    total_uncompressed = 0
    infos = zf.infolist()
    for info in infos:
        if info.is_dir():
            continue
        name_candidates = [info.filename]
        # 3.11+ 可用 metadata_encoding；旧版常见：未设 UTF-8 位时中文名被当成 cp437，需 gbk 再解一次
        if sys.version_info < (3, 11) and not (getattr(info, "flag_bits", 0) & 0x800):
            try:
                alt = info.filename.encode("cp437").decode("gbk")
                if alt and alt not in name_candidates:
                    name_candidates.append(alt)
            except (UnicodeEncodeError, UnicodeDecodeError, LookupError):
                pass
        inner = None
        for cand in name_candidates:
            inner = _zip_inner_path_safe(cand)
            if inner is not None:
                break
        if inner is None:
            continue
        ext = os.path.splitext(inner)[1].lower()
        if ext not in ALLOWED_EXT:
            continue
        try:
            sz = int(getattr(info, "file_size", 0) or 0)
        except (TypeError, ValueError):
            sz = 0
        if sz > _ARCHIVE_EXTRACT_SINGLE_MAX:
            continue
        if total_uncompressed + sz > _ARCHIVE_EXTRACT_MAX_TOTAL_UNCOMPRESSED:
            logger.warning("zip uncompressed total cap reached rel=%s", zip_rel)
            break
        dest = os.path.join(extract_root, *inner.split("/"))
        abs_dest = os.path.abspath(dest)
        if not (abs_dest == abs_root or abs_dest.startswith(abs_root_prefix)):
            continue
        try:
            os.makedirs(os.path.dirname(dest), exist_ok=True)
        except OSError:
            continue
        try:
            with zf.open(info, "r") as src, open(dest, "wb") as dst:
                shutil.copyfileobj(src, dst, length=65536)
        except Exception as e:
            logger.warning("zip extract member failed %s in %s: %s", inner, zip_rel, e)
            try:
                if os.path.isfile(dest):
                    os.remove(dest)
            except OSError:
                pass
            continue
        total_uncompressed += sz
        rel_out = zip_rel.rstrip("/") + "/" + inner
        out.append((dest, rel_out))
        if len(out) >= _ARCHIVE_EXTRACT_MAX_FILES:
            logger.warning("zip extract file count cap rel=%s", zip_rel)
            break
    if not out:
        n_nd = sum(1 for i in infos if not i.is_dir())
        if n_nd:
            exts = []
            seen = set()
            for i in infos:
                if i.is_dir():
                    continue
                raw = _zip_strip_leading_windows_drive(str(i.filename))
                base = raw.split("/")[-1] if raw else ""
                ext = os.path.splitext(base)[1].lower() or "(无扩展名)"
                if ext not in seen:
                    seen.add(ext)
                    exts.append(ext)
                if len(exts) >= 12:
                    break
            logger.warning(
                "zip rel=%s: 包内约 %d 个文件项，但未加入批处理（扩展名须为 doc/docx/xls/xlsx/pdf 等；或为路径/解压失败）。样例扩展名: %s",
                zip_rel,
                n_nd,
                ", ".join(exts) if exts else "—",
            )
    return out


def _expand_zip_for_batch(tmp_dir: str, zip_path: str, zip_rel: str):
    """
    解压 zip，将其中符合 ALLOWED_EXT 的文件加入批处理。
    返回 [(绝对路径, 展示用相对名「压缩包名/包内路径」), ...]
    """
    zip_rel = zip_rel.replace("\\", "/").strip()

    def _open_and_expand(metadata_encoding=None):
        """返回 (能否打开 zip 结构, 解压出的文档列表)。"""
        extract_root = os.path.join(tmp_dir, "_unzip", uuid.uuid4().hex[:16])
        try:
            os.makedirs(extract_root, exist_ok=True)
        except OSError:
            return False, []
        try:
            kwargs = {}
            if metadata_encoding is not None and sys.version_info >= (3, 11):
                kwargs["metadata_encoding"] = metadata_encoding
            with zipfile.ZipFile(zip_path, "r", **kwargs) as zf:
                return True, _expand_zipfile_members(zf, extract_root, zip_rel)
        except (zipfile.BadZipFile, OSError) as e:
            if metadata_encoding is None:
                logger.warning("zip open failed rel=%s err=%s", zip_rel, e)
            return False, []

    ok, out = _open_and_expand(None)
    if not ok:
        return []
    if out:
        return out
    if sys.version_info >= (3, 11):
        for enc in ("gbk", "utf-8", "cp437"):
            ok2, out2 = _open_and_expand(enc)
            if ok2 and out2:
                logger.info("zip rel=%s used metadata_encoding=%s", zip_rel, enc)
                return out2
    return []


def _collect_allowed_from_extract_dir(extract_root: str, archive_rel: str):
    """
    遍历已解压目录，收集 ALLOWED_EXT 文件；返回 [(绝对路径, archive_rel/相对路径), ...]
    """
    archive_rel = archive_rel.replace("\\", "/").strip()
    out = []
    try:
        abs_er = os.path.abspath(extract_root)
    except OSError:
        return []
    if not os.path.isdir(abs_er):
        return []
    prefix = abs_er if abs_er.endswith(os.sep) else abs_er + os.sep
    total_bytes = 0
    for root, _, files in os.walk(abs_er):
        for fn in sorted(files):
            if len(out) >= _ARCHIVE_EXTRACT_MAX_FILES:
                return out
            full = os.path.join(root, fn)
            try:
                st = os.stat(full)
            except OSError:
                continue
            if st.st_size > _ARCHIVE_EXTRACT_SINGLE_MAX:
                continue
            ext = os.path.splitext(fn)[1].lower()
            if ext not in ALLOWED_EXT:
                continue
            try:
                rel_inner = os.path.relpath(full, abs_er).replace("\\", "/")
            except ValueError:
                continue
            if ".." in rel_inner.split("/"):
                continue
            abs_full = os.path.abspath(full)
            if not (abs_full == abs_er or abs_full.startswith(prefix)):
                continue
            if total_bytes + st.st_size > _ARCHIVE_EXTRACT_MAX_TOTAL_UNCOMPRESSED:
                logger.warning("archive walk total cap archive=%s", archive_rel)
                return out
            total_bytes += st.st_size
            rel_out = archive_rel.rstrip("/") + "/" + rel_inner
            out.append((full, rel_out))
    return out


def _resolve_7zip_cli() -> Optional[str]:
    for key in ("SEVEN_ZIP_EXE", "7ZIP_EXE", "7Z_EXE"):
        p = (os.environ.get(key) or "").strip().strip('"').strip("'")
        if p and os.path.isfile(p):
            return p
    for w in (shutil.which("7z"), shutil.which("7z.exe")):
        if w and os.path.isfile(w):
            return w
    for p in (
        r"C:\Program Files\7-Zip\7z.exe",
        r"C:\Program Files (x86)\7-Zip\7z.exe",
    ):
        if os.path.isfile(p):
            return p
    return None


def _expand_via_7zip_cli(tmp_dir: str, archive_path: str, archive_rel: str):
    """使用 7-Zip 命令行解压（支持 .7z / .rar / .zip 等）。"""
    exe = _resolve_7zip_cli()
    if not exe:
        return []
    extract_root = os.path.join(tmp_dir, "_7x", uuid.uuid4().hex[:16])
    try:
        os.makedirs(extract_root, exist_ok=True)
    except OSError:
        return []
    cmd = [exe, "x", os.path.abspath(archive_path), f"-o{extract_root}", "-aoa", "-y", "-bb0"]
    try:
        kwargs = {
            "capture_output": True,
            "text": True,
            "errors": "replace",
            "timeout": 600,
        }
        if hasattr(subprocess, "CREATE_NO_WINDOW"):
            kwargs["creationflags"] = subprocess.CREATE_NO_WINDOW
        r = subprocess.run(cmd, **kwargs)
    except (OSError, subprocess.TimeoutExpired) as e:
        logger.warning("7z cli failed %s: %s", archive_rel, e)
        return []
    if r.returncode != 0:
        logger.warning(
            "7z exit %s archive=%s err=%s",
            r.returncode,
            archive_rel,
            (r.stderr or r.stdout or "")[:500],
        )
        return []
    return _collect_allowed_from_extract_dir(extract_root, archive_rel)


def _expand_7z_py7zr(tmp_dir: str, archive_path: str, archive_rel: str):
    """无 7z.exe 时，用 py7zr 解压 .7z（纯 Python）。"""
    try:
        import py7zr
    except ImportError:
        return []
    extract_root = os.path.join(tmp_dir, "_p7z", uuid.uuid4().hex[:16])

    def _do_extract() -> None:
        os.makedirs(extract_root, exist_ok=True)
        with py7zr.SevenZipFile(archive_path, mode="r") as z:
            z.extractall(path=extract_root)

    try:
        from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutTimeout

        with ThreadPoolExecutor(max_workers=1) as ex:
            fut = ex.submit(_do_extract)
            fut.result(timeout=600)
    except FutTimeout:
        logger.warning("py7zr extract timeout archive=%s", archive_rel)
        return []
    except Exception as e:
        logger.warning("py7zr extract failed %s: %s", archive_rel, e)
        return []
    return _collect_allowed_from_extract_dir(extract_root, archive_rel)


def _resolve_unrar_cli() -> Optional[str]:
    for key in ("UNRAR_EXE", "UNRAR"):
        p = (os.environ.get(key) or "").strip().strip('"').strip("'")
        if p and os.path.isfile(p):
            return p
    for w in (shutil.which("unrar"), shutil.which("UnRAR.exe"), shutil.which("unrar.exe")):
        if w and os.path.isfile(w):
            return w
    for p in (
        os.path.join(os.environ.get("ProgramFiles", r"C:\Program Files"), "WinRAR", "UnRAR.exe"),
        r"C:\Program Files\WinRAR\UnRAR.exe",
        r"C:\Program Files (x86)\WinRAR\UnRAR.exe",
    ):
        if p and os.path.isfile(p):
            return p
    return None


def _expand_rar_unrar(tmp_dir: str, archive_path: str, archive_rel: str):
    """无 7z 时使用官方 UnRAR 解压 .rar。"""
    unrar = _resolve_unrar_cli()
    if not unrar:
        return []
    extract_root = os.path.join(tmp_dir, "_rar", uuid.uuid4().hex[:16])
    try:
        os.makedirs(extract_root, exist_ok=True)
    except OSError:
        return []
    dest = extract_root
    if not dest.endswith(os.sep):
        dest = dest + os.sep
    cmd = [unrar, "x", "-y", "-o+", os.path.abspath(archive_path), dest]
    try:
        kwargs = {
            "capture_output": True,
            "text": True,
            "errors": "replace",
            "timeout": 600,
        }
        if hasattr(subprocess, "CREATE_NO_WINDOW"):
            kwargs["creationflags"] = subprocess.CREATE_NO_WINDOW
        r = subprocess.run(cmd, **kwargs)
    except (OSError, subprocess.TimeoutExpired) as e:
        logger.warning("unrar failed %s: %s", archive_rel, e)
        return []
    if r.returncode != 0:
        logger.warning(
            "unrar exit %s archive=%s err=%s",
            r.returncode,
            archive_rel,
            (r.stderr or r.stdout or "")[:500],
        )
        return []
    return _collect_allowed_from_extract_dir(extract_root, archive_rel)


def _expand_nonzip_archive(tmp_dir: str, archive_path: str, archive_rel: str):
    """解压 .7z / .rar：优先 7-Zip 命令行，其次 py7zr（仅 7z）或 UnRAR（仅 rar）。"""
    ext = os.path.splitext(archive_rel)[1].lower()
    out = _expand_via_7zip_cli(tmp_dir, archive_path, archive_rel)
    if out:
        return out
    if ext == ".7z":
        return _expand_7z_py7zr(tmp_dir, archive_path, archive_rel)
    if ext == ".rar":
        return _expand_rar_unrar(tmp_dir, archive_path, archive_rel)
    return []


def _expand_archive_for_batch(tmp_dir: str, path: str, rel: str):
    """按扩展名解压压缩包并返回待批处理文件列表。"""
    ext = os.path.splitext(rel)[1].lower()
    if ext == ".zip":
        out = _expand_zip_for_batch(tmp_dir, path, rel)
        if out:
            return out
        # 中文 Windows「发送到压缩文件夹」等产生的 ZIP，标准库可能解码不出成员名；用 7-Zip 再试
        out7 = _expand_via_7zip_cli(tmp_dir, path, rel)
        if out7:
            logger.info("zip rel=%s expanded via 7-Zip CLI (stdlib found 0 documents)", rel)
        return out7
    if ext in (".7z", ".rar"):
        return _expand_nonzip_archive(tmp_dir, path, rel)
    return []


def _upload_relpath(client_filename):
    """
    Parse upload filename into safe relative path, e.g. "a/b/c.ext".
    Return None for invalid paths.
    """
    if not client_filename:
        return None
    norm = str(client_filename).replace("\\", "/").strip()
    if ".." in norm:
        return None
    parts = [p for p in norm.split("/") if p and p not in (".", "..")]
    if not parts:
        return None
    safe_parts = []
    for p in parts:
        s = p
        for ch in '<>:"|?*':
            s = s.replace(ch, "_")
        s = s.strip()
        if not s:
            return None
        safe_parts.append(s)
    rel = "/".join(safe_parts)
    if len(rel) > 300:
        rel = rel[-300:]
    return rel


def _write_original_backup(tmp_dir, rel, saved_path):
    """保存上传后立即复制一份原文，供单文件超时后 ZIP 仍包含未处理版本。"""
    backup = os.path.join(tmp_dir, "_aiprint_orig", *rel.split("/"))
    os.makedirs(os.path.dirname(backup), exist_ok=True)
    shutil.copy2(saved_path, backup)
    return backup


def _zip_arcname_timeout_original(original_rel):
    """ZIP 内超时条目：在文件名（不含扩展名）后加 _【超时原文】，保留相对目录结构。"""
    s = (original_rel or "").replace("\\", "/").strip()
    if not s:
        return "【超时原文】.bin"
    parts = [p for p in s.split("/") if p and p not in (".", "..")]
    if not parts:
        return "【超时原文】.bin"
    base, ext = os.path.splitext(parts[-1])
    if ext:
        parts[-1] = base + "_【超时原文】" + ext
    else:
        parts[-1] = parts[-1] + "_【超时原文】"
    return "/".join(parts)


def _detail_is_timeout_skip(d):
    """是否为单文件总超时跳过（ZIP 内改用上传备份并改文件名）。"""
    if not isinstance(d, dict):
        return False
    if d.get("timeout_skip"):
        return True
    return "【超时跳过】" in (d.get("message") or "")


def _get_printers():
    """Return available printers on Windows."""
    try:
        import win32print
        out = []
        for p in win32print.EnumPrinters(win32print.PRINTER_ENUM_LOCAL | win32print.PRINTER_ENUM_CONNECTIONS):
            out.append(p[2])
        return out
    except Exception:
        return []


@app.route("/")
def index():
    return send_from_directory(os.path.join(ROOT, "static"), "index.html")


@app.route("/favicon.ico")
def favicon():
    """Return empty favicon response to avoid 404."""
    return "", 204


@app.route("/api/printers")
def api_printers():
    """Return available printer list."""
    return jsonify({"printers": _get_printers()})


@app.route("/api/batch-print", methods=["POST"])
def api_batch_print():
    """
    Receive upload and options, execute batch print, return JSON.
    Fields: files, check_signature, accept_revisions, printer_name, copies, dry_run.
    """
    files = request.files.getlist("files") or request.files.getlist("files[]")
    if not files or not any(f.filename for f in files):
        return jsonify({"ok": False, "error": "?????????"}), 400

    try:
        opts = _parse_batch_opts_from_form(request.form)
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400

    tmp_dir = None
    try:
        tmp_dir = tempfile.mkdtemp(prefix="aiprintword_")
        saved_paths = []
        original_names = []
        original_backup_paths = []
        for f in files:
            rel = _upload_relpath(f.filename)
            if not rel:
                continue
            path = os.path.join(tmp_dir, *rel.split("/"))
            os.makedirs(os.path.dirname(path), exist_ok=True)
            f.save(path)
            if _is_archive_upload(rel):
                for p2, r2 in _expand_archive_for_batch(tmp_dir, path, rel):
                    saved_paths.append(p2)
                    original_names.append(r2)
                    bk = _write_original_backup(tmp_dir, r2, p2)
                    original_backup_paths.append(bk)
            elif _allowed_file(rel):
                saved_paths.append(path)
                original_names.append(rel)
                bk = _write_original_backup(tmp_dir, rel, path)
                original_backup_paths.append(bk)

        if not saved_paths:
            return jsonify(
                {
                    "ok": False,
                    "error": "未找到可处理文件：压缩包内需含 .doc/.docx/.xls/.xlsx/.pdf 等文档。中文 Windows 自带的「压缩文件夹」ZIP 若解压不到文件，请安装 7-Zip（7z.exe 在 PATH 或设置 SEVEN_ZIP_EXE）。.7z 无 7z 时可 pip 安装 py7zr；.rar 需 7-Zip 或 UnRAR。",
                }
            ), 400

        from batch_print import run_batch, build_batch_modification_zip_text
        result = run_batch(
            saved_paths,
            recursive=False,
            check_formal=opts["check_formal"],
            check_signature=opts["check_signature"],
            accept_revisions=opts["accept_revisions"],
            word_content_preserve=opts["word_content_preserve"],
            word_preserve_page_count=opts["word_preserve_page_count"],
            word_image_risk_guard=opts["word_image_risk_guard"],
            word_step_timeout_sec=opts.get("word_step_timeout_sec"),
            word_skip_file_on_timeout=opts.get("word_skip_file_on_timeout"),
            file_timeout_sec=opts.get("file_timeout_sec"),
            word_font_profile=opts.get("doc_font_profile", "mixed"),
            printer_name=opts["printer_name"],
            copies=opts["copies"],
            dry_run=opts["dry_run"],
            skip_print=opts["skip_print"],
            raw_print=opts["raw_print"],
            incremental_output_dir=opts.get("incremental_output_dir"),
            relative_names=original_names,
            incremental_exists_action=opts.get("incremental_exists_action", "overwrite"),
        )
        for i, d in enumerate(result["details"]):
            d["filename"] = original_names[i] if i < len(original_names) else os.path.basename(d["path"])
        if opts.get("skip_print") and result.get("total"):
            try:
                token = uuid.uuid4().hex
                mod_txt = build_batch_modification_zip_text(result["details"])
                _, zip_ftp_ok, zip_ftp_err = _zip_batch_exports(
                    result["details"],
                    # 取消/中断时不打包未处理文件（这些仍是原稿副本）
                    saved_paths[: len(result.get("details") or [])],
                    original_names[: len(result.get("details") or [])],
                    original_backup_paths[: len(result.get("details") or [])],
                    token,
                    modification_report_text=mod_txt,
                )
                result["download_token"] = token
                result["download_filename"] = "processed_documents.zip"
                result["zip_ftp_uploaded"] = bool(zip_ftp_ok)
                if zip_ftp_err:
                    result["zip_ftp_error"] = zip_ftp_err
            except Exception as e:
                return jsonify({"ok": False, "error": _format_com_error(e)}), 500
        return jsonify({"ok": True, "result": result})
    except Exception as e:
        return jsonify({"ok": False, "error": _format_com_error(e)}), 500
    finally:
        if tmp_dir and os.path.isdir(tmp_dir):
            try:
                shutil.rmtree(tmp_dir, ignore_errors=True)
            except Exception:
                pass


def _parse_sse_heartbeat_sec(form):
    """
    批处理 SSE：队列 get 超时时间 = 无 progress 时推送 heartbeat 的间隔。
    默认 120 秒；可由表单 sse_heartbeat_sec 或环境变量 AIPRINTWORD_SSE_HEARTBEAT_SEC 覆盖；限制 10–600 秒。
    """
    try:
        raw = None
        if form is not None and hasattr(form, "get"):
            raw = form.get("sse_heartbeat_sec")
        if raw is None or str(raw).strip() == "":
            try:
                from runtime_settings.resolve import get_setting

                raw = get_setting("AIPRINTWORD_SSE_HEARTBEAT_SEC")
            except Exception:
                raw = os.environ.get("AIPRINTWORD_SSE_HEARTBEAT_SEC", "120")
        v = float(str(raw).strip())
    except Exception:
        v = 120.0
    return max(10.0, min(600.0, v))


def _parse_batch_opts_from_form(form):
    """Parse run_mode: formal_save | standard_print | raw_print."""
    incremental_output_dir = _parse_incremental_output_dir(form)
    run_mode = (form.get("run_mode") or "standard_print").strip().lower()
    if run_mode not in ("formal_save", "standard_print", "raw_print"):
        run_mode = "standard_print"
    word_content_preserve = form.get("word_content_preserve", "true").lower() == "true"
    word_preserve_page_count = form.get("word_preserve_page_count", "true").lower() == "true"
    word_image_risk_guard = form.get("word_image_risk_guard", "false").lower() == "true"
    try:
        word_step_timeout_min = float(form.get("word_step_timeout_min", "60") or "60")
        word_step_timeout_min = max(0.0, min(word_step_timeout_min, 24 * 60.0))
    except Exception:
        word_step_timeout_min = 60.0
    word_skip_file_on_timeout = form.get("word_skip_file_on_timeout", "true").lower() == "true"
    try:
        # 默认 45 分钟：为 0 时单文件总超时看门狗不生效（不会自动杀进程跳过）
        file_timeout_min = float(form.get("file_timeout_min", "45") or "45")
        file_timeout_min = max(0.0, min(file_timeout_min, 24 * 60.0))
    except Exception:
        file_timeout_min = 45.0
    incremental_exists_action = (form.get("incremental_exists_action") or "overwrite").strip().lower()
    if incremental_exists_action not in ("overwrite", "skip"):
        incremental_exists_action = "overwrite"
    doc_font_profile = (
        form.get("doc_font_profile")
        or form.get("word_font_profile")
        or "mixed"
    ).strip().lower()
    if doc_font_profile not in ("chinese", "english", "mixed"):
        doc_font_profile = "mixed"
    printer_name = form.get("printer_name") or None
    if printer_name == "":
        printer_name = None
    try:
        copies = int(form.get("copies", "1") or "1")
        copies = max(1, min(copies, 99))
    except ValueError:
        copies = 1

    if run_mode == "formal_save":
        # ?????????????????? ? ???? ? ????/??????????????? ZIP
        return {
            "run_mode": run_mode,
            "check_formal": True,
            "check_signature": True,
            "accept_revisions": True,
            "dry_run": False,
            "skip_print": True,
            "raw_print": False,
            "word_content_preserve": word_content_preserve,
            "word_preserve_page_count": word_preserve_page_count,
            "word_image_risk_guard": word_image_risk_guard,
            "word_step_timeout_sec": word_step_timeout_min * 60.0,
            "word_skip_file_on_timeout": word_skip_file_on_timeout,
            "file_timeout_sec": file_timeout_min * 60.0,
            "doc_font_profile": doc_font_profile,
            "printer_name": printer_name,
            "copies": copies,
            "incremental_output_dir": incremental_output_dir,
            "incremental_exists_action": incremental_exists_action,
        }
    if run_mode == "raw_print":
        return {
            "run_mode": run_mode,
            "check_formal": False,
            "check_signature": False,
            "accept_revisions": False,
            "dry_run": False,
            "skip_print": False,
            "raw_print": True,
            "word_content_preserve": word_content_preserve,
            "word_preserve_page_count": word_preserve_page_count,
            "word_image_risk_guard": word_image_risk_guard,
            "word_step_timeout_sec": word_step_timeout_min * 60.0,
            "word_skip_file_on_timeout": word_skip_file_on_timeout,
            "file_timeout_sec": file_timeout_min * 60.0,
            "doc_font_profile": doc_font_profile,
            "printer_name": printer_name,
            "copies": copies,
            "incremental_output_dir": incremental_output_dir,
            "incremental_exists_action": incremental_exists_action,
        }
    return {
        "run_mode": run_mode,
        "check_formal": True,
        "check_signature": True,
        "accept_revisions": form.get("accept_revisions", "true").lower() == "true",
        "dry_run": form.get("dry_run", "false").lower() == "true",
        "skip_print": False,
        "raw_print": False,
        "word_content_preserve": word_content_preserve,
        "word_preserve_page_count": word_preserve_page_count,
        "word_image_risk_guard": word_image_risk_guard,
        "word_step_timeout_sec": word_step_timeout_min * 60.0,
        "word_skip_file_on_timeout": word_skip_file_on_timeout,
        "file_timeout_sec": file_timeout_min * 60.0,
        "doc_font_profile": doc_font_profile,
        "printer_name": printer_name,
        "copies": copies,
        "incremental_output_dir": incremental_output_dir,
        "incremental_exists_action": incremental_exists_action,
    }


def _zip_arcname_with_processed_ext(original_arcname: str, processed_path: str) -> str:
    """上传名为 .doc/.xls 等时，成品可能已转为 .docx/.xlsx，ZIP 内条目扩展名与磁盘成品一致。"""
    pe = os.path.splitext(processed_path)[1]
    if not pe:
        return original_arcname.replace("\\", "/")
    norm = original_arcname.replace("\\", "/")
    d, base = os.path.split(norm)
    stem = os.path.splitext(base)[0]
    new_base = stem + pe.lower()
    return f"{d}/{new_base}" if d else new_base


def _zip_batch_exports(
    details,
    saved_paths,
    original_names,
    original_backup_paths,
    token,
    modification_report_text=None,
):
    """写入 ZIP：修改明细 + 各文件。成功项必须打包落盘后的成品路径（含 .doc→.docx 等）。"""
    os.makedirs(BATCH_EXPORT_ROOT, exist_ok=True)
    zip_path = os.path.join(BATCH_EXPORT_ROOT, f"{token}.zip")
    backups = original_backup_paths or []
    n = min(len(saved_paths), len(original_names))
    if details is not None:
        n = min(n, len(details))
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        detail_text = (
            str(modification_report_text).strip()
            if modification_report_text is not None
            else ""
        )
        if not detail_text:
            detail_text = "（本批未生成修改说明正文；请核对各文件是否已按预期处理。）"
        zf.writestr(
            "\u4fee\u6539\u660e\u7ec6.txt",
            "\ufeff" + detail_text + "\n",
            compress_type=zipfile.ZIP_DEFLATED,
        )
        for i in range(n):
            name = original_names[i]
            work = saved_paths[i]
            d = details[i] if details and i < len(details) else {}
            backup = backups[i] if i < len(backups) else None
            if _detail_is_timeout_skip(d):
                if backup and os.path.isfile(backup):
                    src = backup
                else:
                    src = work
                arcname = _zip_arcname_timeout_original(name)
            else:
                proc = (d.get("processed_path") or "").strip()
                if d.get("success") and proc and os.path.isfile(proc):
                    src = proc
                    arcname = _zip_arcname_with_processed_ext(name, proc)
                else:
                    src = work
                    arcname = name.replace("\\", "/")
            if os.path.isfile(src):
                zf.write(src, arcname=arcname)
                # 终版成品优先 FTP；失败写入明细供排查，本地 ZIP 仍可用
                if d.get("success") and (d.get("processed_path") or "").strip() and os.path.isfile(src):
                    try:
                        from ftp_store import try_upload_file

                        ftp_p, up_err = try_upload_file(src, f"batch/final/{token}/{arcname}")
                        if ftp_p:
                            d["final_ftp_path"] = ftp_p
                        elif up_err:
                            d["ftp_upload_error"] = up_err
                    except Exception:
                        pass
    # 同步 ZIP 到 FTP（供下载复用；失败不影响主流程）
    zip_ftp_err: Optional[str] = None
    try:
        from ftp_store import try_upload_file

        pz, ez = try_upload_file(zip_path, f"batch/exports/{token}.zip")
        ftp_zip_ok = bool(pz)
        if ez:
            zip_ftp_err = ez
    except Exception:
        ftp_zip_ok = False
    return zip_path, ftp_zip_ok, zip_ftp_err


def _history_index_path():
    os.makedirs(BATCH_HISTORY_ROOT, exist_ok=True)
    return os.path.join(BATCH_HISTORY_ROOT, "index.json")


def _load_history_index():
    try:
        import batch_history_mysql as bhm

        if bhm.enabled():
            return {"entries": bhm.list_summaries(_batch_history_max())}
    except Exception:
        pass
    p = _history_index_path()
    if not os.path.isfile(p):
        return {"entries": []}
    try:
        with open(p, encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, dict) and isinstance(data.get("entries"), list):
            return data
    except Exception:
        pass
    return {"entries": []}


def _save_history_index(data):
    p = _history_index_path()
    try:
        with open(p, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.warning("history index save failed: %s", e)


def _history_trim_folders(keep_ids):
    if not os.path.isdir(BATCH_HISTORY_ROOT):
        return
    for name in os.listdir(BATCH_HISTORY_ROOT):
        if name == "index.json":
            continue
        path = os.path.join(BATCH_HISTORY_ROOT, name)
        if not os.path.isdir(path):
            continue
        if name in keep_ids:
            continue
        try:
            shutil.rmtree(path, ignore_errors=True)
        except Exception:
            pass


def _history_summary_from_record(rec):
    from batch_history_mysql import compute_display_title, _safe_ftp_err_text

    r = rec.get("result") or {}
    if not isinstance(r, dict):
        r = {}
    title = rec.get("display_title") or compute_display_title(
        rec.get("original_names"), r
    )
    ze = _safe_ftp_err_text(r.get("zip_ftp_error")) or _safe_ftp_err_text(
        rec.get("zip_ftp_error")
    )
    has_zip = bool(rec.get("download_token"))
    zup = r.get("zip_ftp_uploaded")
    if zup is None:
        zup = rec.get("zip_ftp_uploaded")
    if not has_zip:
        zip_ftp = None
    elif zup is True or zup == 1:
        zip_ftp = True
    elif zup is False or zup == 0:
        zip_ftp = False
    else:
        zip_ftp = None
    return {
        "id": rec["id"],
        "created_at": rec.get("created_at", ""),
        "run_mode": rec.get("run_mode"),
        "title": title,
        "total": int(r.get("total") or 0),
        "ok": int(r.get("ok") or 0),
        "failed": int(r.get("failed") or 0),
        "has_zip": has_zip,
        "has_stash": bool(rec.get("has_stash")),
        "success": bool(rec.get("payload_ok")),
        "zip_ftp": zip_ftp,
        "zip_ftp_error": (ze or None) if has_zip else None,
    }


def _record_to_default_form(record):
    """合并到 CombinedMultiDict 时在后半部分，仅补全表单未传的项。"""
    from werkzeug.datastructures import MultiDict

    o = record.get("options") or {}
    m = MultiDict()
    rm = record.get("run_mode") or o.get("run_mode") or "standard_print"
    m.set("run_mode", rm)

    def tb(x, default=False):
        return "true" if (x if x is not None else default) else "false"

    m.set("accept_revisions", tb(o.get("accept_revisions"), True))
    m.set("dry_run", tb(o.get("dry_run"), False))
    m.set("word_content_preserve", tb(o.get("word_content_preserve"), True))
    m.set("word_preserve_page_count", tb(o.get("word_preserve_page_count"), True))
    m.set("word_image_risk_guard", tb(o.get("word_image_risk_guard"), False))
    m.set("word_skip_file_on_timeout", tb(o.get("word_skip_file_on_timeout"), True))
    fts = o.get("file_timeout_sec")
    if fts is not None:
        try:
            m.set("file_timeout_min", str(max(0, int(float(fts) / 60.0))))
        except Exception:
            m.set("file_timeout_min", "45")
    wst = o.get("word_step_timeout_sec")
    if wst is not None:
        try:
            m.set("word_step_timeout_min", str(max(0, int(float(wst) / 60.0))))
        except Exception:
            m.set("word_step_timeout_min", "60")
    m.set("doc_font_profile", (o.get("doc_font_profile") or "mixed").strip())
    pn = o.get("printer_name")
    m.set("printer_name", pn if pn else "")
    try:
        m.set("copies", str(max(1, min(99, int(o.get("copies", 1))))))
    except Exception:
        m.set("copies", "1")
    iod = o.get("incremental_output_dir")
    m.set("incremental_output_dir", iod if iod else "")
    m.set("incremental_exists_action", (o.get("incremental_exists_action") or "overwrite").strip())
    m.set("sse_heartbeat_sec", str(int(_parse_sse_heartbeat_sec(None))))
    return m


def _batch_history_persist(stream_state, tmp_dir):
    """任务结束后写入历史；有失败或整体报错时整包暂存 stash 供重试。"""
    try:
        os.makedirs(BATCH_HISTORY_ROOT, exist_ok=True)
    except Exception:
        return None
    payload = stream_state.get("final_json")
    opts = stream_state.get("opts") or {}
    hid = uuid.uuid4().hex
    hist_dir = os.path.join(BATCH_HISTORY_ROOT, hid)
    try:
        os.makedirs(hist_dir, exist_ok=True)
    except Exception:
        return None

    ok = bool(payload.get("ok")) if isinstance(payload, dict) else False
    result = (payload or {}).get("result") if ok else None
    err = (payload or {}).get("error") if not ok else None
    failed = int((result or {}).get("failed") or 0)
    needs_stash = (not ok) or (failed > 0)
    stash_rel = "stash"
    stash_path = os.path.join(hist_dir, stash_rel)

    if needs_stash and tmp_dir and os.path.isdir(tmp_dir):
        try:
            shutil.copytree(tmp_dir, stash_path, dirs_exist_ok=True)
        except TypeError:
            if os.path.isdir(stash_path):
                shutil.rmtree(stash_path, ignore_errors=True)
            shutil.copytree(tmp_dir, stash_path)
        except Exception as e:
            logger.warning("history stash copy failed: %s", e)
            needs_stash = False
            try:
                if os.path.isdir(stash_path):
                    shutil.rmtree(stash_path, ignore_errors=True)
            except Exception:
                pass

    has_stash = needs_stash and os.path.isdir(stash_path)
    from batch_history_mysql import compute_display_title

    record = {
        "id": hid,
        "created_at": time.strftime("%Y-%m-%dT%H:%M:%S", time.localtime()),
        "run_mode": opts.get("run_mode"),
        "options": {
            "run_mode": opts.get("run_mode"),
            "printer_name": opts.get("printer_name"),
            "copies": opts.get("copies"),
            "dry_run": opts.get("dry_run"),
            "skip_print": opts.get("skip_print"),
            "raw_print": opts.get("raw_print"),
            "accept_revisions": opts.get("accept_revisions"),
            "word_content_preserve": opts.get("word_content_preserve"),
            "word_preserve_page_count": opts.get("word_preserve_page_count"),
            "word_image_risk_guard": opts.get("word_image_risk_guard"),
            "doc_font_profile": opts.get("doc_font_profile"),
            "incremental_output_dir": opts.get("incremental_output_dir"),
            "incremental_exists_action": opts.get("incremental_exists_action"),
            "file_timeout_sec": opts.get("file_timeout_sec"),
            "word_step_timeout_sec": opts.get("word_step_timeout_sec"),
            "word_skip_file_on_timeout": opts.get("word_skip_file_on_timeout"),
        },
        "original_names": list(stream_state.get("original_names") or []),
        "payload_ok": ok,
        "error": err,
        "result": result,
        "download_token": (result or {}).get("download_token") if result else None,
        "has_stash": has_stash,
        "zip_ftp_uploaded": bool((result or {}).get("zip_ftp_uploaded"))
        if result
        else False,
        "zip_ftp_error": ((result or {}).get("zip_ftp_error") or "").strip() or None
        if result
        else None,
    }
    # 打印模式也上传“终版成品”到 FTP，并在历史记录中写入下载信息
    try:
        if ok and isinstance(result, dict):
            details = result.get("details") or []
            orig = list(stream_state.get("original_names") or [])
            # 仅在有明细时处理
            if isinstance(details, list) and details:
                from ftp_store import try_upload_file

                for i, d in enumerate(details):
                    if not isinstance(d, dict):
                        continue
                    proc = (d.get("processed_path") or "").strip()
                    if not d.get("success") or not proc or (not os.path.isfile(proc)):
                        continue
                    name = orig[i] if i < len(orig) else (d.get("filename") or d.get("path") or f"file_{i+1}")
                    arcname = _zip_arcname_with_processed_ext(str(name), proc)
                    try:
                        ftp_p, up_err = try_upload_file(proc, f"batch/final/{hid}/{arcname}")
                        if ftp_p:
                            d["final_ftp_path"] = ftp_p
                        elif up_err:
                            d["final_ftp_error"] = up_err
                    except Exception:
                        pass
    except Exception:
        pass
    record["display_title"] = compute_display_title(
        record["original_names"], result
    )
    try:
        import batch_history_mysql as bhm

        if bhm.enabled():
            bhm.save_record(record)
            keep_list = bhm.trim_to_max(_batch_history_max())
            _history_trim_folders(set(keep_list))
            return hid
    except Exception as e:
        logger.warning("batch history MySQL 写入失败，回退本地 record.json：%s", e)

    try:
        with open(os.path.join(hist_dir, "record.json"), "w", encoding="utf-8") as f:
            json.dump(record, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.warning("history record write failed: %s", e)
        return None

    idx = _load_history_index()
    entries = idx.setdefault("entries", [])
    summ = _history_summary_from_record(record)
    entries.insert(0, summ)
    entries[:] = entries[: _batch_history_max()]
    keep = {e["id"] for e in entries}
    _history_trim_folders(keep)
    _save_history_index(idx)
    return hid


def _history_record_path(hid):
    if not _BATCH_EXPORT_TOKEN_RE.match(hid or ""):
        return None
    p = os.path.join(BATCH_HISTORY_ROOT, hid, "record.json")
    return p if os.path.isfile(p) else None


def _load_history_record(hid):
    try:
        import batch_history_mysql as bhm

        if bhm.enabled():
            rec = bhm.get_record(hid)
            if rec:
                return rec
    except Exception:
        pass
    rp = _history_record_path(hid)
    if not rp:
        return None
    try:
        with open(rp, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def _collect_stash_files_for_retry(stash_root, original_names, details, only_failed):
    """从暂存目录收集文件；only_failed 时仅失败条目对应文档 + 全部非文档资源。"""
    failed_set = set()
    for i, d in enumerate(details or []):
        if i < len(original_names) and not d.get("success"):
            failed_set.add(str(original_names[i]).replace("\\", "/"))
    out = []
    for root, dirs, files in os.walk(stash_root):
        rel_root = os.path.relpath(root, stash_root)
        parts = [] if rel_root == "." else rel_root.replace("\\", "/").split("/")
        if "_aiprint_orig" in parts:
            continue
        for fn in files:
            abs_path = os.path.join(root, fn)
            rel = "/".join(parts + [fn]) if parts else fn
            rel = rel.replace("\\", "/")
            ext = os.path.splitext(fn)[1].lower()
            is_doc = ext in ALLOWED_EXT
            if only_failed:
                if is_doc and rel not in failed_set:
                    continue
            out.append((abs_path, rel))
    return out


def _materialize_retry_workspace(stash_root, original_names, details, only_failed):
    pairs = _collect_stash_files_for_retry(stash_root, original_names, details, only_failed)
    if not pairs:
        return None, None, None, None
    tmp_dir = tempfile.mkdtemp(prefix="aiprintword_retry_")
    try:
        for abs_p, rel in pairs:
            dest = os.path.join(tmp_dir, *rel.split("/"))
            os.makedirs(os.path.dirname(dest), exist_ok=True)
            shutil.copy2(abs_p, dest)
        doc_rels = [rel for _abs, rel in pairs if _allowed_file(rel)]
        order_map = {n: i for i, n in enumerate(original_names)}
        doc_rels_sorted = sorted(doc_rels, key=lambda n: order_map.get(n, 10**9))
        saved_paths = []
        original_names_ord = []
        original_backup_paths = []
        for rel in doc_rels_sorted:
            path = os.path.join(tmp_dir, *rel.split("/"))
            if not os.path.isfile(path):
                continue
            saved_paths.append(path)
            original_names_ord.append(rel)
            original_backup_paths.append(_write_original_backup(tmp_dir, rel, path))
        if not saved_paths:
            shutil.rmtree(tmp_dir, ignore_errors=True)
            return None, None, None, None
        return tmp_dir, saved_paths, original_names_ord, original_backup_paths
    except Exception:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        return None, None, None, None


def _admin_settings_auth_error():
    _sync_admin_token_from_disk()
    dotenv_path = _resolved_dotenv_path()
    exp = (os.environ.get("AIPRINTWORD_ADMIN_TOKEN") or "").strip()
    if not exp:
        exp = _read_key_from_env_file(dotenv_path, "AIPRINTWORD_ADMIN_TOKEN")
        if exp:
            os.environ["AIPRINTWORD_ADMIN_TOKEN"] = exp
    if not exp:
        exists = os.path.isfile(dotenv_path)
        tf_default = os.path.isfile(_ADMIN_TOKEN_FILE_DEFAULT)
        logger.warning(
            "管理 API 仍无口令：ROOT=%s dotenv=%s 存在=%s 默认口令文件存在=%s cwd=%s",
            ROOT,
            dotenv_path,
            exists,
            tf_default,
            os.getcwd(),
        )
        return (
            {
                "error": (
                    "服务端未配置管理口令：请在运行服务的机器上配置 AIPRINTWORD_ADMIN_TOKEN，"
                    "或设置 AIPRINTWORD_DOTENV_PATH 指向 .env 绝对路径，或创建 data/admin_token.txt。"
                    " 部署后本响应应含 diag 字段；若仍只有短 JSON（约 138 字节），说明未加载最新 app.py。"
                ),
                "diag": {
                    "auth_api_build": AIPRINTWORD_WEB_BUILD,
                    "app_root": ROOT,
                    "dotenv_path": dotenv_path,
                    "dotenv_exists": exists,
                    "cwd": os.getcwd(),
                    "admin_token_file_default": _ADMIN_TOKEN_FILE_DEFAULT,
                    "admin_token_file_exists": tf_default,
                },
            },
            503,
        )
    tok = (request.headers.get("X-Admin-Token") or request.args.get("token") or "").strip()
    if tok != exp:
        if not tok:
            return (
                "未授权：未携带管理口令。请在设置页「管理口令」输入与服务器 .env 中 "
                "AIPRINTWORD_ADMIN_TOKEN 完全相同的值后再点「加载」（请求头 X-Admin-Token）。",
                401,
            )
        return (
            "未授权：口令与服务器不一致。请核对页面输入与 .env 中 AIPRINTWORD_ADMIN_TOKEN "
            "是否完全一致（区分大小写、首尾勿多空格）。",
            401,
        )
    return None, None


@app.route("/settings")
def settings_page():
    return send_from_directory(os.path.join(ROOT, "static"), "settings.html")


@app.route("/api/aiprintword-build")
def api_aiprintword_build():
    """
    无鉴权：确认当前进程实际加载的 app.py 路径与构建号。
    若此处 build>=4 而 /api/admin/settings 仍为 138 字节 503，则几乎不可能（请清 CDN/代理缓存或核对是否多台服务）。
    """
    return jsonify(
        {
            "ok": True,
            "build": AIPRINTWORD_WEB_BUILD,
            "app_py": os.path.abspath(__file__),
            "resolved_dotenv": _resolved_dotenv_path(),
        }
    )


@app.route("/api/admin/settings", methods=["GET"])
def api_admin_settings_get():
    err, code = _admin_settings_auth_error()
    if err:
        if isinstance(err, dict):
            return jsonify({"ok": False, **err}), code
        return jsonify({"ok": False, "error": err}), code
    from runtime_settings.resolve import list_all_settings

    return jsonify({"ok": True, "items": list_all_settings(mask_secrets=True)})


@app.route("/api/admin/settings", methods=["PUT"])
def api_admin_settings_put():
    err, code = _admin_settings_auth_error()
    if err:
        if isinstance(err, dict):
            return jsonify({"ok": False, **err}), code
        return jsonify({"ok": False, "error": err}), code
    try:
        body = request.get_json(force=True, silent=True) or {}
    except Exception:
        body = {}
    if not isinstance(body, dict):
        return jsonify({"ok": False, "error": "请求体须为 JSON 对象"}), 400
    updates = {}
    for k, v in body.items():
        ks = str(k).strip()
        if not ks:
            continue
        if v is None:
            updates[ks] = ""
        elif isinstance(v, bool):
            updates[ks] = "1" if v else "0"
        else:
            updates[ks] = str(v).strip()
    try:
        from runtime_settings.resolve import set_settings

        changed = set_settings(updates)
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    except RuntimeError as e:
        return jsonify({"ok": False, "error": str(e)}), 503
    except Exception as e:
        return jsonify({"ok": False, "error": _format_com_error(e)}), 500
    try:
        from doc_handlers.word_handler import apply_resolved_word_base_settings

        apply_resolved_word_base_settings()
    except Exception:
        pass
    return jsonify({"ok": True, "updated": changed})


@app.route("/api/admin/batch-history/migrate-from-disk", methods=["POST"])
def api_admin_batch_history_migrate_from_disk():
    """将 data/batch_history/<id>/record.json 中尚未入库的记录导入 MySQL，并补全 display_title。"""
    err, code = _admin_settings_auth_error()
    if err:
        if isinstance(err, dict):
            return jsonify({"ok": False, **err}), code
        return jsonify({"ok": False, "error": err}), code
    import batch_history_mysql as bhm

    if not bhm.enabled():
        return (
            jsonify(
                {
                    "ok": False,
                    "error": "未配置 MySQL（MYSQL_HOST），无法写入历史库",
                }
            ),
            400,
        )
    try:
        st = bhm.migrate_from_disk(BATCH_HISTORY_ROOT, _batch_history_max())
        backfilled = bhm.backfill_display_titles()
        return jsonify({"ok": True, **st, "backfilled": backfilled})
    except Exception as e:
        logger.exception("batch history migrate-from-disk failed")
        return jsonify({"ok": False, "error": _format_com_error(e)}), 500


@app.route("/api/admin/sign/migrate-mysql-blobs-to-ftp", methods=["POST"])
def api_admin_sign_migrate_mysql_blobs_to_ftp():
    """将在线签名模块历史遗留的 MySQL BLOB 文件迁移到 FTP。"""
    err, code = _admin_settings_auth_error()
    if err:
        if isinstance(err, dict):
            return jsonify({"ok": False, **err}), code
        return jsonify({"ok": False, "error": err}), code
    try:
        from sign_handlers import mysql_store

        if not mysql_store.mysql_sign_enabled():
            return jsonify({"ok": False, "error": "未配置 MySQL（MYSQL_HOST），无需迁移或无法读取"}), 400
        data = request.get_json(silent=True) or {}
        batch_size = int(data.get("batch_size", 2000) or 2000)
        max_total = int(data.get("max_total", 200000) or 200000)
        clear_blob = bool(data.get("clear_blob", True))
        st = mysql_store.migrate_mysql_blobs_to_ftp(
            batch_size=batch_size, max_total=max_total, clear_blob=clear_blob
        )
        st2 = mysql_store.migrate_signer_strokes_blobs_to_ftp(
            batch_size=batch_size, max_total=max_total, clear_blob=clear_blob
        )
        st3 = mysql_store.migrate_stroke_items_blobs_to_ftp(
            batch_size=batch_size, max_total=max_total, clear_blob=clear_blob
        )
        st4 = mysql_store.verify_and_backfill_ftp_files(limit=batch_size)
        return jsonify({"ok": True, "files": st, "legacy_strokes": st2, "stroke_items": st3, "verify": st4})
    except Exception as e:
        return jsonify({"ok": False, "error": _format_com_error(e)}), 500


@app.route("/api/admin/sign/reload-slot-layout-rules", methods=["POST"])
def api_admin_sign_reload_slot_layout_rules():
    """运行时热重载签字位版式规则（sign_slot_layout_rules.json）。"""
    err, code = _admin_settings_auth_error()
    if err:
        if isinstance(err, dict):
            return jsonify({"ok": False, **err}), code
        return jsonify({"ok": False, "error": err}), code
    try:
        from sign_handlers import config as sign_config

        sign_config.reload_sign_slot_layout_rules_from_disk()
        rules = sign_config.SIGN_SLOT_LAYOUT_RULES
        slot = rules.get("replace_prefilled_slot") or {}
        cfg_path = os.path.join(
            ROOT, "sign_handlers", "sign_slot_layout_rules.json"
        )
        return jsonify(
            {
                "ok": True,
                "reloaded": True,
                "config_path": cfg_path,
                "schema_version": int(rules.get("schema_version", 1) or 1),
                "replace_prefilled_slot": {
                    "enabled": bool(slot.get("enabled", True)),
                    "max_text_len": int(slot.get("max_text_len", 48) or 48),
                    "fullmatch_patterns_count": len(slot.get("fullmatch_patterns") or []),
                    "search_patterns_count": len(slot.get("search_patterns") or []),
                },
            }
        )
    except Exception as e:
        logger.exception("reload slot layout rules failed")
        return jsonify({"ok": False, "error": _format_com_error(e)}), 500


@app.route("/api/admin/sign/slot-layout-rules", methods=["GET"])
def api_admin_sign_slot_layout_rules_get():
    """查看当前生效的签字位版式规则与磁盘配置。"""
    err, code = _admin_settings_auth_error()
    if err:
        if isinstance(err, dict):
            return jsonify({"ok": False, **err}), code
        return jsonify({"ok": False, "error": err}), code
    try:
        import json as _json
        from sign_handlers import config as sign_config

        rules = sign_config.SIGN_SLOT_LAYOUT_RULES
        slot = rules.get("replace_prefilled_slot") or {}
        cfg_path = os.path.join(ROOT, "sign_handlers", "sign_slot_layout_rules.json")

        disk_rules = None
        disk_load_error = None
        disk_validation_error = None
        try:
            with open(cfg_path, "r", encoding="utf-8") as f:
                disk_rules = _json.load(f)
            from sign_handlers import config as _scfg

            _scfg.validate_sign_slot_layout_rules_payload(disk_rules)
        except Exception as e:
            if disk_rules is None:
                disk_load_error = _format_com_error(e)
            else:
                disk_validation_error = _format_com_error(e)

        return jsonify(
            {
                "ok": True,
                "config_path": cfg_path,
                "config_exists": os.path.isfile(cfg_path),
                "disk_load_error": disk_load_error,
                "disk_validation_error": disk_validation_error,
                "runtime_rules": {
                    "schema_version": int(rules.get("schema_version", 1) or 1),
                    "replace_prefilled_slot": {
                        "enabled": bool(slot.get("enabled", True)),
                        "max_text_len": int(slot.get("max_text_len", 48) or 48),
                        "fullmatch_patterns": [
                            getattr(p, "pattern", str(p))
                            for p in (slot.get("fullmatch_patterns") or [])
                        ],
                        "search_patterns": [
                            getattr(p, "pattern", str(p))
                            for p in (slot.get("search_patterns") or [])
                        ],
                    },
                },
                "disk_rules": disk_rules,
            }
        )
    except Exception as e:
        logger.exception("get slot layout rules failed")
        return jsonify({"ok": False, "error": _format_com_error(e)}), 500


@app.route("/api/admin/sign/slot-layout-rules/download", methods=["GET"])
def api_admin_sign_slot_layout_rules_download():
    """下载当前磁盘上的签字位版式规则 JSON。"""
    err, code = _admin_settings_auth_error()
    if err:
        if isinstance(err, dict):
            return jsonify({"ok": False, **err}), code
        return jsonify({"ok": False, "error": err}), code
    cfg_path = os.path.join(ROOT, "sign_handlers", "sign_slot_layout_rules.json")
    if not os.path.isfile(cfg_path):
        return jsonify({"ok": False, "error": "规则文件不存在"}), 404
    return send_file(
        cfg_path,
        as_attachment=True,
        download_name="sign_slot_layout_rules.json",
        mimetype="application/json",
        max_age=0,
    )


@app.route("/api/admin/sign/slot-layout-rules/upload", methods=["POST"])
def api_admin_sign_slot_layout_rules_upload():
    """上传并覆盖签字位版式规则 JSON；默认上传后立即热加载。"""
    err, code = _admin_settings_auth_error()
    if err:
        if isinstance(err, dict):
            return jsonify({"ok": False, **err}), code
        return jsonify({"ok": False, "error": err}), code
    cfg_path = os.path.join(ROOT, "sign_handlers", "sign_slot_layout_rules.json")
    try:
        f = request.files.get("file")
        if not f:
            return jsonify({"ok": False, "error": "缺少上传文件（file）"}), 400
        raw = f.read()
        if not raw:
            return jsonify({"ok": False, "error": "上传文件为空"}), 400
        try:
            text = raw.decode("utf-8")
        except Exception:
            text = raw.decode("utf-8-sig")
        obj = json.loads(text)
        from sign_handlers import config as sign_config

        normalized = sign_config.validate_sign_slot_layout_rules_payload(obj)
        with open(cfg_path, "w", encoding="utf-8") as fp:
            json.dump(normalized, fp, ensure_ascii=False, indent=2)
            fp.write("\n")

        sign_config.reload_sign_slot_layout_rules_from_disk()
        rules = sign_config.SIGN_SLOT_LAYOUT_RULES
        slot = rules.get("replace_prefilled_slot") or {}
        return jsonify(
            {
                "ok": True,
                "uploaded": True,
                "reloaded": True,
                "config_path": cfg_path,
                "schema_version": int(rules.get("schema_version", 1) or 1),
                "replace_prefilled_slot": {
                    "enabled": bool(slot.get("enabled", True)),
                    "max_text_len": int(slot.get("max_text_len", 48) or 48),
                    "fullmatch_patterns_count": len(slot.get("fullmatch_patterns") or []),
                    "search_patterns_count": len(slot.get("search_patterns") or []),
                },
            }
        )
    except json.JSONDecodeError as e:
        return jsonify({"ok": False, "error": f"JSON 格式错误: {e}"}), 400
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    except Exception as e:
        logger.exception("upload slot layout rules failed")
        return jsonify({"ok": False, "error": _format_com_error(e)}), 500


@app.route("/api/batch-export/<token>")
def api_batch_export(token):
    """?????????? ZIP?token ???????????"""
    if not _BATCH_EXPORT_TOKEN_RE.match(token or ""):
        return jsonify({"ok": False, "error": "invalid token"}), 404
    path = os.path.join(BATCH_EXPORT_ROOT, token + ".zip")
    if not os.path.isfile(path):
        try:
            from ftp_store import download_bytes

            b = download_bytes(f"batch/exports/{token}.zip")
            os.makedirs(BATCH_EXPORT_ROOT, exist_ok=True)
            with open(path, "wb") as fp:
                fp.write(b)
        except Exception:
            return jsonify({"ok": False, "error": "not found or expired"}), 404
    return send_file(
        path,
        as_attachment=True,
        download_name="processed_documents.zip",
        max_age=0,
    )


def _run_batch_with_progress(
    saved_paths, original_names, original_backup_paths, opts, queue
):
    """Run batch in background and push progress to queue."""
    try:
        import pythoncom
        pythoncom.CoInitialize()
    except Exception:
        pass
    job = opts.get("_job") if isinstance(opts, dict) else None
    cancel_event = getattr(job, "cancel_event", None) if job else None
    pause_event = getattr(job, "pause_event", None) if job else None
    job_id = getattr(job, "job_id", None) if job else None
    skip_current_event = getattr(job, "skip_current_event", None) if job else None

    try:
        from batch_print import run_batch, build_batch_modification_zip_text, _format_duration_cn

        def progress_callback(
            step,
            file_index,
            file_total,
            file_name,
            percent=None,
            eta=None,
            processing_meta=None,
        ):
            name = original_names[file_index - 1] if file_index <= len(original_names) else file_name
            eta_log = ""
            if isinstance(eta, dict) and eta.get("eta_remaining_sec") is not None:
                eta_log = " eta_rem≈%s avg=%ss" % (
                    _format_duration_cn(eta["eta_remaining_sec"]),
                    eta.get("avg_sec_per_file", "?"),
                )
            mode_log = ""
            if isinstance(processing_meta, dict) and processing_meta.get("processingModeLabel"):
                mode_log = " mode=%s" % (processing_meta.get("processingModeLabel"),)
            logger.info(
                "progress idx=%s/%s pct=%s step=%s file=%s%s%s",
                file_index,
                file_total,
                percent,
                step,
                name,
                eta_log,
                mode_log,
            )
            payload = {
                "step": step,
                "fileIndex": file_index,
                "fileTotal": file_total,
                "fileName": name,
                "percent": percent,
            }
            if isinstance(eta, dict):
                payload["eta"] = eta
            if isinstance(processing_meta, dict):
                if processing_meta.get("processingMode") is not None:
                    payload["processingMode"] = processing_meta["processingMode"]
                if processing_meta.get("processingModeLabel"):
                    payload["processingModeLabel"] = processing_meta["processingModeLabel"]
                pts = processing_meta.get("modificationPoints")
                if pts:
                    payload["modificationPoints"] = pts
            queue.put(("progress", payload))

        logger.info(
            "batch worker start total=%s mode=%s dry_run=%s skip_print=%s raw_print=%s check_formal=%s check_signature=%s",
            len(saved_paths),
            opts.get("run_mode"),
            opts.get("dry_run"),
            opts.get("skip_print"),
            opts.get("raw_print"),
            opts.get("check_formal"),
            opts.get("check_signature"),
        )
        result = run_batch(
            saved_paths,
            recursive=False,
            check_formal=opts["check_formal"],
            check_signature=opts["check_signature"],
            accept_revisions=opts["accept_revisions"],
            word_content_preserve=opts["word_content_preserve"],
            word_preserve_page_count=opts["word_preserve_page_count"],
            word_image_risk_guard=opts["word_image_risk_guard"],
            word_step_timeout_sec=opts.get("word_step_timeout_sec"),
            word_skip_file_on_timeout=opts.get("word_skip_file_on_timeout"),
            file_timeout_sec=opts.get("file_timeout_sec"),
            word_font_profile=opts.get("doc_font_profile", "mixed"),
            printer_name=opts["printer_name"],
            copies=opts["copies"],
            dry_run=opts["dry_run"],
            skip_print=opts["skip_print"],
            raw_print=opts["raw_print"],
            progress_callback=progress_callback,
            # cancel/pause 来自 job
            cancel_event=cancel_event,
            pause_event=pause_event,
            incremental_output_dir=opts.get("incremental_output_dir"),
            relative_names=original_names,
            incremental_exists_action=opts.get("incremental_exists_action", "overwrite"),
            skip_current_event=skip_current_event,
        )
        for i, d in enumerate(result["details"]):
            d["filename"] = original_names[i] if i < len(original_names) else os.path.basename(d["path"])
        logger.info(
            "batch worker done total=%s ok=%s failed=%s elapsed=%ss summary=%s",
            result.get("total"),
            result.get("ok"),
            result.get("failed"),
            result.get("batch_elapsed_sec"),
            result.get("eta_summary"),
        )
        if opts.get("skip_print") and result.get("total"):
            try:
                token = uuid.uuid4().hex
                # 取消时仅打包已完成的文件（details 中已有顺序）；未完成的不纳入。
                done_n = len(result.get("details") or [])
                pack_paths = saved_paths[:done_n]
                pack_names = original_names[:done_n]
                pack_backups = original_backup_paths[:done_n]
                mod_txt = build_batch_modification_zip_text(result["details"])
                _, zip_ftp_ok, zip_ftp_err = _zip_batch_exports(
                    result["details"],
                    pack_paths,
                    pack_names,
                    pack_backups,
                    token,
                    modification_report_text=mod_txt,
                )
                result["download_token"] = token
                result["download_filename"] = "processed_documents.zip"
                result["zip_ftp_uploaded"] = bool(zip_ftp_ok)
                if zip_ftp_err:
                    result["zip_ftp_error"] = zip_ftp_err
            except Exception as e:
                logger.exception("batch zip failed: %s", e)
                queue.put(("result", {"ok": False, "error": _format_com_error(e)}))
                return
        # 本轮（local）也提供终版下载：上传成功项终版到 FTP 并下发 final_download_key
        try:
            details = result.get("details") or []
            if isinstance(details, list) and details:
                key = uuid.uuid4().hex
                from ftp_store import try_upload_file

                items = []
                for i, d in enumerate(details):
                    if not isinstance(d, dict):
                        continue
                    proc = (d.get("processed_path") or "").strip()
                    if not d.get("success") or not proc or (not os.path.isfile(proc)):
                        continue
                    name = original_names[i] if i < len(original_names) else (d.get("filename") or d.get("path") or f"file_{i+1}")
                    arcname = _zip_arcname_with_processed_ext(str(name), proc)
                    try:
                        ftp_p, up_err = try_upload_file(proc, f"batch/final/{key}/{arcname}")
                        if ftp_p:
                            d["final_ftp_path"] = ftp_p
                            items.append({"ftp_path": ftp_p, "filename": os.path.basename(arcname) or f"file_{i+1}"})
                        elif up_err:
                            d["final_ftp_error"] = up_err
                    except Exception:
                        pass
                if items:
                    _FINAL_DOWNLOAD_CACHE[key] = {"created_at": time.time(), "items": items}
                    result["final_download_key"] = key
        except Exception:
            pass
        queue.put(("result", {"ok": True, "result": result}))
    except Exception as e:
        logger.exception("batch worker failed: %s", e)
        queue.put(("result", {"ok": False, "error": _format_com_error(e)}))
    finally:
        try:
            import pythoncom
            pythoncom.CoUninitialize()
        except Exception:
            pass
        queue.put((None, None))  # ????

        if job and job_id:
            try:
                _BATCH_JOB_REGISTRY.pop(job_id, None)
            except Exception:
                pass


def _sse_message(event, data):
    """Build one SSE message."""
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


def _sse_batch_stream_body(queue, tmp_dir, stream_state_ref, heartbeat_sec=120.0):
    """从队列读事件并 yield SSE；结束时持久化历史并删除 tmp。heartbeat_sec：无进度时推送间隔（秒）。"""
    hb = max(10.0, min(600.0, float(heartbeat_sec)))
    started_at = time.time()
    last_emit = time.time()
    try:
        while True:
            try:
                event_type, payload = queue.get(timeout=hb)
            except Empty:
                waited = int(time.time() - last_emit)
                yield _sse_message(
                    "heartbeat",
                    {
                        "ts": int(time.time()),
                        "heartbeatSec": int(hb),
                        "message": f"仍在处理中（已 {waited} 秒无进度推送，心跳间隔 {int(hb)} 秒）",
                        "runningForSec": int(time.time() - started_at),
                    },
                )
                continue
            if event_type is None:
                break
            if event_type == "progress":
                last_emit = time.time()
                yield _sse_message("progress", payload)
            elif event_type == "result":
                last_emit = time.time()
                stream_state_ref["final_json"] = payload
                yield _sse_message("result", payload)
                break
    finally:
        try:
            _batch_history_persist(stream_state_ref, tmp_dir)
        except Exception:
            logger.exception("batch history persist error")
        if tmp_dir and os.path.isdir(tmp_dir):
            try:
                shutil.rmtree(tmp_dir, ignore_errors=True)
            except Exception:
                pass


@app.route("/api/batch-control/<job_id>", methods=["POST"])
def api_batch_control(job_id):
    """暂停/继续/取消当前批处理任务。"""
    if not _BATCH_EXPORT_TOKEN_RE.match(job_id or ""):
        return jsonify({"ok": False, "error": "invalid job id"}), 404
    job = _BATCH_JOB_REGISTRY.get(job_id)
    if not job:
        return jsonify({"ok": False, "error": "job not found or finished"}), 404
    # 不能写成 form.get or json.get if is_json else ''：multipart 时 is_json 为 False，整式会变成 ''，导致永远 invalid action
    if request.is_json:
        raw = (request.get_json(silent=True) or {}).get("action")
    else:
        raw = request.form.get("action")
    action = str(raw or "").strip().lower()
    if action == "cancel":
        job.cancel_event.set()
        job.pause_event.clear()
        return jsonify({"ok": True, "jobId": job_id, "state": "cancelling"})
    if action == "pause":
        job.pause_event.set()
        return jsonify({"ok": True, "jobId": job_id, "state": "paused"})
    if action in ("resume", "continue"):
        job.pause_event.clear()
        return jsonify({"ok": True, "jobId": job_id, "state": "running"})
    if action in ("skip_current", "skip_file"):
        job.skip_current_event.set()
        return jsonify({"ok": True, "jobId": job_id, "state": "skip_requested"})
    return jsonify({"ok": False, "error": "invalid action"}), 400


@app.route("/api/batch-print-stream", methods=["POST"])
def api_batch_print_stream():
    """
    Streaming batch print via text/event-stream.
    progress: { step, fileIndex, fileTotal, fileName, percent, eta?, processingMode?, processingModeLabel?, modificationPoints? }
    result: { ok, result? or error? }
    """
    files = request.files.getlist("files") or request.files.getlist("files[]")
    if not files or not any(f.filename for f in files):
        return jsonify({"ok": False, "error": "?????????"}), 400

    try:
        opts = _parse_batch_opts_from_form(request.form)
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400

    tmp_dir = None
    try:
        tmp_dir = tempfile.mkdtemp(prefix="aiprintword_")
        saved_paths = []
        original_names = []
        original_backup_paths = []
        for f in files:
            rel = _upload_relpath(f.filename)
            if not rel:
                continue
            path = os.path.join(tmp_dir, *rel.split("/"))
            os.makedirs(os.path.dirname(path), exist_ok=True)
            f.save(path)
            if _is_archive_upload(rel):
                for p2, r2 in _expand_archive_for_batch(tmp_dir, path, rel):
                    saved_paths.append(p2)
                    original_names.append(r2)
                    bk = _write_original_backup(tmp_dir, r2, p2)
                    original_backup_paths.append(bk)
            elif _allowed_file(rel):
                saved_paths.append(path)
                original_names.append(rel)
                bk = _write_original_backup(tmp_dir, rel, path)
                original_backup_paths.append(bk)

        if not saved_paths:
            return jsonify(
                {
                    "ok": False,
                    "error": "未找到可处理文件：压缩包内需含 .doc/.docx/.xls/.xlsx/.pdf 等文档；ZIP 异常时可安装 7-Zip 后重试；.7z/.rar 需 7-Zip 或 py7zr/UnRAR。",
                }
            ), 400
        logger.info("stream request accepted files=%s", len(saved_paths))

        job_id = uuid.uuid4().hex
        job = _BatchJob(job_id=job_id, cancel_event=threading.Event(), pause_event=threading.Event())
        _BATCH_JOB_REGISTRY[job_id] = job

        queue = Queue()
        opts["_job"] = job
        thread = threading.Thread(
            target=_run_batch_with_progress,
            args=(saved_paths, original_names, original_backup_paths, opts, queue),
            daemon=True,
        )
        thread.start()

        stream_state = {
            "opts": {k: v for k, v in opts.items() if not str(k).startswith("_")},
            "saved_paths": saved_paths,
            "original_names": original_names,
            "original_backup_paths": original_backup_paths,
            "final_json": None,
        }
        heartbeat_sec = _parse_sse_heartbeat_sec(request.form)
        logger.info("batch-print-stream sse heartbeat interval=%ss", int(heartbeat_sec))

        def generate_with_job():
            yield _sse_message("job", {"jobId": job_id})
            yield from _sse_batch_stream_body(
                queue, tmp_dir, stream_state, heartbeat_sec=heartbeat_sec
            )

        return Response(
            generate_with_job(),
            mimetype="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
                "Connection": "keep-alive",
            },
        )
    except Exception as e:
        if tmp_dir and os.path.isdir(tmp_dir):
            try:
                shutil.rmtree(tmp_dir, ignore_errors=True)
            except Exception:
                pass
        return jsonify({"ok": False, "error": _format_com_error(e)}), 500


def _history_entries_ensure_title(entries):
    """旧版 index 摘要无 title 时，按需加载 record 补全。"""
    from batch_history_mysql import compute_display_title

    out = []
    for e in entries:
        if not isinstance(e, dict):
            continue
        if e.get("title"):
            # 即便已有 title，也可能存在 zip_ftp/zip_ftp_error 字段过期（例如迁移脚本更新了 record.json）
            e2 = dict(e)
            hid = e2.get("id")
            rec = _load_history_record(hid) if hid else None
            if rec:
                try:
                    r = rec.get("result") if isinstance(rec, dict) else None
                    r = r if isinstance(r, dict) else {}
                    zup = r.get("zip_ftp_uploaded")
                    if zup is None:
                        zup = rec.get("zip_ftp_uploaded") if isinstance(rec, dict) else None
                    if zup is True or zup == 1:
                        e2["zip_ftp"] = True
                    elif zup is False or zup == 0:
                        e2["zip_ftp"] = False
                    # 只在 record 有明确值时覆盖 error，避免把前端已有提示清空
                    ze = (r.get("zip_ftp_error") or rec.get("zip_ftp_error") or "").strip() if isinstance(rec, dict) else ""
                    if ze:
                        e2["zip_ftp_error"] = ze
                    else:
                        # 如果 record 里明确为空，则清空
                        if "zip_ftp_error" in e2:
                            e2["zip_ftp_error"] = None
                except Exception:
                    pass
            out.append(e2)
            continue
        e2 = dict(e)
        hid = e2.get("id")
        rec = _load_history_record(hid) if hid else None
        if rec:
            e2["title"] = rec.get("display_title") or compute_display_title(
                rec.get("original_names"), rec.get("result")
            )
            # 补全 zip_ftp/zip_ftp_error（从 record 派生，避免 index.json 过期）
            try:
                r = rec.get("result") if isinstance(rec, dict) else None
                r = r if isinstance(r, dict) else {}
                zup = r.get("zip_ftp_uploaded")
                if zup is None:
                    zup = rec.get("zip_ftp_uploaded") if isinstance(rec, dict) else None
                if zup is True or zup == 1:
                    e2["zip_ftp"] = True
                elif zup is False or zup == 0:
                    e2["zip_ftp"] = False
                ze = (r.get("zip_ftp_error") or rec.get("zip_ftp_error") or "").strip() if isinstance(rec, dict) else ""
                e2["zip_ftp_error"] = ze or None
            except Exception:
                pass
        else:
            rm = e2.get("run_mode") or "批处理"
            tot = int(e2.get("total") or 0)
            n = tot if tot > 0 else 1
            e2["title"] = f"{rm}等{n}份文件处理记录"
        out.append(e2)
    return out


@app.route("/api/batch-history", methods=["GET"])
def api_batch_history_list():
    idx = _load_history_index()
    entries = _history_entries_ensure_title(idx.get("entries", []))
    return jsonify({"ok": True, "entries": entries})


@app.route("/api/batch-history/<hid>", methods=["GET"])
def api_batch_history_get(hid):
    rec = _load_history_record(hid)
    if not rec:
        return jsonify({"ok": False, "error": "记录不存在"}), 404
    return jsonify({"ok": True, "record": rec})


@app.route("/api/batch-history/<hid>/download", methods=["GET"])
def api_batch_history_download(hid):
    rec = _load_history_record(hid)
    if not rec:
        return jsonify({"ok": False, "error": "记录不存在"}), 404
    token = rec.get("download_token")
    if not token or not _BATCH_EXPORT_TOKEN_RE.match(str(token)):
        return jsonify({"ok": False, "error": "该次任务无打包下载"}), 404
    path = os.path.join(BATCH_EXPORT_ROOT, str(token) + ".zip")
    if not os.path.isfile(path):
        try:
            from ftp_store import download_bytes

            b = download_bytes(f"batch/exports/{token}.zip")
            os.makedirs(BATCH_EXPORT_ROOT, exist_ok=True)
            with open(path, "wb") as fp:
                fp.write(b)
        except Exception:
            return jsonify({"ok": False, "error": "文件已删除或已过期"}), 404
    return send_file(
        path,
        as_attachment=True,
        download_name="processed_documents.zip",
        max_age=0,
    )


@app.route("/api/batch-history/<hid>/final/<int:idx>", methods=["GET"])
def api_batch_history_final_download(hid, idx: int):
    """下载某次历史任务的单文件终版成品（来自 FTP）。"""
    rec = _load_history_record(hid)
    if not rec:
        return jsonify({"ok": False, "error": "记录不存在"}), 404
    res = (rec.get("result") or {}) if isinstance(rec, dict) else {}
    details = res.get("details") or []
    if not isinstance(details, list) or idx < 0 or idx >= len(details):
        return jsonify({"ok": False, "error": "无效的文件索引"}), 404
    d = details[idx] if isinstance(details[idx], dict) else {}
    ftp_p = (d.get("final_ftp_path") or "").strip()
    if not ftp_p:
        return jsonify({"ok": False, "error": "该文件无可下载终版"}), 404
    try:
        from ftp_store import download_bytes

        b = download_bytes(ftp_p)
    except Exception as e:
        return jsonify({"ok": False, "error": "FTP 下载失败"}), 502
    # 下载文件名：优先 filename/path，再退回 idx
    nm = (d.get("filename") or d.get("path") or f"file_{idx+1}").replace("\\", "/")
    base = os.path.basename(nm) or f"file_{idx+1}"
    return send_file(
        io.BytesIO(b),
        as_attachment=True,
        download_name=base,
        mimetype="application/octet-stream",
        max_age=0,
    )


@app.route("/api/batch-final/<key>/<int:idx>", methods=["GET"])
def api_batch_final_download_local(key: str, idx: int):
    """下载本轮（local）终版成品：通过内存 key 映射到 FTP 路径。"""
    if not _BATCH_EXPORT_TOKEN_RE.match(key or ""):
        return jsonify({"ok": False, "error": "invalid key"}), 404
    rec = _FINAL_DOWNLOAD_CACHE.get(key)
    if not rec:
        return jsonify({"ok": False, "error": "该下载已过期，请刷新历史列表后从历史记录下载"}), 404
    items = rec.get("items") or []
    if idx < 0 or idx >= len(items):
        return jsonify({"ok": False, "error": "无效的文件索引"}), 404
    it = items[idx] if isinstance(items[idx], dict) else {}
    ftp_p = (it.get("ftp_path") or "").strip()
    fn = (it.get("filename") or f"file_{idx+1}").strip()
    if not ftp_p:
        return jsonify({"ok": False, "error": "无可下载终版"}), 404
    try:
        from ftp_store import download_bytes

        b = download_bytes(ftp_p)
    except Exception:
        return jsonify({"ok": False, "error": "FTP 下载失败"}), 502
    return send_file(
        io.BytesIO(b),
        as_attachment=True,
        download_name=fn,
        mimetype="application/octet-stream",
        max_age=0,
    )


@app.route("/api/batch-history/<hid>/retry", methods=["POST"])
def api_batch_history_retry(hid):
    from werkzeug.datastructures import CombinedMultiDict

    rec = _load_history_record(hid)
    if not rec:
        return jsonify({"ok": False, "error": "记录不存在"}), 404
    if not rec.get("has_stash"):
        return jsonify({"ok": False, "error": "无可用暂存，请重新上传后再处理"}), 400
    stash = os.path.join(BATCH_HISTORY_ROOT, hid, "stash")
    if not os.path.isdir(stash):
        return jsonify({"ok": False, "error": "暂存目录已丢失"}), 400
    only_failed = (request.form.get("only_failed") or "true").strip().lower() in (
        "1",
        "true",
        "yes",
        "on",
    )
    defaults = _record_to_default_form(rec)
    combined = CombinedMultiDict([request.form, defaults])
    try:
        opts = _parse_batch_opts_from_form(combined)
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400

    original_names = rec.get("original_names") or []
    details = (rec.get("result") or {}).get("details") or []
    tmp_dir, saved_paths, original_names_ord, original_backup_paths = _materialize_retry_workspace(
        stash, original_names, details, only_failed
    )
    if not tmp_dir or not saved_paths:
        return jsonify({"ok": False, "error": "未能从暂存中恢复待重试文件"}), 400

    logger.info(
        "history retry hid=%s only_failed=%s docs=%s",
        hid,
        only_failed,
        len(saved_paths),
    )
    job_id = uuid.uuid4().hex
    job = _BatchJob(job_id=job_id, cancel_event=threading.Event(), pause_event=threading.Event())
    _BATCH_JOB_REGISTRY[job_id] = job
    queue = Queue()
    opts["_job"] = job
    thread = threading.Thread(
        target=_run_batch_with_progress,
        args=(saved_paths, original_names_ord, original_backup_paths, opts, queue),
        daemon=True,
    )
    thread.start()

    stream_state = {
        "opts": {k: v for k, v in opts.items() if not str(k).startswith("_")},
        "saved_paths": saved_paths,
        "original_names": original_names_ord,
        "original_backup_paths": original_backup_paths,
        "final_json": None,
    }
    heartbeat_sec = _parse_sse_heartbeat_sec(combined)

    def generate_with_job():
        yield _sse_message("job", {"jobId": job_id})
        yield from _sse_batch_stream_body(
            queue, tmp_dir, stream_state, heartbeat_sec=heartbeat_sec
        )

    return Response(
        generate_with_job(),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


# ---------- ????????????????????----------
SIGN_ALLOWED_EXT = {".docx", ".xlsx"}
SIGN_WORD_SOURCE_EXT = {".doc", ".docx", ".docm"}
SIGN_EXCEL_SOURCE_EXT = {".xls", ".xlsx", ".xlsm"}
SIGN_SOURCE_DOC_EXT = SIGN_WORD_SOURCE_EXT | SIGN_EXCEL_SOURCE_EXT
SIGN_ARCHIVE_ALLOWED_EXT = {".zip", ".7z", ".rar"}
SIGN_UPLOAD_ALLOWED_EXT = SIGN_SOURCE_DOC_EXT | SIGN_ARCHIVE_ALLOWED_EXT
SIGN_INBOX_ROOT = os.path.join(ROOT, "data", "sign_inbox")
SIGN_DETECT_CORRECTIONS_ROOT = os.path.join(ROOT, "data", "sign_detect_corrections")
SIGN_MAX_SAVED_FILES = 50
_SIGN_FILE_ID_RE = re.compile(r"^[0-9a-f]{32}$")


def _sign_using_mysql():
    """Use MySQL store when MYSQL_HOST is configured."""
    try:
        from sign_handlers import mysql_store

        return mysql_store.mysql_sign_enabled()
    except Exception:
        return False


def _sign_ensure_session_inbox():
    """Ensure per-session inbox directory exists."""
    if "sign_inbox_sid" not in session:
        session["sign_inbox_sid"] = uuid.uuid4().hex
    if "sign_files" not in session:
        session["sign_files"] = []
    sid = session["sign_inbox_sid"]
    inbox_dir = os.path.join(SIGN_INBOX_ROOT, sid)
    os.makedirs(inbox_dir, exist_ok=True)
    return sid, inbox_dir


def _sign_saved_disk_path(sid: str, file_id: str, ext: str) -> str:
    return os.path.join(SIGN_INBOX_ROOT, sid, file_id + ext.lower())


def _parse_detect_correction_image_upload():
    """解析参考图上传请求，返回 (blob, ext, display_name) 或 (None, error_msg)。"""
    f = request.files.get("file") or request.files.get("image")
    if not f or not getattr(f, "filename", None):
        return None, None, None, "未选择图片"
    raw_name = os.path.basename(str(f.filename).replace("\\", "/"))
    ext = os.path.splitext(raw_name)[1].lower()
    if ext not in {".png", ".jpg", ".jpeg", ".webp"}:
        return None, None, None, "仅支持 PNG/JPEG/WebP"
    blob = f.read()
    if not blob:
        return None, None, None, "空文件"
    if len(blob) > 2 * 1024 * 1024:
        return None, None, None, "单张图片不能超过 2MB"
    return blob, ext, raw_name, None


def _get_file_detect_correction(file_id: str) -> dict:
    from sign_handlers.file_session_cache import trim_detect_correction

    if _sign_using_mysql():
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        data = mysql_store.get_file_detect_correction(file_id)
    else:
        _sign_ensure_session_inbox()
        sid = session["sign_inbox_sid"]
        from sign_handlers.sign_library_local import get_file_detect_correction as local_get

        data = local_get(SIGN_INBOX_ROOT, sid, file_id)
    return trim_detect_correction(data) if data else {}


def _get_sign_file_source_name(file_id: str) -> str:
    """待签文件展示名（用于规则 pattern 推断）。"""
    if _sign_using_mysql():
        try:
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            row = mysql_store.get_file_row(file_id)
            if row:
                return str(row.get("name") or "").strip()
        except Exception:
            pass
    else:
        try:
            _sign_ensure_session_inbox()
            rec = _sign_find_record(file_id)
            if rec:
                return str(rec.get("name") or "").strip()
        except Exception:
            pass
    return ""


def _sync_detect_correction_to_rules(
    file_id: str, correction: dict, *, export_md: bool = False
) -> dict:
    """将标误登记同步到角色/签字位规则 JSON；MD 导出由调用方批量合并触发。"""
    try:
        from sign_handlers.detect_correction_rules import sync_rules_from_correction

        src = _get_sign_file_source_name(file_id)
        if not src:
            return {"ok": False, "error": "无法获取文件名"}
        return sync_rules_from_correction(src, correction, export_md=export_md)
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _export_detect_correction_markdown_after_sync(corrections: list) -> dict:
    """标误保存后至多导出一次角色/签字位 MD，避免每条登记启子进程卡住请求。"""
    need_role_md = False
    need_slot_md = False
    for corr in corrections or []:
        if not isinstance(corr, dict):
            continue
        scopes = (
            corr.get("correction_save")
            if isinstance(corr.get("correction_save"), dict)
            else {}
        )
        save_roles = scopes.get("roles", True) if scopes else True
        save_slot = scopes.get("slot", True) if scopes else True
        wrong = str(corr.get("wrong_description") or "").strip()
        esl = corr.get("expected_slot_layout")
        has_slot = isinstance(esl, dict) and bool(esl)
        if save_roles and wrong:
            need_role_md = True
        if save_slot and (has_slot or wrong):
            need_slot_md = True
    if not need_role_md and not need_slot_md:
        return {"skipped": True}
    out: dict = {}

    def _run_exports() -> None:
        if need_role_md:
            try:
                from sign_handlers.detect_correction_rules import export_rules_markdown

                out["role_md"] = export_rules_markdown()
            except Exception as e:
                out["role_md"] = {"ok": False, "error": str(e)[:500]}
        if need_slot_md:
            try:
                from sign_handlers.detect_correction_slot_rules import (
                    export_slot_layout_markdown,
                )

                out["slot_md"] = export_slot_layout_markdown()
            except Exception as e:
                out["slot_md"] = {"ok": False, "error": str(e)[:500]}

    skip_bg = (os.environ.get("SIGN_DETECT_CORRECTION_EXPORT_MD_SYNC") or "").strip().lower() in (
        "1",
        "true",
        "yes",
    )
    if skip_bg:
        _run_exports()
        return out
    import threading

    threading.Thread(target=_run_exports, daemon=True).start()
    return {"async": True, "need_role_md": need_role_md, "need_slot_md": need_slot_md}


def _set_file_detect_correction(file_id: str, correction: dict) -> None:
    from sign_handlers.file_session_cache import trim_detect_correction

    payload = trim_detect_correction(correction)
    if _sign_using_mysql():
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        mysql_store.set_file_detect_correction(file_id, payload)
    else:
        _sign_ensure_session_inbox()
        sid = session["sign_inbox_sid"]
        from sign_handlers.sign_library_local import set_file_detect_correction as local_set

        local_set(SIGN_INBOX_ROOT, sid, file_id, payload)


def _sign_upload_display_name(client_filename, allowed_exts=None):
    """
    ??multipart ????????????????????????????    ??.doc/.docx/.docm/.xls/.xlsx/.xlsm ?? (None, None)??    """
    if not client_filename:
        return None, None
    norm = str(client_filename).replace("\\", "/").strip()
    if ".." in norm:
        norm = os.path.basename(norm)
    parts = [p for p in norm.split("/") if p and p not in (".", "..")]
    if not parts:
        return None, None
    last = parts[-1]
    ext = os.path.splitext(last)[1].lower()
    allowed = allowed_exts or SIGN_ALLOWED_EXT
    if ext not in allowed:
        return None, None
    display = "/".join(parts)
    if len(display) > 200:
        display = display[-200:]
    return display, ext


def _sign_extract_zip_upload_items(zip_display_name: str, zip_bytes: bytes) -> tuple[list, list, dict]:
    """从 zip 字节中提取签字候选文档（Word/Excel）。"""
    out = []
    warnings = []
    stats = {
        "archive_name": zip_display_name,
        "archive_ext": ".zip",
        "total_members": 0,
        "source_members": 0,
        "added_candidates": 0,
        "skipped_members": 0,
        "skipped_by_ext": {},
    }

    def _inc_skip(ext_key: str):
        ek = ext_key or "(no_ext)"
        stats["skipped_by_ext"][ek] = int(stats["skipped_by_ext"].get(ek) or 0) + 1
        stats["skipped_members"] += 1

    try:
        with zipfile.ZipFile(io.BytesIO(zip_bytes), "r") as zf:
            total_uncompressed = 0
            hit_file_cap = False
            hit_total_cap = False
            for info in zf.infolist():
                if info.is_dir():
                    continue
                stats["total_members"] += 1
                if len(out) >= _ARCHIVE_EXTRACT_MAX_FILES:
                    hit_file_cap = True
                    break
                inner = _zip_inner_path_safe(info.filename)
                if not inner:
                    _inc_skip("(invalid_path)")
                    continue
                ext = os.path.splitext(inner)[1].lower()
                if ext not in SIGN_SOURCE_DOC_EXT:
                    _inc_skip(ext or "(no_ext)")
                    continue
                stats["source_members"] += 1
                fsize = int(getattr(info, "file_size", 0) or 0)
                if fsize <= 0:
                    _inc_skip(ext or "(empty)")
                    continue
                if fsize > _ARCHIVE_EXTRACT_SINGLE_MAX:
                    warnings.append(f"压缩包成员过大，已跳过：{inner}")
                    _inc_skip(ext or "(oversize)")
                    continue
                if total_uncompressed + fsize > _ARCHIVE_EXTRACT_MAX_TOTAL_UNCOMPRESSED:
                    hit_total_cap = True
                    break
                try:
                    raw = zf.read(info)
                except Exception as e:
                    warnings.append(f"读取压缩包成员失败，已跳过：{inner}（{e}）")
                    _inc_skip(ext or "(read_error)")
                    continue
                if not raw:
                    _inc_skip(ext or "(empty)")
                    continue
                if len(raw) > _ARCHIVE_EXTRACT_SINGLE_MAX:
                    warnings.append(f"压缩包成员解压后超限，已跳过：{inner}")
                    _inc_skip(ext or "(oversize)")
                    continue
                total_uncompressed += len(raw)
                disp = inner.replace("\\", "/").strip()
                if not disp:
                    _inc_skip(ext or "(invalid_path)")
                    continue
                disp = _sign_upload_base_name(disp)
                if len(disp) > 200:
                    disp = disp[-200:]
                out.append({"name": disp, "ext": ext, "raw": raw})
                stats["added_candidates"] += 1
            if hit_file_cap:
                warnings.append(f"压缩包可提取文件数超过上限（{_ARCHIVE_EXTRACT_MAX_FILES}），其余已跳过")
            if hit_total_cap:
                warnings.append("压缩包解压总量超过上限，部分成员已跳过")
    except (zipfile.BadZipFile, OSError) as e:
        warnings.append(f"压缩包无法读取：{zip_display_name}（{e}）")
    return out, warnings, stats


def _sign_finalize_collected_items(items: list) -> list:
    """入库/判重统一用展示名基名（与 insert_file 的 normalize 一致）。"""
    for item in items or []:
        if not isinstance(item, dict):
            continue
        item["name"] = _sign_upload_base_name(str(item.get("name") or ""))
    return items


def _sign_extract_archive_upload_items(
    archive_display_name: str, archive_bytes: bytes, archive_ext: str
) -> tuple[list, list, dict]:
    """从压缩包字节中提取签字候选文档（Word/Excel）。"""
    ext = (archive_ext or "").lower().strip()
    if ext == ".zip":
        return _sign_extract_zip_upload_items(archive_display_name, archive_bytes)

    warnings = []
    out = []
    stats = {
        "archive_name": archive_display_name,
        "archive_ext": ext or "(unknown)",
        "total_members": 0,
        "source_members": 0,
        "added_candidates": 0,
        "skipped_members": 0,
        "skipped_by_ext": {},
    }
    tmp_dir = tempfile.mkdtemp(prefix="aiprintword_sign_archive_")
    archive_path = os.path.join(tmp_dir, "upload" + ext)
    try:
        with open(archive_path, "wb") as fp:
            fp.write(archive_bytes or b"")
        expanded = _expand_archive_for_batch(tmp_dir, archive_path, archive_display_name)
        for p, rel in expanded:
            try:
                stats["total_members"] += 1
                e = os.path.splitext(str(rel or ""))[1].lower()
                if e not in SIGN_SOURCE_DOC_EXT:
                    ek = e or "(no_ext)"
                    stats["skipped_by_ext"][ek] = int(stats["skipped_by_ext"].get(ek) or 0) + 1
                    stats["skipped_members"] += 1
                    continue
                stats["source_members"] += 1
                if not p or not os.path.isfile(p):
                    stats["skipped_members"] += 1
                    stats["skipped_by_ext"][e or "(missing)"] = int(
                        stats["skipped_by_ext"].get(e or "(missing)") or 0
                    ) + 1
                    continue
                raw = Path(p).read_bytes()
                if not raw:
                    stats["skipped_members"] += 1
                    stats["skipped_by_ext"][e or "(empty)"] = int(
                        stats["skipped_by_ext"].get(e or "(empty)") or 0
                    ) + 1
                    continue
                if len(raw) > _ARCHIVE_EXTRACT_SINGLE_MAX:
                    warnings.append(f"压缩包成员解压后超限，已跳过：{rel}")
                    stats["skipped_members"] += 1
                    stats["skipped_by_ext"][e or "(oversize)"] = int(
                        stats["skipped_by_ext"].get(e or "(oversize)") or 0
                    ) + 1
                    continue
                disp = str(rel or "").replace("\\", "/").strip()
                if not disp:
                    stats["skipped_members"] += 1
                    stats["skipped_by_ext"][e or "(invalid_path)"] = int(
                        stats["skipped_by_ext"].get(e or "(invalid_path)") or 0
                    ) + 1
                    continue
                disp = _sign_upload_base_name(disp)
                if len(disp) > 200:
                    disp = disp[-200:]
                out.append({"name": disp, "ext": e, "raw": raw})
                stats["added_candidates"] += 1
            except Exception as ie:
                warnings.append(f"读取压缩包成员失败，已跳过：{rel}（{ie}）")
                ek = os.path.splitext(str(rel or ""))[1].lower() or "(read_error)"
                stats["skipped_members"] += 1
                stats["skipped_by_ext"][ek] = int(stats["skipped_by_ext"].get(ek) or 0) + 1
        if not out:
            if ext in (".7z", ".rar") and not _resolve_7zip_cli():
                warnings.append(
                    f"未检测到 7-Zip 命令行，无法解包：{archive_display_name}（请安装 7z/7za 并加入 PATH）"
                )
            else:
                warnings.append(f"压缩包内未找到可签字文档：{archive_display_name}")
    except Exception as e:
        warnings.append(f"压缩包无法读取：{archive_display_name}（{e}）")
    finally:
        try:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass
    return out, warnings, stats


def _sign_convert_source_item_to_signable(item: dict) -> tuple[Optional[dict], Optional[str]]:
    """将 .doc/.docm/.xls/.xlsm 转为 .docx/.xlsx；原生 .docx/.xlsx 直接透传。"""
    name = str((item or {}).get("name") or "").strip()
    ext = str((item or {}).get("ext") or "").lower().strip()
    raw = (item or {}).get("raw")
    if not name or not ext or not isinstance(raw, (bytes, bytearray)):
        return None, "无效上传项"
    if ext in SIGN_ALLOWED_EXT:
        return {"name": name, "ext": ext, "raw": bytes(raw)}, None

    tmp_dir = tempfile.mkdtemp(prefix="aiprintword_sign_convert_")
    try:
        src_name = os.path.basename(name.replace("\\", "/")) or ("document" + ext)
        src_path = os.path.join(tmp_dir, src_name)
        with open(src_path, "wb") as fp:
            fp.write(bytes(raw))

        out_path = None
        if ext in SIGN_WORD_SOURCE_EXT:
            from doc_handlers.word_handler import convert_doc_to_docx

            out_path = convert_doc_to_docx(src_path)
        elif ext in SIGN_EXCEL_SOURCE_EXT:
            from doc_handlers.excel_handler import convert_xls_to_xlsx

            out_path = convert_xls_to_xlsx(src_path)
        else:
            return None, f"暂不支持的文档类型：{ext}"

        if not out_path or not os.path.isfile(out_path):
            return None, f"转换失败（未生成输出文件）：{name}"

        out_ext = os.path.splitext(out_path)[1].lower()
        if out_ext not in SIGN_ALLOWED_EXT:
            return None, f"转换后格式不支持：{name} -> {out_ext}"
        out_raw = Path(out_path).read_bytes()
        if not out_raw:
            return None, f"转换结果为空：{name}"

        stem = os.path.splitext(name)[0]
        disp = (stem + out_ext).replace("\\", "/")
        if len(disp) > 200:
            disp = disp[-200:]
        return {"name": disp, "ext": out_ext, "raw": out_raw}, None
    except Exception as e:
        return None, f"文档转换失败：{name}（{e}）"
    finally:
        try:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass


def _sign_upload_base_name(name: str) -> str:
    from sign_handlers.filename_util import normalize_display_filename

    return normalize_display_filename(name or "")


def _sign_record_same_project(rec: dict, proj: dict) -> bool:
    want = str((proj or {}).get("id") or "").strip()
    got = str((rec or {}).get("project_id") or "").strip()
    if want:
        return got == want
    return not got


def _sign_session_capture_replace_caches(
    sid: str, inbox_dir: str, proj: dict, parsed: list
) -> dict:
    """Session 模式：覆盖上传前抓取旧 file_id 的工作台缓存（按 basename）。"""
    from sign_handlers.sign_library_local import get_file_role_map, list_file_session_caches

    bases = {
        _sign_upload_base_name(str(item.get("name") or ""))
        for item in (parsed or [])
        if str(item.get("name") or "").strip()
    }
    bases.discard("")
    if not bases:
        return {}
    caches_all = list_file_session_caches(inbox_dir, sid, lite=False)
    captured: dict = {}
    for rec in session.get("sign_files") or []:
        if not isinstance(rec, dict) or not rec.get("id"):
            continue
        if not _sign_record_same_project(rec, proj):
            continue
        base = _sign_upload_base_name(str(rec.get("name") or ""))
        if base not in bases:
            continue
        fid = str(rec["id"])
        ent = dict(caches_all.get(fid) or {})
        role_map = get_file_role_map(inbox_dir, sid, fid)
        if role_map:
            ent["map"] = role_map
        if ent:
            captured[base] = ent
    return captured


def _sign_session_apply_replace_caches(
    sid: str, inbox_dir: str, file_id: str, payload: dict
) -> None:
    if not file_id or not isinstance(payload, dict) or not payload:
        return
    from sign_handlers.sign_library_local import (
        set_file_detect_correction,
        set_file_detect_snapshot,
        set_file_role_map,
        set_file_workbench_state,
    )

    if isinstance(payload.get("detect"), dict) and payload["detect"]:
        set_file_detect_snapshot(inbox_dir, sid, file_id, payload["detect"])
    if isinstance(payload.get("workbench"), dict) and payload["workbench"]:
        set_file_workbench_state(inbox_dir, sid, file_id, payload["workbench"])
    if isinstance(payload.get("detect_correction"), dict) and payload["detect_correction"]:
        set_file_detect_correction(inbox_dir, sid, file_id, payload["detect_correction"])
    if isinstance(payload.get("map"), dict) and payload["map"]:
        set_file_role_map(inbox_dir, sid, file_id, payload["map"])


def _sign_mysql_insert_parsed_items(
    parsed: list, proj: dict, *, replace_duplicates: bool = False
) -> tuple[list, Optional[dict], Optional[str]]:
    """批量写入 MySQL（并行 FTP，缩短大压缩包解压后的落库时间）。"""
    from sign_handlers import mysql_store

    if not parsed:
        return [], None, None
    captured_replace: dict = {}
    if replace_duplicates:
        proj_id = str((proj or {}).get("id") or "") or None
        names = [item.get("name") or "" for item in parsed]
        captured_replace = mysql_store.capture_file_replace_caches_by_basenames_in_project(
            proj_id, names
        )
        mysql_store.delete_inbox_by_basenames_in_project(proj_id, names)
    added_ids: list = []
    last_rec = None
    errors: list = []

    def _insert_one(item: dict) -> tuple[str, dict]:
        file_id = uuid.uuid4().hex
        display_name = item["name"]
        ext = item["ext"]
        raw = item["raw"]
        if not raw:
            raise ValueError(f"文件为空：{display_name}")
        mysql_store.insert_file(
            file_id,
            display_name,
            ext,
            raw,
            project_id=str(proj.get("id") or ""),
            project_name=str(proj.get("name") or ""),
            project_key=str(proj.get("project_key") or proj.get("name") or ""),
        )
        caches = captured_replace.get(_sign_upload_base_name(display_name))
        if caches:
            try:
                mysql_store.apply_file_replace_caches(file_id, caches)
            except Exception as e:
                logger.warning(
                    "sign upload replace cache migrate failed for %s: %s",
                    display_name,
                    e,
                )
        return file_id, {
            "id": file_id,
            "name": display_name,
            "ext": ext,
            "project_id": proj.get("id"),
            "project_name": proj.get("name"),
            "project_key": proj.get("project_key"),
            "project_label": proj.get("label") or proj.get("name"),
        }

    workers = min(4, max(1, len(parsed)))
    if len(parsed) <= 2:
        for item in parsed:
            try:
                fid, rec = _insert_one(item)
                added_ids.append(fid)
                last_rec = rec
            except Exception as e:
                errors.append(str(e))
    else:
        from concurrent.futures import ThreadPoolExecutor, as_completed

        with ThreadPoolExecutor(max_workers=workers) as ex:
            futs = {ex.submit(_insert_one, item): item for item in parsed}
            for fut in as_completed(futs):
                item = futs[fut]
                try:
                    fid, rec = fut.result()
                    added_ids.append(fid)
                    last_rec = rec
                except Exception as e:
                    errors.append(f"{item.get('name')}: {e}")
    if errors and not added_ids:
        return [], None, errors[0]
    if errors:
        return added_ids, last_rec, "部分文件写入失败：" + errors[0][:200]
    return added_ids, last_rec, None


def _sign_collect_upload_items(uploads: list) -> tuple[list, list, bool, int]:
    """统一收集上传项（普通文档 + zip 内文档）。"""
    items = []
    warnings = []
    upload_has_archive = False
    archive_expanded = 0
    archive_summary = {
        "archives": 0,
        "total_members": 0,
        "source_members": 0,
        "added_candidates": 0,
        "added_signable": 0,
        "skipped_members": 0,
        "skipped_by_ext": {},
        "per_archive": [],
    }

    def _merge_skip_counts(dst: dict, src: dict):
        for k, v in (src or {}).items():
            kk = str(k or "(unknown)")
            dst[kk] = int(dst.get(kk) or 0) + int(v or 0)
    for upload in uploads:
        if not upload or not upload.filename:
            continue
        display_name, ext = _sign_upload_display_name(upload.filename, SIGN_UPLOAD_ALLOWED_EXT)
        if not display_name:
            warnings.append(f"已忽略不支持类型：{upload.filename}")
            continue
        raw = upload.read()
        if not raw:
            warnings.append(f"空文件已忽略：{display_name}")
            continue
        if ext in SIGN_SOURCE_DOC_EXT:
            one, err = _sign_convert_source_item_to_signable(
                {"name": display_name, "ext": ext, "raw": raw}
            )
            if one:
                items.append(one)
            elif err:
                warnings.append(err)
            continue
        if ext in SIGN_ARCHIVE_ALLOWED_EXT:
            upload_has_archive = True
            ex_items, ex_warn, ex_stats = _sign_extract_archive_upload_items(display_name, raw, ext)
            archive_summary["archives"] += 1
            archive_summary["total_members"] += int((ex_stats or {}).get("total_members") or 0)
            archive_summary["source_members"] += int((ex_stats or {}).get("source_members") or 0)
            archive_summary["added_candidates"] += int((ex_stats or {}).get("added_candidates") or 0)
            archive_summary["skipped_members"] += int((ex_stats or {}).get("skipped_members") or 0)
            _merge_skip_counts(archive_summary["skipped_by_ext"], (ex_stats or {}).get("skipped_by_ext") or {})
            if ex_stats:
                archive_summary["per_archive"].append(ex_stats)
            if ex_warn:
                warnings.extend(ex_warn)
            if ex_items:
                archive_expanded += len(ex_items)
                for ex in ex_items:
                    one, err = _sign_convert_source_item_to_signable(ex)
                    if one:
                        items.append(one)
                        archive_summary["added_signable"] += 1
                    elif err:
                        warnings.append(err)
            continue
        warnings.append(f"暂不支持该压缩格式：{display_name}")
    _sign_finalize_collected_items(items)
    return items, warnings, upload_has_archive, archive_expanded, archive_summary


def _safe_display_filename_keep_unicode(name: str) -> str:
    """
    保留中文等非 ASCII 字符的“安全显示文件名”净化。
    仅替换 Windows 非法文件名字符，避免 secure_filename 把中文清空。
    """
    s = str(name or "").strip()
    if not s:
        return "document"
    # 仅取最后一段，避免路径穿越/目录显示
    s = os.path.basename(s.replace("\\", "/"))
    # Windows 非法字符：<>:"/\\|?* 以及控制字符
    s = re.sub(r"[<>:\"/\\\\|?*]+", "_", s)
    s = re.sub(r"[\x00-\x1f]+", "_", s)
    # 尾部点/空格在 Windows 里会被裁剪；这里做个兜底
    s = s.rstrip(". ")
    return s or "document"


def _sign_find_record(file_id: str):
    fid = str(file_id or "")
    for rec in session.get("sign_files") or []:
        if str(rec.get("id") or "") == fid:
            return rec
    return None


def _sign_norm_ext(ext: Optional[str]) -> str:
    e = (ext or ".docx").lower().strip()
    if not e.startswith("."):
        e = "." + e
    return e if e in SIGN_ALLOWED_EXT else ".docx"


def _sign_saved_file_exists(sid: str, file_id: str, ext: Optional[str]) -> bool:
    """会话收件箱中该 id 是否仍有对应磁盘文件（兼容扩展名与记录不一致）。"""
    fid = str(file_id or "").strip()
    if not fid or not _SIGN_FILE_ID_RE.match(fid):
        return False
    primary = _sign_norm_ext(ext)
    path = _sign_saved_disk_path(sid, fid, primary)
    if os.path.isfile(path):
        return True
    for alt in SIGN_ALLOWED_EXT:
        if alt == primary:
            continue
        p2 = _sign_saved_disk_path(sid, fid, alt)
        if os.path.isfile(p2):
            return True
    return False


def _sign_record_saved_at_iso(sid: str, rec: dict) -> Optional[str]:
    """待签文件保存时间（列表展示/筛选）；优先记录字段，否则用磁盘 mtime。"""
    for key in ("saved_at", "created_at"):
        raw = str(rec.get(key) or "").strip()
        if raw:
            return raw
    fid = str(rec.get("id") or "").strip()
    if not fid:
        return None
    path = _sign_saved_disk_path(sid, fid, rec.get("ext"))
    if path and os.path.isfile(path):
        try:
            return time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(os.path.getmtime(path)))
        except OSError:
            pass
    return None


def _sign_enrich_session_file_record(sid: str, rec: dict) -> dict:
    out = dict(rec or {})
    ts = _sign_record_saved_at_iso(sid, out)
    if ts:
        out["created_at"] = ts
        out["saved_at"] = ts
    return out


def _sign_prune_session_files_to_disk(sid: str) -> list:
    """
    以磁盘为准修剪 session['sign_files']：已删文件但 Cookie 未写回时，刷新列表仍会从库里删掉条目。
    """
    records = list(session.get("sign_files") or [])
    pruned = []
    for r in records:
        if not _sign_saved_file_exists(sid, str(r.get("id") or ""), r.get("ext")):
            continue
        pruned.append(_sign_enrich_session_file_record(sid, r))
    if len(pruned) != len(records):
        session["sign_files"] = pruned
        session.modified = True
    return pruned


def _sign_remove_disk_files_for_id(sid: str, file_id: str, ext: Optional[str]) -> None:
    """删除收件箱内该 file_id 可能存在的 .docx/.xlsx（避免扩展名记录与磁盘不一致导致删不干净）。"""
    fid = str(file_id or "").strip()
    if not fid:
        return
    tried = set()
    for e in (_sign_norm_ext(ext), ".docx", ".xlsx"):
        if e in tried:
            continue
        tried.add(e)
        path = _sign_saved_disk_path(sid, fid, e)
        if os.path.isfile(path):
            try:
                os.remove(path)
            except OSError:
                pass


def _decode_png_data_url_or_b64(s):
    """Decode canvas PNG data URL or raw base64 string."""
    if not s or not str(s).strip():
        return None
    s = str(s).strip()
    if s.startswith("data:image"):
        parts = s.split(",", 1)
        s = parts[1] if len(parts) > 1 else ""
    try:
        return base64.b64decode(s, validate=False)
    except Exception:
        try:
            return base64.b64decode(s)
        except Exception:
            return None


def _keywords_from_label_preview(text, max_items=6):
    t = str(text or "").strip()
    if not t:
        return []
    out = []
    seen = set()
    for part in re.split(r"[|｜/:：,，;；\s]+", t):
        p = str(part or "").strip()
        if not p:
            continue
        if len(p) < 2 or len(p) > 24:
            continue
        if p in seen:
            continue
        seen.add(p)
        out.append(p)
        if len(out) >= max_items:
            break
    return out


def _parse_docx_table_cell_loc(loc_text: str):
    """table#2.r10.c3 → {table, row, col}（1-based）。"""
    s = str(loc_text or "").strip()
    m = re.match(r"^\s*table#?(\d+)\.r(\d+)\.c(\d+)\s*$", s, re.IGNORECASE)
    if not m:
        return None
    return {
        "table": int(m.group(1)),
        "row": int(m.group(2)),
        "col": int(m.group(3)),
    }


def _parse_xlsx_table_cell_loc(loc_text: str):
    s = str(loc_text or "").strip()
    m = re.match(r"^\s*sheet(\d+)!\s*r(\d+)\.c(\d+)\s*$", s, re.IGNORECASE)
    if not m:
        return None
    return {
        "sheet": int(m.group(1)),
        "row": int(m.group(2)),
        "col": int(m.group(3)),
    }


def _layout_loc_to_source_hints(loc_text):
    """把 signature_layout 的 name_loc/date_loc 统一映射为 sign_engine 可消费的 row hint。"""
    s = str(loc_text or "").strip()
    if not s:
        return []
    out = []
    # docx: table#2.r10.c1 -> table2.row10
    m = re.match(r"^\s*table#?(\d+)\.r(\d+)\.c\d+\s*$", s, re.IGNORECASE)
    if m:
        out.append(f"table{int(m.group(1))}.row{int(m.group(2))}")
    # docx: paragraph#5 -> paragraph5
    m = re.match(r"^\s*paragraph#?(\d+)\s*$", s, re.IGNORECASE)
    if m:
        out.append(f"paragraph{int(m.group(1))}")
    # xlsx: sheet1!r20.c3 -> sheet1.row20
    m = re.match(r"^\s*sheet(\d+)!\s*r(\d+)\.c\d+\s*$", s, re.IGNORECASE)
    if m:
        out.append(f"sheet{int(m.group(1))}.row{int(m.group(2))}")
    # 兜底：已有提示格式直接透传
    if re.match(r"^\s*(?:sheet\d+|table\d+)\.row\d+\s*$", s, re.IGNORECASE):
        out.append(s)
    uniq = []
    seen = set()
    for h in out:
        if h in seen:
            continue
        seen.add(h)
        uniq.append(h)
    return uniq


def _layout_to_plan_types(layout_item):
    """把 signature_layout 的关系字段映射为落位策略标签。"""
    if not isinstance(layout_item, dict):
        return []
    rel = str(layout_item.get("date_relation") or "").strip()
    pos = str(layout_item.get("date_position") or "").strip()
    out = []
    if rel == "same_cell":
        out.append("same_cell_inline")
    if rel == "different_cell" and pos == "right":
        out.append("adjacent_right_cell")
    if pos == "below":
        out.append("footer_sig_above_date")
        out.append("below_cell")
    return out


def _build_placement_plan_from_detect(det, role_ids):
    """
    从 detect snapshot 提取“角色落位提示关键词”。
    只做弱约束提示：签字引擎优先使用这些关键词，但仍保留兜底关键词扫描。
    """
    plan = {}
    if not isinstance(det, dict):
        return plan
    role_ev = det.get("role_evidence") if isinstance(det.get("role_evidence"), dict) else {}
    sl = det.get("signature_layout") if isinstance(det.get("signature_layout"), dict) else {}
    sl_roles = sl.get("role_layouts") if isinstance(sl.get("role_layouts"), dict) else {}
    for rid in role_ids or []:
        rid_s = str(rid or "").strip()
        if not rid_s:
            continue
        arr = role_ev.get(rid_s)
        kws = []
        hints = []
        layout_types = []
        if isinstance(arr, list):
            for ev in arr[:5]:
                if not isinstance(ev, dict):
                    continue
                kws.extend(_keywords_from_label_preview(ev.get("label_preview") or ""))
                sh = str(ev.get("source_hint") or "").strip()
                if sh:
                    hints.append(sh)
                sf = str(ev.get("slot_form") or ev.get("layout_form") or "").strip()
                if "两行签批表" in sf:
                    layout_types.append("two_row_signoff_table")
                elif ("同字段" in sf) and ("占位" in sf):
                    layout_types.append("same_cell_inline")
                elif ("右侧" in sf) and ("单元格" in sf):
                    layout_types.append("adjacent_right_cell")
        if not kws and isinstance(det.get("blocks"), list):
            for b in det.get("blocks")[:20]:
                if not isinstance(b, dict):
                    continue
                fields = b.get("fields") if isinstance(b.get("fields"), list) else []
                has_role = False
                for f in fields:
                    if not isinstance(f, dict):
                        continue
                    if str(f.get("type") or "") == "role_id" and str(f.get("name") or "") == rid_s:
                        has_role = True
                        break
                if not has_role:
                    continue
                kws.extend(_keywords_from_label_preview(b.get("label_preview") or ""))
                sh = str(b.get("source_hint") or "").strip()
                if sh:
                    hints.append(sh)
                sf = str(b.get("slot_form") or b.get("layout_form") or "").strip()
                if "两行签批表" in sf:
                    layout_types.append("two_row_signoff_table")
                elif ("同字段" in sf) and ("占位" in sf):
                    layout_types.append("same_cell_inline")
                elif ("右侧" in sf) and ("单元格" in sf):
                    layout_types.append("adjacent_right_cell")
                if len(kws) >= 8:
                    break
        if kws or hints or layout_types:
            uniq = []
            seen = set()
            for kw in kws:
                if kw in seen:
                    continue
                seen.add(kw)
                uniq.append(kw)
                if len(uniq) >= 8:
                    break
            huniq = []
            hseen = set()
            for h in hints:
                if h in hseen:
                    continue
                hseen.add(h)
                huniq.append(h)
                if len(huniq) >= 12:
                    break
            lt_uniq = []
            lt_seen = set()
            for lt in layout_types:
                if lt in lt_seen:
                    continue
                lt_seen.add(lt)
                lt_uniq.append(lt)
            plan[rid_s] = {"keywords": uniq, "source_hints": huniq, "layout_types": lt_uniq}
        # 追加 signature_layout 的角色级提示，避免“版式识别对了但落位仍找不到”。
        li = sl_roles.get(rid_s) if isinstance(sl_roles, dict) else None
        if isinstance(li, dict):
            p = plan.get(rid_s) if isinstance(plan.get(rid_s), dict) else {}
            p_kws = list(p.get("keywords") or [])
            p_hints = list(p.get("source_hints") or [])
            p_lts = list(p.get("layout_types") or [])
            for lk in _keywords_from_label_preview(li.get("label_preview") or ""):
                if lk not in p_kws:
                    p_kws.append(lk)
            for loc_key in ("name_loc", "date_loc"):
                for hh in _layout_loc_to_source_hints(li.get(loc_key) or ""):
                    if hh not in p_hints:
                        p_hints.append(hh)
            for lt in _layout_to_plan_types(li):
                if lt not in p_lts:
                    p_lts.append(lt)
            nc = _parse_docx_table_cell_loc(li.get("name_loc") or "")
            dc = _parse_docx_table_cell_loc(li.get("date_loc") or "")
            xnc = _parse_xlsx_table_cell_loc(li.get("name_loc") or "")
            xdc = _parse_xlsx_table_cell_loc(li.get("date_loc") or "")
            if p_kws or p_hints or p_lts or nc or dc or xnc or xdc:
                entry = {
                    "keywords": p_kws[:8],
                    "source_hints": p_hints[:12],
                    "layout_types": p_lts[:6],
                    "date_slot": bool(li.get("date_slot", True)),
                }
                if nc:
                    entry["name_cell"] = nc
                if dc:
                    entry["date_cell"] = dc
                if xnc:
                    entry["xlsx_name_cell"] = xnc
                if xdc:
                    entry["xlsx_date_cell"] = xdc
                plan[rid_s] = entry
    return plan


def _normalize_sign_role_ids(role_ids) -> list:
    from sign_handlers import ROLE_ID_TO_KEYWORD
    from sign_handlers.config import canonical_sign_role_id

    out = []
    seen = set()
    for rid in role_ids or []:
        rr = canonical_sign_role_id(str(rid or "").strip())
        if not rr or rr in seen or rr not in ROLE_ID_TO_KEYWORD:
            continue
        seen.add(rr)
        out.append(rr)
    return out


def _resolve_actual_roles_for_file(file_id: Optional[str], fallback_roles: list) -> list:
    roles = _normalize_sign_role_ids(fallback_roles or [])
    fid = str(file_id or "").strip()
    if not fid:
        return roles
    corr = _get_file_detect_correction(fid) or {}
    exp = corr.get("expected_roles")
    if isinstance(exp, list) and exp:
        exp_roles = _normalize_sign_role_ids(exp)
        if exp_roles:
            return exp_roles
    return roles


_CORE_SIGNOFF_ROLES = ("author", "reviewer", "approver")
_SIGN_SLOT_PLACEHOLDER_RE = re.compile(
    r"(?:_{3,}|\.{3,}|·{3,}|□|■|[（(]\s*[年Y]\s*[）)]\s*[（(]\s*[月M]\s*[）)]\s*[（(]\s*[日D]\s*[）)])"
)
_SIGN_SLOT_LABEL_RE = re.compile(r"(?:编制|编写|作者|审核|复核|批准|签字|签名|日期|Date|Signed)", re.I)
_SIGNOFF_ROW_HINT_RE = re.compile(
    r"(?:table(\d+)\.row(\d+)|paragraph(\d+)|header|footer|sheet(\d+)\.row(\d+))",
    re.IGNORECASE,
)


def _looks_like_sign_slot_placeholder(text: str) -> bool:
    t = str(text or "").strip()
    if not t:
        return False
    if _SIGN_SLOT_PLACEHOLDER_RE.search(t):
        return True
    # 角色/日期标签后跟长空白，通常是待签字位（Word/Excel 常见）
    if re.search(r"(编制|编写|作者|审核|复核|批准|签字|签名|日期)\s*[:：]?\s{3,}", t):
        return True
    has_label = bool(_SIGN_SLOT_LABEL_RE.search(t))
    if not has_label:
        return False
    # 自适应：不枚举具体符号，按“标签 + 待填写形态”判断
    if re.search(r"(?:编制|编写|作者|审核|复核|批准|签字|签名|日期|Date|Signed)\s*[:：]\s*$", t, re.I):
        return True
    for seg in re.findall(r"[^\w\u4e00-\u9fff]{3,}", t):
        if seg.strip():
            return True
    if re.search(r"[（(]\s*[^\w\u4e00-\u9fff]?\s*[）)]", t):
        return True
    return False


def _signoff_row_bias(source_hint: str) -> int:
    s = str(source_hint or "")
    m = _SIGNOFF_ROW_HINT_RE.search(s)
    if not m:
        return 0
    if m.group(0).lower() in ("header", "footer"):
        return 2
    table_no = int(m.group(1)) if m.group(1) else 0
    row_no = int(m.group(2)) if m.group(2) else 0
    para_no = int(m.group(3)) if m.group(3) else 0
    sheet_row_no = int(m.group(5)) if m.group(5) else 0
    if para_no and para_no <= 140:
        return 1
    if sheet_row_no and sheet_row_no <= 60:
        return 1
    if table_no and (table_no <= 6 or row_no <= 8):
        return 1
    return 0


def _collect_block_role_stats(det: dict) -> dict:
    stats = {}
    blocks = det.get("blocks") if isinstance(det.get("blocks"), list) else []
    for b in blocks:
        if not isinstance(b, dict):
            continue
        fields = b.get("fields") if isinstance(b.get("fields"), list) else []
        role_ids = []
        has_date = False
        has_action = False
        for f in fields:
            if not isinstance(f, dict):
                continue
            ftype = str(f.get("type") or "").strip()
            if ftype == "date":
                has_date = True
            elif ftype == "action":
                has_action = True
            elif ftype == "role_id":
                role_ids.append(str(f.get("name") or "").strip())
        if not role_ids:
            continue
        txt = (str(b.get("label_preview") or "") + " " + str(b.get("source_hint") or "")).strip()
        matched = [str(x or "").strip() for x in (b.get("matched_rules") or [])]
        has_placeholder = _looks_like_sign_slot_placeholder(txt)
        has_strong_rule = any(
            x in (
                "strong_block_rule",
                "triad_rule",
                "triad_text_hint",
                "role_date_rule",
                "table_header_rule",
                "docx_role_with_date_row",
                "multi_role_rule",
                "org_seal_rule",
            )
            for x in matched
        )
        row_bias = _signoff_row_bias(b.get("source_hint") or "")
        contextual = has_date or has_action or has_placeholder or has_strong_rule
        is_strong = bool(contextual and row_bias >= 1) or bool(has_strong_rule and contextual)
        for rid in role_ids:
            if not rid:
                continue
            one = stats.setdefault(
                rid,
                {
                    "total_blocks": 0,
                    "strong_blocks": 0,
                    "date_blocks": 0,
                    "action_blocks": 0,
                    "placeholder_blocks": 0,
                },
            )
            one["total_blocks"] += 1
            if is_strong:
                one["strong_blocks"] += 1
            if has_date:
                one["date_blocks"] += 1
            if has_action:
                one["action_blocks"] += 1
            if has_placeholder:
                one["placeholder_blocks"] += 1
    return stats


def _collect_layout_candidate_roles(det: dict) -> list:
    from sign_handlers import ROLE_ID_TO_KEYWORD
    from sign_handlers.config import canonical_sign_role_id

    out = []
    seen = set()
    for rr in det.get("roles") or []:
        if not isinstance(rr, dict):
            continue
        rid = canonical_sign_role_id(str(rr.get("id") or "").strip())
        if rid and rid not in seen and rid in ROLE_ID_TO_KEYWORD:
            seen.add(rid)
            out.append(rid)
    stats = _collect_block_role_stats(det)
    for rid_raw, st in stats.items():
        rid = canonical_sign_role_id(rid_raw)
        if not rid or rid in seen or rid not in ROLE_ID_TO_KEYWORD:
            continue
        if st.get("strong_blocks", 0) > 0:
            seen.add(rid)
            out.append(rid)
    # 限制角色集合规模，避免把正文噪音角色全带入布局分析。
    return out[:8]


def _layout_role_is_strong(layout_item: dict) -> bool:
    if not isinstance(layout_item, dict):
        return False
    if not bool(layout_item.get("name_slot")):
        return False
    if bool(layout_item.get("date_slot")):
        return True
    rel = str(layout_item.get("date_relation") or "").strip()
    sep = str(layout_item.get("separator") or "").strip()
    return rel in ("same_cell", "different_cell", "paragraph_inline") or sep in (
        "slash",
        "space",
        "empty_cell",
        "cell",
        "newline",
        "adjacent",
    )


def _looks_like_formal_signoff_case(
    role_conf: dict, stats: dict, layout_roles: dict, debug_summary: dict
) -> bool:
    core_hit = 0
    for rid in _CORE_SIGNOFF_ROLES:
        st = stats.get(rid, {})
        if (
            role_conf.get(rid, 0.0) >= 0.56
            or _layout_role_is_strong(layout_roles.get(rid) or {})
            or int(st.get("strong_blocks", 0)) > 0
        ):
            core_hit += 1
    edge_filtered = bool((debug_summary or {}).get("edge_filter_applied"))
    if core_hit >= 3:
        return True
    if core_hit >= 2 and edge_filtered:
        return True
    return core_hit >= 2 and sum(int((stats.get(r) or {}).get("date_blocks", 0)) for r in _CORE_SIGNOFF_ROLES) >= 2


def _reconcile_detect_roles_with_layout(det: dict, layout_info: dict, source_name: str = "") -> dict:
    from sign_handlers import ROLE_ID_TO_KEYWORD
    from sign_handlers.config import canonical_sign_role_id

    if not isinstance(det, dict) or not det.get("ok"):
        return det
    out = dict(det)
    role_conf = {}
    for rr in out.get("roles") or []:
        if not isinstance(rr, dict):
            continue
        rid = canonical_sign_role_id(str(rr.get("id") or "").strip())
        if rid and rid in ROLE_ID_TO_KEYWORD:
            role_conf[rid] = max(role_conf.get(rid, 0.0), float(rr.get("confidence") or 0.0))
    stats = _collect_block_role_stats(out)
    layout_roles_raw = (
        layout_info.get("role_layouts")
        if isinstance(layout_info, dict) and isinstance(layout_info.get("role_layouts"), dict)
        else {}
    )
    layout_roles = {}
    for rid_raw, item in layout_roles_raw.items():
        rid = canonical_sign_role_id(str(rid_raw or "").strip())
        if rid and rid in ROLE_ID_TO_KEYWORD and isinstance(item, dict):
            layout_roles[rid] = item
    ds_src = det.get("debug_summary") if isinstance(det.get("debug_summary"), dict) else {}
    signoff_doc = _looks_like_formal_signoff_case(role_conf, stats, layout_roles, ds_src)

    # 布局强证据可补齐/抬高角色置信度。
    for rid, info in layout_roles.items():
        if not _layout_role_is_strong(info):
            continue
        st = stats.get(rid, {})
        boost = 0.84 if rid in _CORE_SIGNOFF_ROLES else 0.74
        boost = max(boost, min(0.92, 0.68 + 0.05 * int(st.get("strong_blocks", 0))))
        role_conf[rid] = max(role_conf.get(rid, 0.0), boost)

    # 报告/方案等固定签批类：至少两核心角色成立时，对第三角色做弱补齐。
    if signoff_doc:
        core_present = [
            rid
            for rid in _CORE_SIGNOFF_ROLES
            if role_conf.get(rid, 0.0) >= 0.56 or _layout_role_is_strong(layout_roles.get(rid) or {})
        ]
        if len(core_present) >= 2:
            for rid in _CORE_SIGNOFF_ROLES:
                st = stats.get(rid, {})
                if rid not in role_conf and (
                    st.get("strong_blocks", 0) > 0
                    or st.get("date_blocks", 0) > 0
                    or rid in layout_roles
                ):
                    role_conf[rid] = 0.66
                elif role_conf.get(rid, 0.0) < 0.62 and (
                    st.get("strong_blocks", 0) > 0 or rid in layout_roles
                ):
                    role_conf[rid] = 0.62

    # 对噪音角色降噪：无布局证据且只有弱文本命中的角色（常见正文列头）剔除。
    for rid in list(role_conf.keys()):
        if rid in _CORE_SIGNOFF_ROLES:
            continue
        st = stats.get(rid, {})
        weak_only = st.get("total_blocks", 0) > 0 and st.get("strong_blocks", 0) <= 0
        has_layout = _layout_role_is_strong(layout_roles.get(rid) or {})
        if weak_only and not has_layout and role_conf.get(rid, 0.0) < 0.7:
            role_conf.pop(rid, None)

    out["roles"] = [{"id": rid, "confidence": role_conf[rid]} for rid in sorted(role_conf.keys())]
    ds = dict(out.get("debug_summary") or {})
    ds["role_layout_reconciled"] = True
    ds["layout_role_count"] = len(layout_roles)
    ds["role_count_after_reconcile"] = len(out["roles"])
    out["debug_summary"] = ds
    return out


_SLOT_PROBE_DUMMY_PNG_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAASsJTYQAAAAASUVORK5CYII="
)


def _probe_detect_slot_placement(
    file_bytes: bytes,
    ext: str,
    base_name: str,
    role_ids: list,
    placement_plan: Optional[dict] = None,
) -> dict:
    """
    用最小 PNG 对已识别角色做“可落位”探测：
    - 不写入 signed_output
    - 仅验证 sign_engine 能否在对应签字位完成落位
    """
    role_ids = [str(x or "").strip() for x in (role_ids or []) if str(x or "").strip()]
    if not role_ids:
        return {"ok": False, "error": "未识别到需签角色", "missing_roles": []}
    try:
        dummy_png = base64.b64decode(_SLOT_PROBE_DUMMY_PNG_B64)
    except Exception:
        return {"ok": False, "error": "构造签字位探测素材失败", "missing_roles": role_ids}

    from sign_handlers import sign_document

    sig_map = {rid: dummy_png for rid in role_ids}
    date_map = {rid: dummy_png for rid in role_ids}
    placement_result = {}
    tmp_dir = tempfile.mkdtemp(prefix="aiprintword_slot_probe_")
    try:
        safe_in_name = secure_filename(os.path.basename(base_name or "")) or ("document" + ext)
        if not safe_in_name.lower().endswith(ext):
            safe_in_name = os.path.splitext(safe_in_name)[0] + ext
        in_path = os.path.join(tmp_dir, f"{uuid.uuid4().hex[:8]}_{safe_in_name}")
        with open(in_path, "wb") as fp:
            fp.write(file_bytes)
        sign_document(
            in_path,
            sig_map,
            date_map,
            placement_plan=placement_plan,
            placement_result=placement_result,
        )
        missing = []
        for rid in role_ids:
            one = placement_result.get(rid) if isinstance(placement_result, dict) else None
            if not (isinstance(one, dict) and one.get("placed")):
                missing.append(rid)
        return {
            "ok": not missing,
            "error": "" if not missing else ("以下角色未找到可落位签字位：" + "、".join(missing)),
            "missing_roles": missing,
            "active_roles": role_ids,
            "per_role_results": placement_result,
        }
    except Exception as e:
        return {
            "ok": False,
            "error": f"签字位可落位探测失败: {e}",
            "missing_roles": role_ids,
            "active_roles": role_ids,
            "per_role_results": placement_result,
        }
    finally:
        try:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass


def _sign_process_document_bytes(
    file_bytes: bytes,
    ext: str,
    base_name: str,
    roles: list,
    sig_map: dict,
    date_map: dict,
    source_file_id: Optional[str],
    batch_id: Optional[str] = None,
    placement_plan: Optional[dict] = None,
    require_role_placement: bool = False,
) -> dict:
    """生成已签名文档字节；MySQL 模式下同时写入 sign_signed_output。返回 dict 含 ok / error / out_bytes 等。

    require_role_placement=True 时：
    - 若所有角色都未落位 → ok=False（无可下载内容）；
    - 若仅部分角色未落位 → 仍 ok=True，文件名追加 `_部分成功` 标记，写入 sign_signed_output 供下载与打包，
      result 含 partial=True 与 missing_roles，便于前端展示与提示。
    """
    def _append_docx_manual_tail_rows(
        raw_docx: bytes, role_ids: list, role_sig_map: dict, role_date_map: dict
    ) -> tuple[bytes, list]:
        if not raw_docx or not role_ids:
            return raw_docx, []
        try:
            from docx import Document
            from docx.shared import Cm
            from sign_handlers.config import role_display_name
        except Exception:
            return raw_docx, []
        appended = []
        try:
            bio = io.BytesIO(raw_docx)
            doc = Document(bio)
            p = doc.add_paragraph()
            p.add_run("自动补签（未识别到签字位）：")
            table = doc.add_table(rows=1, cols=3)
            try:
                table.style = "Table Grid"
            except Exception:
                pass
            table.cell(0, 0).text = "角色"
            table.cell(0, 1).text = "签名"
            table.cell(0, 2).text = "日期"
            for rid in role_ids:
                rid_s = str(rid or "").strip()
                if not rid_s:
                    continue
                sb = role_sig_map.get(rid_s)
                db = role_date_map.get(rid_s)
                if not sb and not db:
                    continue
                row = table.add_row().cells
                row[0].text = role_display_name(rid_s) or rid_s
                if sb:
                    try:
                        row[1].paragraphs[0].add_run().add_picture(io.BytesIO(sb), width=Cm(2.8))
                    except Exception:
                        row[1].text = "签名插入失败"
                if db:
                    try:
                        row[2].paragraphs[0].add_run().add_picture(io.BytesIO(db), width=Cm(2.8))
                    except Exception:
                        row[2].text = "日期插入失败"
                appended.append(rid_s)
            if not appended:
                return raw_docx, []
            out_bio = io.BytesIO()
            doc.save(out_bio)
            return out_bio.getvalue(), appended
        except Exception:
            return raw_docx, []

    tmp_dir = tempfile.mkdtemp(prefix="aiprintword_sign_")
    try:
        # 临时文件名需兼容 Windows/文件系统；显示名称/记录名称保留原始（可含中文）
        safe_in_name = secure_filename(os.path.basename(base_name or "")) or ("document" + ext)
        if not safe_in_name.lower().endswith(ext):
            safe_in_name = os.path.splitext(safe_in_name)[0] + ext
        in_path = os.path.join(tmp_dir, f"{uuid.uuid4().hex[:8]}_{safe_in_name}")
        with open(in_path, "wb") as fp:
            fp.write(file_bytes)
        from sign_handlers import sign_document

        placement_result = {}
        out_path = sign_document(
            in_path,
            sig_map,
            date_map,
            placement_plan=placement_plan,
            placement_result=placement_result,
        )
        dl_name = os.path.splitext(os.path.basename(base_name or "document"))[0] + "_signed" + ext
        with open(out_path, "rb") as fp:
            out_bytes = fp.read()
        missing_roles = []
        active_roles = []
        fallback_roles = []
        placed_roles = []
        for rid in roles or []:
            rid_s = str(rid or "").strip()
            if not rid_s:
                continue
            if (rid_s not in sig_map) and (rid_s not in date_map):
                continue
            active_roles.append(rid_s)
            one = placement_result.get(rid_s) if isinstance(placement_result, dict) else None
            if not (isinstance(one, dict) and one.get("placed")):
                missing_roles.append(rid_s)
                continue
            placed_roles.append(rid_s)
            placed_by_val = str(one.get("placed_by") or "")
            if placed_by_val.startswith("fallback_keywords"):
                fallback_roles.append(rid_s)
        tail_appended_roles = []
        if missing_roles and ext == ".docx":
            out_bytes, tail_appended_roles = _append_docx_manual_tail_rows(
                out_bytes, missing_roles, sig_map, date_map
            )
            if tail_appended_roles and isinstance(placement_result, dict):
                for rid in tail_appended_roles:
                    one = placement_result.get(rid)
                    if isinstance(one, dict) and not one.get("placed"):
                        one["placed_by"] = "manual_tail_append"
                        one["tail_appended"] = True
        missing_roles = [r for r in missing_roles if r not in set(tail_appended_roles)]
        partial = bool(missing_roles) and bool(placed_roles or tail_appended_roles)
        # 仅当“强制落位”且“无任何角色落位且无末页补签”时才记为失败。
        if (
            require_role_placement
            and missing_roles
            and not placed_roles
            and not tail_appended_roles
        ):
            return {
                "ok": False,
                "error": "以下角色未找到可落位签字位：" + "、".join(missing_roles),
                "per_role_results": placement_result,
                "missing_roles": missing_roles,
                "active_roles": active_roles,
                "placed_roles": placed_roles,
                "fallback_roles": fallback_roles,
                "tail_appended_roles": [],
                "partial": False,
            }
        if partial:
            stem, _ext_dl = os.path.splitext(dl_name)
            dl_name = f"{stem}_部分成功{_ext_dl}"
        try:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass
        tmp_dir = None
        signed_row_id = None
        if _sign_using_mysql():
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            signed_row_id = uuid.uuid4().hex
            src_id = (
                source_file_id
                if source_file_id and _SIGN_FILE_ID_RE.match(source_file_id or "")
                else None
            )
            mysql_store.insert_signed_output(
                signed_row_id,
                batch_id,
                src_id,
                base_name,
                dl_name,
                ext,
                json.dumps(roles, ensure_ascii=False),
                out_bytes,
            )
        return {
            "ok": True,
            "partial": partial,
            "out_bytes": out_bytes,
            "dl_name": dl_name,
            "signed_id": signed_row_id,
            "per_role_results": placement_result,
            "missing_roles": missing_roles,
            "active_roles": active_roles,
            "placed_roles": placed_roles,
            "fallback_roles": fallback_roles,
            "tail_appended_roles": tail_appended_roles,
        }
    except Exception as e:
        if tmp_dir and os.path.isdir(tmp_dir):
            try:
                shutil.rmtree(tmp_dir, ignore_errors=True)
            except Exception:
                pass
        return {"ok": False, "error": str(e)}


def _aiword_handoff_secret_expected() -> str:
    v = (os.environ.get("AIWORD_HANDOFF_SECRET") or "").strip()
    if v:
        return v
    try:
        from runtime_settings.resolve import get_setting

        return (str(get_setting("AIWORD_HANDOFF_SECRET") or "")).strip()
    except Exception:
        return ""


def _handoff_safe_token(tok: str) -> bool:
    return bool(tok) and len(tok) <= 64 and re.match(r"^[0-9a-fA-F]+$", tok)


def _handoff_validate_reuse_ftp_path(p: str) -> str:
    """返回规范化后的可复用 FTP 路径；无效则返回空串。"""
    s = (p or "").strip()
    if not s or len(s) > 768 or ".." in s:
        return ""
    if "\x00" in s or "\r" in s or "\n" in s:
        return ""
    if s.startswith("ftp://") or s.startswith("http://") or s.startswith("https://"):
        return ""
    return s


def _handoff_prune_stale() -> None:
    """删除过期交接文件（best-effort）。"""
    if not os.path.isdir(HANDOFF_DIR):
        return
    now = time.time()
    try:
        for name in os.listdir(HANDOFF_DIR):
            if not name.endswith(".json"):
                continue
            jp = os.path.join(HANDOFF_DIR, name)
            try:
                with open(jp, "r", encoding="utf-8") as f:
                    meta = json.load(f)
                exp = float(meta.get("expires_at") or 0)
                if exp and now > exp + 3600:
                    if (meta.get("kind") or "").strip().lower() == "batch":
                        try:
                            os.remove(jp)
                        except OSError:
                            pass
                        continue
                    tok = meta.get("token") or name[:-5]
                    for suf in (".json", ".dat"):
                        p = os.path.join(HANDOFF_DIR, str(tok) + suf)
                        if os.path.isfile(p):
                            try:
                                os.remove(p)
                            except OSError:
                                pass
            except Exception:
                continue
    except Exception:
        pass


def _handoff_parse_aiword_context(raw_ctx: object) -> dict:
    parsed_ctx = None
    if isinstance(raw_ctx, dict):
        parsed_ctx = raw_ctx
    elif isinstance(raw_ctx, str):
        s = raw_ctx.strip()
        if s:
            try:
                parsed_ctx = json.loads(s)
            except Exception:
                parsed_ctx = None
    out: dict = {}
    if isinstance(parsed_ctx, dict):
        alias = {
            "projectId": "project_id",
            "projectName": "project_name",
            "projectCode": "project_code",
            "task_type": "phase",
            "belonging_module": "phase",
        }
        norm_ctx = dict(parsed_ctx)
        for ak, nk in alias.items():
            if nk in norm_ctx:
                continue
            if ak in parsed_ctx:
                norm_ctx[nk] = parsed_ctx.get(ak)
        for k in (
            "editor",
            "writer",
            "reviewer",
            "approver",
            "doc_date",
            "country",
            "phase",
            "project_id",
            "project_name",
            "project_code",
        ):
            v = norm_ctx.get(k)
            if v is None:
                continue
            sv = str(v).strip()
            if sv and len(sv) <= 500:
                out[k] = sv
    return out


def _normalize_handoff_display_filename(name: str) -> str:
    from sign_handlers.filename_util import normalize_display_filename

    return normalize_display_filename(name or "")


def _handoff_resolve_filename(meta: dict) -> str:
    """交接展示名；内部缓存名时尝试从 reuse_ftp_path 末段恢复中文名。"""
    from sign_handlers.filename_util import is_internal_cache_filename

    fname = _normalize_handoff_display_filename(
        (meta.get("filename") or "document.docx").strip() or "document.docx"
    )
    if not is_internal_cache_filename(fname):
        return fname
    reuse = (meta.get("reuse_ftp_path") or "").strip()
    if reuse:
        base = os.path.basename(reuse.replace("\\", "/"))
        if base and not is_internal_cache_filename(base):
            return _normalize_handoff_display_filename(base)
    return fname


def _handoff_create_one_token(
    *,
    purpose: str,
    filename: str,
    aiword_ctx: dict,
    reuse_ftp_path: str = "",
    raw: Optional[bytes] = None,
) -> dict:
    p = (purpose or "sign").strip().lower()
    if p not in ("sign", "print"):
        p = "sign"
    name = _normalize_handoff_display_filename(filename or "document.docx")
    reuse_ok = _handoff_validate_reuse_ftp_path(reuse_ftp_path or "")
    if not reuse_ok:
        if raw is None:
            raise ValueError("缺少文件内容")
        if not raw:
            raise ValueError("文件为空")
        if len(raw) > _HANDOFF_MAX_BYTES:
            raise ValueError("文件过大")

    token = uuid.uuid4().hex
    data_path = os.path.join(HANDOFF_DIR, token + ".dat")
    meta_path = os.path.join(HANDOFF_DIR, token + ".json")
    expires_at = time.time() + _HANDOFF_TTL_SEC
    meta_obj: dict = {
        "token": token,
        "filename": name,
        "purpose": p,
        "expires_at": expires_at,
    }
    if reuse_ok:
        meta_obj["reuse_ftp_path"] = reuse_ok
    if aiword_ctx:
        meta_obj["aiword_context"] = aiword_ctx
    with open(meta_path, "w", encoding="utf-8") as mf:
        json.dump(meta_obj, mf, ensure_ascii=False)
    if (not reuse_ok) and raw is not None:
        with open(data_path, "wb") as out:
            out.write(raw)
    return {
        "token": token,
        "filename": name,
        "purpose": p,
        "expires_at": expires_at,
        "reuse_ftp": bool(reuse_ok),
    }


@app.route("/api/handoff", methods=["POST"])
def api_handoff_create():
    """
    aiword 服务端将任务文档 POST 到此接口，换取一次性 token；浏览器再带 token 打开 /sign 或 / 由前端拉取文件。
    鉴权：请求头 X-Aiword-Handoff-Secret 须与环境变量 AIWORD_HANDOFF_SECRET（或 runtime_settings）一致。
    若提供 reuse_ftp_path（与 aiword 任务模板同一套 FTP），可不传文件体，由签字端直接登记该路径，避免重复上传。
    """
    expected = _aiword_handoff_secret_expected()
    got = (request.headers.get("X-Aiword-Handoff-Secret") or "").strip()
    if not expected or got != expected:
        return jsonify({"ok": False, "error": "未授权或未配置 AIWORD_HANDOFF_SECRET"}), 401

    purpose = (request.form.get("purpose") or "sign").strip().lower()
    if purpose not in ("sign", "print"):
        purpose = "sign"
    filename = _normalize_handoff_display_filename(
        (request.form.get("filename") or "document.docx").strip() or "document.docx"
    )
    reuse_raw = (request.form.get("reuse_ftp_path") or "").strip()
    reuse_ok = _handoff_validate_reuse_ftp_path(reuse_raw)

    aiword_ctx = _handoff_parse_aiword_context((request.form.get("handoff_context") or "").strip())
    os.makedirs(HANDOFF_DIR, exist_ok=True)
    _handoff_prune_stale()
    raw = None
    if not reuse_ok:
        up = request.files.get("file") or request.files.get("files")
        if not up or not getattr(up, "filename", None):
            return jsonify({"ok": False, "error": "缺少文件字段 file"}), 400
        raw = up.read()
    try:
        created = _handoff_create_one_token(
            purpose=purpose,
            filename=filename,
            aiword_ctx=aiword_ctx,
            reuse_ftp_path=reuse_ok,
            raw=raw,
        )
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    except OSError as e:
        return jsonify({"ok": False, "error": f"无法写入交接目录：{e}"}), 500

    return jsonify(
        {
            "ok": True,
            "token": created["token"],
            "expires_in_sec": _HANDOFF_TTL_SEC,
            "reuse_ftp": bool(created.get("reuse_ftp")),
        }
    )


@app.route("/api/handoff/batch", methods=["POST"])
def api_handoff_batch_create():
    expected = _aiword_handoff_secret_expected()
    got = (request.headers.get("X-Aiword-Handoff-Secret") or "").strip()
    if not expected or got != expected:
        return jsonify({"ok": False, "error": "未授权或未配置 AIWORD_HANDOFF_SECRET"}), 401

    manifest_raw = (request.form.get("manifest") or "").strip()
    if not manifest_raw:
        return jsonify({"ok": False, "error": "缺少 manifest"}), 400
    try:
        manifest = json.loads(manifest_raw)
    except Exception:
        return jsonify({"ok": False, "error": "manifest 不是有效 JSON"}), 400
    if not isinstance(manifest, list) or not manifest:
        return jsonify({"ok": False, "error": "manifest 需为非空数组"}), 400
    if len(manifest) > 200:
        return jsonify({"ok": False, "error": "manifest 条目过多（最多 200）"}), 400

    os.makedirs(HANDOFF_DIR, exist_ok=True)
    _handoff_prune_stale()

    items: list[dict] = []
    failures: list[dict] = []
    with _HANDOFF_LOCK:
        for idx, it in enumerate(manifest):
            if not isinstance(it, dict):
                failures.append({"index": idx, "error": "条目必须为对象"})
                continue
            purpose = (it.get("purpose") or request.form.get("purpose") or "sign").strip().lower()
            filename = _normalize_handoff_display_filename(
                (it.get("filename") or f"document_{idx+1}.docx").strip() or f"document_{idx+1}.docx"
            )
            aiword_ctx = _handoff_parse_aiword_context(it.get("handoff_context"))
            reuse_ok = _handoff_validate_reuse_ftp_path(str(it.get("reuse_ftp_path") or ""))
            raw = None
            if not reuse_ok:
                ff = str(it.get("file_field") or f"file_{idx}")
                up = request.files.get(ff)
                if not up or not getattr(up, "filename", None):
                    failures.append({"index": idx, "filename": filename, "error": f"缺少文件字段 {ff}"})
                    continue
                raw = up.read()
            try:
                created = _handoff_create_one_token(
                    purpose=purpose,
                    filename=filename,
                    aiword_ctx=aiword_ctx,
                    reuse_ftp_path=reuse_ok,
                    raw=raw,
                )
            except ValueError as e:
                failures.append({"index": idx, "filename": filename, "error": str(e)})
                continue
            except OSError as e:
                failures.append({"index": idx, "filename": filename, "error": f"无法写入交接目录：{e}"})
                continue
            items.append(
                {
                    "index": idx,
                    "token": created["token"],
                    "filename": created["filename"],
                    "purpose": created["purpose"],
                    "reuse_ftp": bool(created.get("reuse_ftp")),
                }
            )

        if not items:
            return jsonify({"ok": False, "error": "批量交接失败", "failures": failures}), 400

        batch_token = uuid.uuid4().hex
        batch_meta_path = os.path.join(HANDOFF_DIR, f"batch_{batch_token}.json")
        exp = time.time() + _HANDOFF_TTL_SEC
        try:
            with open(batch_meta_path, "w", encoding="utf-8") as bf:
                json.dump(
                    {
                        "kind": "batch",
                        "token": batch_token,
                        "expires_at": exp,
                        "items": items,
                    },
                    bf,
                    ensure_ascii=False,
                )
        except OSError as e:
            return jsonify({"ok": False, "error": f"无法写入批量交接目录：{e}"}), 500

    return jsonify(
        {
            "ok": True,
            "batch_token": batch_token,
            "items": items,
            "failures": failures,
            "success_count": len(items),
            "failure_count": len(failures),
            "expires_in_sec": _HANDOFF_TTL_SEC,
        }
    )


@app.route("/api/handoff/batch/<batch_token>", methods=["GET"])
def api_handoff_batch_get(batch_token: str):
    tok = (batch_token or "").strip()
    if not _handoff_safe_token(tok):
        return jsonify({"ok": False, "error": "无效的 batch_token"}), 400
    p = os.path.join(HANDOFF_DIR, f"batch_{tok}.json")
    with _HANDOFF_LOCK:
        if not os.path.isfile(p):
            return jsonify({"ok": False, "error": "批量交接不存在或已过期"}), 404
        try:
            with open(p, "r", encoding="utf-8") as f:
                meta = json.load(f)
        except Exception:
            return jsonify({"ok": False, "error": "批量交接元数据损坏"}), 500
        if float(meta.get("expires_at") or 0) < time.time():
            try:
                os.remove(p)
            except OSError:
                pass
            return jsonify({"ok": False, "error": "批量交接已过期"}), 410
        items = meta.get("items")
        if not isinstance(items, list):
            return jsonify({"ok": False, "error": "批量交接元数据损坏"}), 500
    return jsonify({"ok": True, "batch_token": tok, "items": items})


@app.route("/api/handoff/<token>/claim-sign", methods=["POST"])
def api_handoff_claim_sign(token: str):
    """浏览器一次性将交接登记进签字文件列表。"""
    out, status = _handoff_claim_sign_token(token)
    return jsonify(out), status


def _handoff_claim_sign_token(token: str, *, include_files_list: bool = True) -> tuple[dict, int]:
    tok = (token or "").strip()
    if not _handoff_safe_token(tok):
        return {"ok": False, "error": "无效的 token"}, 400
    meta_path = os.path.join(HANDOFF_DIR, tok + ".json")
    data_path = os.path.join(HANDOFF_DIR, tok + ".dat")
    snap: dict = {}
    with _HANDOFF_LOCK:
        if not os.path.isfile(meta_path):
            return {"ok": False, "error": "交接不存在或已使用"}, 404
        try:
            with open(meta_path, "r", encoding="utf-8") as f:
                meta = json.load(f)
        except Exception:
            return {"ok": False, "error": "交接元数据损坏"}, 500
        if float(meta.get("expires_at") or 0) < time.time():
            try:
                if os.path.isfile(data_path):
                    os.remove(data_path)
                os.remove(meta_path)
            except OSError:
                pass
            return {"ok": False, "error": "交接已过期"}, 410
        if (meta.get("purpose") or "sign") != "sign":
            return {"ok": False, "error": "非签字交接"}, 400
        reuse_p = (meta.get("reuse_ftp_path") or "").strip()
        fname = _handoff_resolve_filename(meta)
        ac = meta.get("aiword_context")
        ctx_out: dict = {}
        if isinstance(ac, dict):
            ctx_out = dict(ac)
        snap = {"reuse_p": reuse_p, "fname": fname, "ctx": ctx_out}
        if not reuse_p:
            if not os.path.isfile(data_path):
                return {"ok": False, "error": "交接不存在或已使用"}, 404
            try:
                snap["raw"] = Path(data_path).read_bytes()
            except OSError as e:
                return {"ok": False, "error": str(e)}, 500

    fname = snap["fname"]
    reuse_p = snap["reuse_p"]
    ctx_out = snap.get("ctx") or {}
    raw = snap.get("raw")

    ext = os.path.splitext(fname)[1].lower() or ".docx"
    if ext not in SIGN_ALLOWED_EXT:
        return {"ok": False, "error": f"不支持的扩展名：{ext}"}, 400

    file_id = uuid.uuid4().hex
    try:
        if _sign_using_mysql():
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            cur_n = mysql_store.count_files()
            _max_f = _sign_mysql_max_files()
            if cur_n + 1 > _max_f:
                return {"ok": False, "error": f"文件数已达上限 {_max_f}"}, 400
            if reuse_p:
                mysql_store.insert_file_from_external_ftp(file_id, fname, ext, reuse_p)
            else:
                mysql_store.insert_file(file_id, fname, ext, raw or b"")
            last_rec = {"id": file_id, "name": fname, "ext": ext}
            _sign_apply_ctx_project_to_file(file_id, ctx_out, last_rec)
            out = {
                "ok": True,
                "file": last_rec,
                "context": ctx_out,
            }
            if include_files_list:
                out["files"] = mysql_store.list_files()
        else:
            sid, inbox_dir = _sign_ensure_session_inbox()
            # 无 MySQL：同任务多次「去签字」会重复追加同名文件；去掉旧同名条仅保留本次
            fname_norm = (fname or "").strip()
            prev = list(session.get("sign_files") or [])
            kept = []
            for old in prev:
                oname = ((old or {}).get("name") or "").strip()
                if fname_norm and oname == fname_norm:
                    oid = (old or {}).get("id")
                    if oid:
                        _sign_remove_disk_files_for_id(sid, str(oid), (old or {}).get("ext"))
                    continue
                kept.append(old)
            if len(kept) >= SIGN_MAX_SAVED_FILES:
                return {"ok": False, "error": f"文件数已达上限 {SIGN_MAX_SAVED_FILES}"}, 400
            last_rec = {"id": file_id, "name": fname, "ext": ext}
            _sign_apply_ctx_project_to_file(file_id, ctx_out, last_rec)
            kept.append(last_rec)
            records = kept
            file_blob = raw
            if reuse_p:
                from ftp_store import download_bytes

                file_blob = download_bytes(reuse_p)
                if not file_blob:
                    return {"ok": False, "error": "FTP 下载结果为空"}, 502
            dest = os.path.join(inbox_dir, file_id + ext)
            with open(dest, "wb") as fp:
                fp.write(file_blob or b"")
            session["sign_files"] = records
            session.modified = True
            out = {"ok": True, "file": last_rec, "files": records, "context": ctx_out}
    except Exception as e:
        return {"ok": False, "error": str(e)}, 500

    with _HANDOFF_LOCK:
        try:
            if os.path.isfile(data_path):
                os.remove(data_path)
            if os.path.isfile(meta_path):
                os.remove(meta_path)
        except OSError:
            pass
    return out, 200


@app.route("/api/handoff/batch/<batch_token>/claim-sign", methods=["POST"])
def api_handoff_batch_claim_sign(batch_token: str):
    tok = (batch_token or "").strip()
    if not _handoff_safe_token(tok):
        return jsonify({"ok": False, "error": "无效的 batch_token"}), 400
    p = os.path.join(HANDOFF_DIR, f"batch_{tok}.json")
    items: list[dict] = []
    with _HANDOFF_LOCK:
        if not os.path.isfile(p):
            return jsonify({"ok": False, "error": "批量交接不存在或已过期"}), 404
        try:
            with open(p, "r", encoding="utf-8") as f:
                meta = json.load(f)
        except Exception:
            return jsonify({"ok": False, "error": "批量交接元数据损坏"}), 500
        if float(meta.get("expires_at") or 0) < time.time():
            try:
                os.remove(p)
            except OSError:
                pass
            return jsonify({"ok": False, "error": "批量交接已过期"}), 410
        items = meta.get("items") if isinstance(meta, dict) else []
        if not isinstance(items, list):
            return jsonify({"ok": False, "error": "批量交接元数据损坏"}), 500

    successes: list[dict] = []
    failures: list[dict] = []
    merged_files: list[dict] = []
    for idx, it in enumerate(items):
        tk = str((it or {}).get("token") or "").strip() if isinstance(it, dict) else ""
        if not tk:
            failures.append({"index": idx, "error": "缺少 token"})
            continue
        out, status = _handoff_claim_sign_token(tk, include_files_list=False)
        if status != 200 or not out.get("ok"):
            failures.append(
                {
                    "index": idx,
                    "token": tk,
                    "error": str(out.get("error") or "claim failed"),
                    "status": status,
                }
            )
            continue
        one = {
            "index": idx,
            "token": tk,
            "file": out.get("file"),
            "context": out.get("context") if isinstance(out.get("context"), dict) else {},
        }
        successes.append(one)

    if _sign_using_mysql():
        try:
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            merged_files = mysql_store.list_files()
        except Exception:
            merged_files = merged_files or []

    with _HANDOFF_LOCK:
        try:
            if os.path.isfile(p):
                os.remove(p)
        except OSError:
            pass

    if not successes:
        return jsonify({"ok": False, "error": "批量认领失败", "failures": failures}), 400
    return jsonify(
        {
            "ok": True,
            "items": successes,
            "files": merged_files,
            "success_count": len(successes),
            "failure_count": len(failures),
            "failures": failures,
        }
    )


@app.route("/api/handoff/<token>/file", methods=["GET"])
def api_handoff_download(token: str):
    """一次性下载：返回文件流后立即删除交接文件。支持仅 meta + reuse_ftp_path（无 .dat）时从 FTP 拉流。"""
    tok = (token or "").strip()
    if not _handoff_safe_token(tok):
        return jsonify({"ok": False, "error": "无效的 token"}), 400
    meta_path = os.path.join(HANDOFF_DIR, tok + ".json")
    data_path = os.path.join(HANDOFF_DIR, tok + ".dat")
    meta: dict = {}
    blob: bytes = b""
    fname = "document.docx"
    with _HANDOFF_LOCK:
        if not os.path.isfile(meta_path):
            return jsonify({"ok": False, "error": "交接不存在或已使用"}), 404
        try:
            with open(meta_path, "r", encoding="utf-8") as f:
                meta = json.load(f)
        except Exception:
            return jsonify({"ok": False, "error": "交接元数据损坏"}), 500
        if float(meta.get("expires_at") or 0) < time.time():
            try:
                if os.path.isfile(data_path):
                    os.remove(data_path)
                os.remove(meta_path)
            except OSError:
                pass
            return jsonify({"ok": False, "error": "交接已过期"}), 410
        fname = _handoff_resolve_filename(meta)
        reuse_p = (meta.get("reuse_ftp_path") or "").strip()
        if reuse_p:
            try:
                from ftp_store import download_bytes

                blob = download_bytes(reuse_p)
            except Exception as e:
                try:
                    os.remove(meta_path)
                except OSError:
                    pass
                return jsonify({"ok": False, "error": f"FTP 读取失败：{e}"}), 502
        else:
            if not os.path.isfile(data_path):
                return jsonify({"ok": False, "error": "交接不存在或已使用"}), 404
            try:
                blob = Path(data_path).read_bytes()
            except OSError as e:
                return jsonify({"ok": False, "error": str(e)}), 500
        try:
            if os.path.isfile(data_path):
                os.remove(data_path)
            os.remove(meta_path)
        except OSError:
            pass

    from urllib.parse import quote

    safe = os.path.basename(fname.replace("\\", "/")) or "document.docx"
    disp = f"attachment; filename*=UTF-8''{quote(safe)}"
    hdrs = {"Content-Disposition": disp}
    ctx_out = meta.get("aiword_context") if isinstance(meta, dict) else None
    if isinstance(ctx_out, dict) and ctx_out:
        try:
            b64 = base64.b64encode(json.dumps(ctx_out, ensure_ascii=False).encode("utf-8")).decode("ascii")
            if len(b64) <= 6144:
                hdrs["X-Aiword-Handoff-Context"] = b64
        except Exception:
            pass
    return Response(
        blob,
        mimetype="application/octet-stream",
        headers=hdrs,
    )


@app.route("/sign")
def sign_page():
    """Serve online signature page (default: file signing)."""
    return send_from_directory(os.path.join(ROOT, "static"), "sign_file.html")


@app.route("/sign/materials")
def sign_materials_page():
    """Serve online signature materials page (stroke library input)."""
    return send_from_directory(os.path.join(ROOT, "static"), "sign_materials.html")


@app.route("/api/sign/files/file-caches", methods=["GET"])
def api_sign_file_caches_list():
    """各文件已保存的识别结果、工作台行状态、角色映射（页面打开时恢复）。"""
    lite = (request.args.get("lite") or "").strip().lower() in ("1", "true", "yes")
    try:
        if _sign_using_mysql():
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            caches = mysql_store.list_file_session_caches(lite=lite)
            return jsonify({"ok": True, "lite": lite, "caches": caches})
        _sign_ensure_session_inbox()
        sid = session["sign_inbox_sid"]
        from sign_handlers.sign_library_local import list_file_session_caches as local_list

        return jsonify({"ok": True, "lite": lite, "caches": local_list(SIGN_INBOX_ROOT, sid, lite=lite)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/files/<file_id>/file-cache", methods=["PUT"])
def api_sign_file_cache_put(file_id):
    if not _SIGN_FILE_ID_RE.match(file_id or ""):
        return jsonify({"ok": False, "error": "无效的文件 id"}), 400
    data = request.get_json(silent=True) or {}
    detect = data.get("detect")
    workbench = data.get("workbench")
    detect_correction = data.get("detect_correction")
    try:
        if _sign_using_mysql():
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            if isinstance(detect, dict):
                mysql_store.set_file_detect_snapshot(file_id, detect)
            if isinstance(workbench, dict):
                mysql_store.set_file_workbench_state(file_id, workbench)
            if isinstance(detect_correction, dict):
                mysql_store.set_file_detect_correction(file_id, detect_correction)
        else:
            _sign_ensure_session_inbox()
            sid = session["sign_inbox_sid"]
            from sign_handlers.sign_library_local import (
                set_file_detect_snapshot as local_set_detect,
                set_file_workbench_state as local_set_wb,
                set_file_detect_correction as local_set_corr,
            )

            if isinstance(detect, dict):
                local_set_detect(SIGN_INBOX_ROOT, sid, file_id, detect)
            if isinstance(workbench, dict):
                local_set_wb(SIGN_INBOX_ROOT, sid, file_id, workbench)
            if isinstance(detect_correction, dict):
                local_set_corr(SIGN_INBOX_ROOT, sid, file_id, detect_correction)
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/files/batch-file-cache", methods=["PUT"])
def api_sign_files_batch_file_cache():
    """批量保存工作台元数据（文档日期/版本/编审批等），避免逐文件 PUT。"""
    data = request.get_json(silent=True) or {}
    raw_items = data.get("items")
    if not isinstance(raw_items, dict) or not raw_items:
        return jsonify({"ok": False, "error": "缺少 items（file_id -> 缓存片段）"}), 400
    items: dict = {}
    for k, v in raw_items.items():
        fid = str(k or "").strip()
        if not fid or not _SIGN_FILE_ID_RE.match(fid):
            continue
        if isinstance(v, dict) and v:
            items[fid] = v
    if not items:
        return jsonify({"ok": False, "error": "items 无有效 file_id"}), 400
    meta_only = data.get("meta_only") in (True, "true", "1", 1, "yes", "on")
    try:
        if _sign_using_mysql():
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            n = mysql_store.apply_files_session_cache_batch(
                items, workbench_only=meta_only
            )
            return jsonify(
                {"ok": True, "updated": n, "count": len(items), "meta_only": meta_only}
            )
        _sign_ensure_session_inbox()
        sid = session["sign_inbox_sid"]
        from sign_handlers.sign_library_local import (
            set_file_detect_correction as local_set_corr,
            set_file_detect_snapshot as local_set_detect,
            set_file_workbench_state as local_set_wb,
        )

        n = 0
        for fid, payload in items.items():
            if isinstance(payload.get("workbench"), dict):
                local_set_wb(SIGN_INBOX_ROOT, sid, fid, payload["workbench"])
                n += 1
            if isinstance(payload.get("detect"), dict):
                local_set_detect(SIGN_INBOX_ROOT, sid, fid, payload["detect"])
            if isinstance(payload.get("detect_correction"), dict):
                local_set_corr(SIGN_INBOX_ROOT, sid, fid, payload["detect_correction"])
        session.modified = True
        return jsonify({"ok": True, "updated": n, "count": len(items)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/files/<file_id>/detect-correction", methods=["GET", "PUT"])
def api_sign_file_detect_correction(file_id):
    """读取/保存人工识别纠正登记（错在哪、正确角色、标签提示、参考图元数据）。"""
    if not _SIGN_FILE_ID_RE.match(file_id or ""):
        return jsonify({"ok": False, "error": "无效的文件 id"}), 400
    if request.method == "GET":
        return jsonify({"ok": True, "correction": _get_file_detect_correction(file_id)})
    data = request.get_json(silent=True) or {}
    correction = data.get("correction")
    if not isinstance(correction, dict):
        return jsonify({"ok": False, "error": "缺少 correction 对象"}), 400
    try:
        from datetime import datetime, timezone

        correction = dict(correction)
        correction["updated_at"] = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
        _set_file_detect_correction(file_id, correction)
        from sign_handlers.file_session_cache import trim_detect_correction

        rule_sync = _sync_detect_correction_to_rules(file_id, correction, export_md=False)
        md_export = _export_detect_correction_markdown_after_sync([correction])
        return jsonify(
            {
                "ok": True,
                "correction": trim_detect_correction(correction),
                "rule_sync": rule_sync,
                "md_export": md_export,
            }
        )
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/files/detect-correction/batch", methods=["PUT"])
def api_sign_detect_correction_batch():
    """批量保存识别纠正登记（单次请求，避免前端逐条 PUT 超时）。"""
    data = request.get_json(silent=True) or {}
    items = data.get("items")
    if not isinstance(items, list) or not items:
        return jsonify({"ok": False, "error": "缺少 items 数组"}), 400
    if len(items) > 500:
        return jsonify({"ok": False, "error": "单次最多登记 500 个文件"}), 400
    from datetime import datetime, timezone

    ts = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    ok_ids = []
    failed = []
    batch_pairs = []
    for it in items:
        if not isinstance(it, dict):
            failed.append({"file_id": "", "error": "条目格式无效"})
            continue
        fid = str(it.get("file_id") or "").strip()
        corr = it.get("correction")
        if not _SIGN_FILE_ID_RE.match(fid):
            failed.append({"file_id": fid, "error": "无效的文件 id"})
            continue
        if not isinstance(corr, dict):
            failed.append({"file_id": fid, "error": "缺少 correction"})
            continue
        payload = dict(corr)
        payload["updated_at"] = ts
        batch_pairs.append((fid, payload))
        ok_ids.append(fid)
    if batch_pairs:
        try:
            if _sign_using_mysql():
                from sign_handlers import mysql_store

                mysql_store.ensure_sign_mysql()
                mysql_store.set_file_detect_correction_batch(batch_pairs)
            else:
                for fid, payload in batch_pairs:
                    _set_file_detect_correction(fid, payload)
        except Exception as e:
            return jsonify(
                {
                    "ok": False,
                    "error": "批量写入失败：" + str(e),
                    "ok_count": 0,
                    "fail_count": len(batch_pairs),
                }
            ), 500
    rule_syncs = []
    for fid, payload in batch_pairs:
        rs = _sync_detect_correction_to_rules(fid, payload, export_md=False)
        rs["file_id"] = fid
        rule_syncs.append(rs)
    md_export = _export_detect_correction_markdown_after_sync(
        [p for _fid, p in batch_pairs]
    )
    return jsonify(
        {
            "ok": True,
            "ok_count": len(ok_ids),
            "fail_count": len(failed),
            "ok_ids": ok_ids,
            "failed": failed[:50],
            "rule_syncs": rule_syncs[:50],
            "md_export": md_export,
        }
    )


@app.route("/api/sign/detect-correction/reference-image", methods=["POST"])
def api_sign_detect_correction_image_upload_shared():
    """上传共享参考图（FTP），供批量登记时多条记录共用同一 ftp_path。"""
    blob, ext, raw_name, err = _parse_detect_correction_image_upload()
    if err:
        return jsonify({"ok": False, "error": err}), 400
    try:
        from sign_handlers.detect_correction_storage import upload_reference_image_bytes

        meta = upload_reference_image_bytes(
            blob, ext, shared=True, filename=raw_name
        )
        return jsonify({"ok": True, "image": meta})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route(
    "/api/sign/files/<file_id>/detect-correction/reference-image",
    methods=["POST"],
)
def api_sign_file_detect_correction_image_upload(file_id):
    """上传「正确签批样式」参考截图至 FTP（PNG/JPEG/WebP，单张≤2MB，每文件最多6张）。"""
    if not _SIGN_FILE_ID_RE.match(file_id or ""):
        return jsonify({"ok": False, "error": "无效的文件 id"}), 400
    blob, ext, raw_name, err = _parse_detect_correction_image_upload()
    if err:
        return jsonify({"ok": False, "error": err}), 400
    try:
        from datetime import datetime, timezone
        from sign_handlers.detect_correction_storage import upload_reference_image_bytes

        corr = _get_file_detect_correction(file_id)
        imgs = list(corr.get("reference_images") or [])
        if len(imgs) >= 6:
            return jsonify({"ok": False, "error": "每文件最多保存 6 张参考图"}), 400
        meta = upload_reference_image_bytes(
            blob, ext, file_id=file_id, shared=False, filename=raw_name
        )
        imgs.append(meta)
        corr["reference_images"] = imgs
        corr["updated_at"] = meta.get("uploaded_at") or datetime.now(
            timezone.utc
        ).strftime("%Y-%m-%dT%H:%M:%SZ")
        _set_file_detect_correction(file_id, corr)
        return jsonify({"ok": True, "image": meta})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route(
    "/api/sign/files/<file_id>/detect-correction/reference-image/<image_id>",
    methods=["GET", "DELETE"],
)
def api_sign_file_detect_correction_image(file_id, image_id):
    if not _SIGN_FILE_ID_RE.match(file_id or ""):
        return jsonify({"ok": False, "error": "无效的文件 id"}), 400
    if not re.match(r"^[0-9a-f]{8,32}$", str(image_id or "")):
        return jsonify({"ok": False, "error": "无效的图片 id"}), 400
    corr = _get_file_detect_correction(file_id)
    meta = None
    for x in corr.get("reference_images") or []:
        if str(x.get("id") or "") == image_id:
            meta = x
            break
    if request.method == "DELETE":
        try:
            from sign_handlers.detect_correction_storage import (
                delete_reference_image_on_ftp,
            )

            if meta:
                delete_reference_image_on_ftp(meta, file_id=file_id)
            imgs = [
                x
                for x in (corr.get("reference_images") or [])
                if str(x.get("id") or "") != image_id
            ]
            corr["reference_images"] = imgs
            _set_file_detect_correction(file_id, corr)
            return jsonify({"ok": True})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500
    ftp_path_q = (request.args.get("ftp_path") or "").strip()
    try:
        from sign_handlers.detect_correction_storage import fetch_reference_image_bytes

        data, mime = fetch_reference_image_bytes(
            meta=meta,
            file_id=file_id,
            image_id=image_id,
            ftp_path_override=ftp_path_q,
        )
        return Response(data, mimetype=mime)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 404
    from sign_handlers.detect_correction_storage import legacy_local_path

    legacy = legacy_local_path(SIGN_DETECT_CORRECTIONS_ROOT, file_id, image_id)
    if legacy and os.path.isfile(legacy):
        mime = "image/png"
        if legacy.endswith(".jpg") or legacy.endswith(".jpeg"):
            mime = "image/jpeg"
        elif legacy.endswith(".webp"):
            mime = "image/webp"
        return send_from_directory(
            os.path.dirname(legacy), os.path.basename(legacy), mimetype=mime
        )
    return jsonify({"ok": False, "error": "图片不存在"}), 404


def _sign_project_from_request() -> tuple[Optional[dict], Optional[tuple]]:
    """解析上传/关联时选择的项目。返回 (project_dict, error_response)。"""
    project_id = (request.form.get("project_id") or "").strip()
    if not project_id and request.is_json:
        data = request.get_json(silent=True) or {}
        project_id = str(data.get("project_id") or "").strip()
    if not project_id:
        return None, (jsonify({"ok": False, "error": "请先选择关联项目"}), 400)
    try:
        from sign_handlers import mysql_store

        if mysql_store.mysql_sign_enabled():
            proj = mysql_store.get_project_by_id(project_id)
            if not proj:
                return None, (
                    jsonify({"ok": False, "error": "项目不存在，请先从 aiword 同步项目列表"}),
                    400,
                )
            return proj, None
    except Exception:
        pass
    from sign_handlers import project_store_local

    proj = project_store_local.get_project_by_id(project_id)
    if not proj:
        return None, (
            jsonify({"ok": False, "error": "项目不存在，请先从 aiword 同步项目列表"}),
            400,
        )
    return proj, None


def _sign_attach_project_to_record(rec: dict, proj: dict) -> None:
    if not rec or not proj:
        return
    rec["project_id"] = str(proj.get("id") or "").strip() or None
    rec["project_name"] = str(proj.get("name") or "").strip() or None
    rec["project_key"] = str(proj.get("project_key") or proj.get("name") or "").strip() or None
    rec["project_label"] = str(proj.get("label") or proj.get("name") or "").strip() or None


def _sign_apply_ctx_project_to_file(file_id: str, ctx: dict, rec: Optional[dict] = None) -> None:
    pid = str((ctx or {}).get("project_id") or "").strip()
    if not pid:
        return
    pname = str((ctx or {}).get("project_name") or "").strip() or None
    pkey = str((ctx or {}).get("project_key") or pname or "").strip() or None
    if _sign_using_mysql():
        try:
            from sign_handlers import mysql_store

            mysql_store.set_file_project(
                file_id, project_id=pid, project_name=pname, project_key=pkey
            )
        except Exception:
            pass
        return
    if rec is not None:
        rec["project_id"] = pid
        rec["project_name"] = pname
        rec["project_key"] = pkey
        rec["project_label"] = pname or pkey


@app.route("/api/sign/projects", methods=["GET"])
def api_sign_projects_list():
    """本地缓存的项目列表（含待签文件数）。"""
    try:
        from sign_handlers import mysql_store

        if mysql_store.mysql_sign_enabled():
            mysql_store.ensure_sign_mysql()
            items = mysql_store.list_projects_with_counts()
            return jsonify({"ok": True, "projects": items})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

    from sign_handlers import project_store_local

    _sign_ensure_session_inbox()
    pruned = _sign_prune_session_files_to_disk(session["sign_inbox_sid"])
    items = project_store_local.list_projects_with_counts(pruned)
    return jsonify({"ok": True, "projects": items})


@app.route("/api/sign/projects/sync", methods=["POST"])
def api_sign_projects_sync():
    """从 aiword 拉取并刷新项目缓存。"""
    from sign_handlers.aiword_projects import sync_projects_to_store

    st = sync_projects_to_store()
    if not st.get("ok"):
        return jsonify(st), 400
    items = st.get("projects")
    if not isinstance(items, list):
        try:
            from sign_handlers import mysql_store

            if mysql_store.mysql_sign_enabled():
                items = mysql_store.list_projects_with_counts(include_file_counts=False)
            else:
                from sign_handlers import project_store_local

                _sign_ensure_session_inbox()
                pruned = _sign_prune_session_files_to_disk(session["sign_inbox_sid"])
                items = project_store_local.list_projects_with_counts(pruned)
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500
    return jsonify(
        {
            "ok": True,
            "count": st.get("count", 0),
            "inserted": st.get("inserted", 0),
            "updated": st.get("updated", 0),
            "storage": st.get("storage"),
            "projects": items,
        }
    )


@app.route("/api/sign/files/<file_id>/project", methods=["PUT"])
def api_sign_file_set_project(file_id):
    if not _SIGN_FILE_ID_RE.match(file_id or ""):
        return jsonify({"ok": False, "error": "无效 file_id"}), 400
    data = request.get_json(silent=True) or {}
    project_id = str(data.get("project_id") or "").strip()
    if _sign_using_mysql():
        try:
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            if not project_id:
                n = mysql_store.set_file_project(file_id, project_id=None)
                if not n:
                    return jsonify({"ok": False, "error": "文件不存在"}), 404
                files = mysql_store.list_files()
                return jsonify({"ok": True, "files": files})
            n = mysql_store.set_file_project(file_id, project_id=project_id)
            if not n:
                return jsonify({"ok": False, "error": "文件不存在"}), 404
            files = mysql_store.list_files()
            return jsonify({"ok": True, "files": files})
        except ValueError as ve:
            return jsonify({"ok": False, "error": str(ve)}), 400
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    _sign_ensure_session_inbox()
    rec = _sign_find_record(file_id)
    if not rec:
        return jsonify({"ok": False, "error": "文件不存在"}), 404
    if not project_id:
        rec["project_id"] = None
        rec["project_name"] = None
        rec["project_key"] = None
        rec["project_label"] = None
        session.modified = True
        sid = session["sign_inbox_sid"]
        pruned = _sign_prune_session_files_to_disk(sid)
        return jsonify({"ok": True, "files": pruned})

    from sign_handlers import project_store_local

    proj = project_store_local.get_project_by_id(project_id)
    if not proj:
        return jsonify({"ok": False, "error": "项目不存在"}), 400
    _sign_attach_project_to_record(rec, proj)
    session.modified = True
    sid = session["sign_inbox_sid"]
    pruned = _sign_prune_session_files_to_disk(sid)
    return jsonify({"ok": True, "files": pruned})


@app.route("/api/sign/files/batch-project", methods=["PUT"])
def api_sign_files_batch_project():
    data = request.get_json(silent=True) or {}
    arr = data.get("file_ids")
    project_id = str(data.get("project_id") or "").strip()
    if not isinstance(arr, list) or not arr:
        return jsonify({"ok": False, "error": "缺少 file_ids"}), 400
    if not project_id:
        return jsonify({"ok": False, "error": "缺少 project_id"}), 400
    ids = []
    seen = set()
    for x in arr:
        fid = str(x or "").strip()
        if not fid or fid in seen or not _SIGN_FILE_ID_RE.match(fid):
            continue
        seen.add(fid)
        ids.append(fid)
    if not ids:
        return jsonify({"ok": False, "error": "file_ids 无效"}), 400
    if _sign_using_mysql():
        try:
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            n = mysql_store.set_files_project(ids, project_id)
            return jsonify({"ok": True, "updated": n, "files": mysql_store.list_files()})
        except ValueError as ve:
            return jsonify({"ok": False, "error": str(ve)}), 400
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    from sign_handlers import project_store_local

    proj = project_store_local.get_project_by_id(project_id)
    if not proj:
        return jsonify({"ok": False, "error": "项目不存在"}), 400
    _sign_ensure_session_inbox()
    updated = 0
    for fid in ids:
        rec = _sign_find_record(fid)
        if not rec:
            continue
        _sign_attach_project_to_record(rec, proj)
        updated += 1
    session.modified = True
    pruned = _sign_prune_session_files_to_disk(session["sign_inbox_sid"])
    return jsonify({"ok": True, "updated": updated, "files": pruned})


@app.route("/api/sign/files", methods=["GET"])
def api_sign_files_list():
    """List pending files for signature."""
    if _sign_using_mysql():
        try:
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            files = mysql_store.list_files()
            return jsonify({"ok": True, "files": files})
        except Exception as e:
            return jsonify({"ok": False, "error": f"MySQL ????: {e}"}), 500
    _sign_ensure_session_inbox()
    sid = session["sign_inbox_sid"]
    pruned = _sign_prune_session_files_to_disk(sid)
    files = []
    for rec in pruned:
        rec = _sign_enrich_session_file_record(sid, rec)
        files.append(
            {
                "id": rec.get("id"),
                "name": rec.get("name"),
                "ext": rec.get("ext"),
                "created_at": rec.get("created_at"),
                "saved_at": rec.get("saved_at") or rec.get("created_at"),
                "project_id": rec.get("project_id"),
                "project_name": rec.get("project_name"),
                "project_key": rec.get("project_key"),
                "project_label": rec.get("project_label") or rec.get("project_name"),
            }
        )
    return jsonify({"ok": True, "files": files})


def _sign_check_upload_duplicates(project_id: str, names: list, has_archive: bool) -> dict:
    """检查待上传文件名在当前项目下是否已存在（按 basename）。"""
    clean_names = [
        _sign_upload_base_name(str(n or ""))
        for n in (names or [])
        if str(n or "").strip()
    ]
    clean_names = [n for n in clean_names if n]
    if has_archive and not clean_names:
        return {"duplicates": [], "archive_deferred": True, "names_checked": []}
    if _sign_using_mysql():
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        dups = mysql_store.find_duplicate_basenames_in_project(project_id or None, clean_names)
        return {
            "duplicates": dups,
            "archive_deferred": bool(has_archive and not clean_names),
            "names_checked": clean_names,
        }
    _sign_ensure_session_inbox()
    sid = session["sign_inbox_sid"]
    pruned = _sign_prune_session_files_to_disk(sid)
    want_bases = set(clean_names)
    hits = []
    seen = set()
    for rec in pruned:
        if not _sign_record_same_project(rec, {"id": project_id}):
            continue
        base = _sign_upload_base_name(str(rec.get("name") or ""))
        if base in want_bases and base not in seen:
            seen.add(base)
            hits.append(base)
    return {
        "duplicates": sorted(hits),
        "archive_deferred": bool(has_archive and not clean_names),
        "names_checked": clean_names,
    }


@app.route("/api/sign/upload/resolve-duplicates", methods=["POST"])
def api_sign_upload_resolve_duplicates():
    """覆盖上传：删除同项目同名旧文件，保留 keep_file_ids（新上传）。"""
    data = request.get_json(silent=True) or {}
    project_id = str(data.get("project_id") or "").strip()
    if not project_id:
        return jsonify({"ok": False, "error": "缺少 project_id"}), 400
    names = data.get("duplicate_names") or data.get("names") or []
    if not isinstance(names, list) or not names:
        return jsonify({"ok": False, "error": "缺少 duplicate_names"}), 400
    keep_ids = data.get("keep_file_ids") or []
    if not isinstance(keep_ids, list):
        keep_ids = []
    if _sign_using_mysql():
        try:
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            mysql_store.migrate_file_replace_caches_to_keep_ids(
                project_id, names, keep_ids
            )
            deleted = mysql_store.delete_inbox_by_basenames_in_project_except(
                project_id, names, keep_ids
            )
            return jsonify(
                {
                    "ok": True,
                    "deleted_ids": deleted,
                    "files": mysql_store.list_files(),
                }
            )
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500
    _sign_ensure_session_inbox()
    sid = session["sign_inbox_sid"]
    from sign_handlers.filename_util import normalize_display_filename

    bases = {normalize_display_filename(str(n or "")) for n in names}
    keep = {str(x or "").strip() for x in keep_ids}
    records = list(session.get("sign_files") or [])
    old_by_base: dict = {}
    new_by_base: dict = {}
    for rec in records:
        if not _sign_record_same_project(rec, {"id": project_id}):
            continue
        fid = str(rec.get("id") or "").strip()
        base = normalize_display_filename(str(rec.get("name") or ""))
        if base not in bases or not fid:
            continue
        if fid in keep:
            new_by_base[base] = fid
        else:
            old_by_base[base] = fid
    for base in bases:
        old_fid = old_by_base.get(base)
        new_fid = new_by_base.get(base)
        if not old_fid or not new_fid or old_fid == new_fid:
            continue
        from sign_handlers.sign_library_local import get_file_role_map, list_file_session_caches
        from sign_handlers.sign_library_local import (
            set_file_detect_correction,
            set_file_detect_snapshot,
            set_file_role_map,
            set_file_workbench_state,
        )

        inbox_dir = os.path.join(SIGN_INBOX_ROOT, sid)
        ent = dict((list_file_session_caches(inbox_dir, sid, lite=False) or {}).get(old_fid) or {})
        role_map = get_file_role_map(inbox_dir, sid, old_fid)
        if role_map:
            ent["map"] = role_map
        if ent.get("detect"):
            set_file_detect_snapshot(inbox_dir, sid, new_fid, ent["detect"])
        if ent.get("workbench"):
            set_file_workbench_state(inbox_dir, sid, new_fid, ent["workbench"])
        if ent.get("detect_correction"):
            set_file_detect_correction(inbox_dir, sid, new_fid, ent["detect_correction"])
        if ent.get("map"):
            set_file_role_map(inbox_dir, sid, new_fid, ent["map"])
    deleted = []
    kept_records = []
    for rec in records:
        fid = str(rec.get("id") or "")
        base = normalize_display_filename(str(rec.get("name") or ""))
        if (
            base in bases
            and _sign_record_same_project(rec, {"id": project_id})
            and fid
            and fid not in keep
        ):
            _sign_remove_disk_files_for_id(sid, fid, rec.get("ext"))
            deleted.append(fid)
        else:
            kept_records.append(rec)
    session["sign_files"] = kept_records
    session.modified = True
    pruned = _sign_prune_session_files_to_disk(sid)
    return jsonify({"ok": True, "deleted_ids": deleted, "files": pruned})


@app.route("/api/sign/upload/check-duplicates", methods=["POST"])
def api_sign_upload_check_duplicates():
    """上传前检查同项目重名（与入库 basename 规则一致）。支持 JSON 或 multipart（含压缩包时服务端解压后按内层文件名判重）。"""
    project_id = str(request.form.get("project_id") or "").strip()
    data = request.get_json(silent=True) or {}
    if not project_id:
        project_id = str(data.get("project_id") or "").strip()
    if not project_id:
        return jsonify({"ok": False, "error": "请先选择关联项目"}), 400

    uploads = request.files.getlist("files")
    if not uploads or not any(f and f.filename for f in uploads):
        one = request.files.get("file")
        if one and one.filename:
            uploads = [one]
    if uploads and any(f and f.filename for f in uploads):
        try:
            parsed, warnings, upload_has_archive, _, _ = _sign_collect_upload_items(uploads)
            names = [str(item.get("name") or "").strip() for item in parsed if item.get("name")]
            info = _sign_check_upload_duplicates(
                project_id, names, has_archive=bool(upload_has_archive)
            )
            if warnings:
                info["warnings"] = warnings[:20]
            return jsonify({"ok": True, **info})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    names = data.get("names")
    if not isinstance(names, list):
        return jsonify({"ok": False, "error": "缺少 names 数组或未上传 files"}), 400
    has_archive = bool(data.get("has_archive"))
    try:
        norm_names = [_sign_upload_base_name(str(n or "")) for n in names if str(n or "").strip()]
        info = _sign_check_upload_duplicates(project_id, norm_names, has_archive)
        return jsonify({"ok": True, **info})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/upload", methods=["POST"])
@sign_heavy_route
def api_sign_upload():
    """Save one or multiple sign docs, supports zip expansion."""
    t0 = time.time()
    uploads = request.files.getlist("files")
    if not uploads or not any(f.filename for f in uploads):
        one = request.files.get("file")
        if one and one.filename:
            uploads = [one]
    if not uploads or not any(f.filename for f in uploads):
        return jsonify({"ok": False, "error": "?????"}), 400

    parsed, warnings, upload_has_archive, archive_expanded, archive_summary = _sign_collect_upload_items(uploads)
    logger.info(
        "sign upload collect done in %.1fs: items=%d archive=%s expanded=%d",
        time.time() - t0,
        len(parsed),
        upload_has_archive,
        archive_expanded,
    )
    if not parsed:
        return jsonify(
            {
                "ok": False,
                "error": "未找到可签字文档。请上传 .doc/.docx/.docm/.xls/.xlsx/.xlsm，或包含这些文件的压缩包。",
                "warnings": warnings,
                "archive_summary": archive_summary,
            }
        ), 400

    proj, proj_err = _sign_project_from_request()
    if proj_err:
        return proj_err

    replace_dup = (request.form.get("replace_duplicates") or "").strip().lower() in (
        "1",
        "true",
        "yes",
        "on",
    )
    dup_before_local: list = []
    if not replace_dup:
        dup_before_local = _sign_check_upload_duplicates(
            str(proj.get("id") or ""),
            [item.get("name") or "" for item in parsed],
            upload_has_archive,
        ).get("duplicates") or []

    if _sign_using_mysql():
        try:
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            cur_n = mysql_store.count_files()
            _max_f = _sign_mysql_max_files()
            if cur_n + len(parsed) > _max_f:
                return jsonify(
                    {
                        "ok": False,
                        "error": (
                            f"当前已保存 {cur_n} 个，本次将新增 {len(parsed)} 个，"
                            f"超过上限 {_max_f}。请先删除部分文件后重试。"
                        ),
                        "warnings": warnings,
                        "archive_summary": archive_summary,
                    }
                ), 400
            dup_before: list = []
            if not replace_dup:
                dup_before = mysql_store.find_duplicate_basenames_in_project(
                    str(proj.get("id") or "") or None,
                    [item.get("name") or "" for item in parsed],
                )
            added_ids, last_rec, ins_err = _sign_mysql_insert_parsed_items(
                parsed, proj, replace_duplicates=replace_dup
            )
            if ins_err and not added_ids:
                return jsonify({"ok": False, "error": ins_err, "warnings": warnings}), 500
            if ins_err:
                warnings.append(ins_err)
            logger.info(
                "sign upload mysql insert done in %.1fs: added=%d replace_dup=%s",
                time.time() - t0,
                len(added_ids),
                replace_dup,
            )
            return jsonify(
                {
                    "ok": True,
                    "file": last_rec,
                    "files": mysql_store.list_files(),
                    "added": len(parsed),
                    "added_ids": added_ids,
                    "replaced_duplicates": replace_dup,
                    "duplicate_names": dup_before if not replace_dup else [],
                    "warnings": warnings,
                    "upload_has_archive": upload_has_archive,
                    "archive_expanded": archive_expanded,
                    "archive_summary": archive_summary,
                }
            )
        except Exception as e:
            return jsonify({"ok": False, "error": f"MySQL ????: {e}"}), 500

    sid, inbox_dir = _sign_ensure_session_inbox()
    records = list(session.get("sign_files") or [])
    captured_replace: dict = {}
    if replace_dup:
        bases = {_sign_upload_base_name(item.get("name") or "") for item in parsed}
        bases.discard("")
        if bases:
            captured_replace = _sign_session_capture_replace_caches(
                sid, inbox_dir, proj, parsed
            )
            records = [
                r
                for r in records
                if not (
                    _sign_upload_base_name(r.get("name") or "") in bases
                    and _sign_record_same_project(r, proj)
                )
            ]
    if len(records) + len(parsed) > SIGN_MAX_SAVED_FILES:
        return jsonify(
            {
                "ok": False,
                "error": (
                    f"当前已保存 {len(records)} 个，本次将新增 {len(parsed)} 个，"
                    f"超过上限 {SIGN_MAX_SAVED_FILES}。请先删除部分文件后重试。"
                ),
                "warnings": warnings,
                "archive_summary": archive_summary,
            }
        ), 400

    last_rec = None
    added_ids = []
    for item in parsed:
        display_name = item["name"]
        ext = item["ext"]
        raw = item["raw"]
        file_id = uuid.uuid4().hex
        dest = os.path.join(inbox_dir, file_id + ext)
        with open(dest, "wb") as fp:
            fp.write(raw)
        saved_ts = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
        last_rec = {
            "id": file_id,
            "name": display_name,
            "ext": ext,
            "created_at": saved_ts,
            "saved_at": saved_ts,
        }
        _sign_attach_project_to_record(last_rec, proj)
        records.append(last_rec)
        added_ids.append(file_id)
        caches = captured_replace.get(_sign_upload_base_name(display_name))
        if caches:
            try:
                _sign_session_apply_replace_caches(sid, inbox_dir, file_id, caches)
            except Exception as e:
                logger.warning(
                    "sign session replace cache migrate failed for %s: %s",
                    display_name,
                    e,
                )
    session["sign_files"] = records
    session.modified = True

    return jsonify(
        {
            "ok": True,
            "file": last_rec,
            "files": records,
            "added": len(parsed),
            "added_ids": added_ids,
            "replaced_duplicates": replace_dup,
            "duplicate_names": dup_before_local if not replace_dup else [],
            "warnings": warnings,
            "upload_has_archive": upload_has_archive,
            "archive_expanded": archive_expanded,
            "archive_summary": archive_summary,
        }
    )


@app.route("/api/sign/files/<file_id>", methods=["DELETE"])
def api_sign_file_delete(file_id):
    """Delete one file record from pending list."""
    if not _SIGN_FILE_ID_RE.match(file_id or ""):
        return jsonify({"ok": False, "error": "??????id"}), 400
    if _sign_using_mysql():
        try:
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            n = mysql_store.delete_file(file_id)
            if not n:
                return jsonify({"ok": False, "error": "??????"}), 404
            return jsonify({"ok": True, "files": mysql_store.list_files()})
        except Exception as e:
            return jsonify({"ok": False, "error": f"MySQL ????: {e}"}), 500
    _sign_ensure_session_inbox()
    sid = session["sign_inbox_sid"]
    rec = _sign_find_record(file_id)
    if not rec:
        return jsonify({"ok": False, "error": "??????"}), 404
    _sign_remove_disk_files_for_id(sid, file_id, rec.get("ext"))
    fid = str(file_id)
    session["sign_files"] = [
        r for r in (session.get("sign_files") or []) if str(r.get("id") or "") != fid
    ]
    session.modified = True
    pruned = _sign_prune_session_files_to_disk(sid)
    return jsonify({"ok": True, "files": pruned})


@app.route("/api/sign/files/batch-delete", methods=["POST"])
def api_sign_files_batch_delete():
    """Batch delete file records from pending sign list."""
    data = request.get_json(silent=True) or {}
    arr = data.get("file_ids")
    if not isinstance(arr, list) or not arr:
        return jsonify({"ok": False, "error": "缺少 file_ids（数组）"}), 400

    dedup = []
    seen = set()
    invalid_ids = []
    for x in arr:
        fid = str(x or "").strip()
        if not fid:
            continue
        if fid in seen:
            continue
        seen.add(fid)
        if not _SIGN_FILE_ID_RE.match(fid):
            invalid_ids.append(fid)
            continue
        dedup.append(fid)
    if not dedup and invalid_ids:
        return jsonify({"ok": False, "error": "file_ids 均无效", "invalid_ids": invalid_ids}), 400

    deleted_ids = []
    missing_ids = []
    refresh_files = data.get("refresh_files", True)
    if isinstance(refresh_files, str):
        refresh_files = refresh_files.strip().lower() not in ("0", "false", "no")
    if _sign_using_mysql():
        try:
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            deleted_ids, missing_ids, delete_errors = mysql_store.delete_files_batch(
                dedup, skip_ftp_cleanup=True
            )
            files_out = mysql_store.list_files() if refresh_files else None
            return jsonify(
                {
                    "ok": True,
                    "deleted_ids": deleted_ids,
                    "missing_ids": missing_ids,
                    "invalid_ids": invalid_ids,
                    "delete_errors": delete_errors,
                    "files": files_out,
                }
            )
        except Exception as e:
            return jsonify({"ok": False, "error": f"MySQL 删除失败: {e}"}), 500

    _sign_ensure_session_inbox()
    sid = session["sign_inbox_sid"]
    records = list(session.get("sign_files") or [])
    by_id = {str(r.get("id") or ""): r for r in records}
    for fid in dedup:
        rec = by_id.get(fid)
        if not rec:
            missing_ids.append(fid)
            continue
        _sign_remove_disk_files_for_id(sid, fid, rec.get("ext"))
        deleted_ids.append(fid)

    if deleted_ids:
        deleted_set = set(deleted_ids)
        session["sign_files"] = [
            r for r in (session.get("sign_files") or []) if str(r.get("id") or "") not in deleted_set
        ]
        session.modified = True
    pruned = _sign_prune_session_files_to_disk(sid)
    return jsonify(
        {
            "ok": True,
            "deleted_ids": deleted_ids,
            "missing_ids": missing_ids,
            "invalid_ids": invalid_ids,
            "files": pruned,
        }
    )


def _humanize_office_package_error(exc) -> str:
    """将 python-docx/openpyxl 的 Package not found 转为可理解的业务提示。"""
    msg = str(exc).strip()
    if "Package not found" in msg:
        return (
            "解析失败：临时文件不是有效的 Word/Excel 文档（Package not found）。"
            "启用 MySQL 时系统会先从 FTP 下载到本机临时目录再识别，并非读取您电脑上的源文件路径；"
            "若持续出现，请检查 FTP 上该文件是否完整，或在列表中删除后重新上传。"
        )
    return msg or type(exc).__name__


@app.route("/api/sign/detect", methods=["GET"])
@sign_detect_route
def api_sign_detect():
    """自动识别文档中的签名位/角色（用于前端自动勾选）。"""
    file_id = (request.args.get("file_id") or "").strip()
    detect_mode_req = (request.args.get("mode") or "auto").strip().lower()
    if detect_mode_req not in ("auto", "full", "light"):
        detect_mode_req = "auto"
    if not file_id or not _SIGN_FILE_ID_RE.match(file_id):
        return jsonify({"ok": False, "error": "缺少或无效的 file_id"}), 400

    ext = None
    source_path = None
    mysql_blob = None

    if _sign_using_mysql():
        try:
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            row = mysql_store.get_file_row(file_id)
            if not row:
                return jsonify({"ok": False, "error": "未找到该文件"}), 404
            if not row.get("file_data"):
                err = (row.get("file_load_error") or "").strip()
                return jsonify(
                    {
                        "ok": False,
                        "error": err
                        or "无法读取文件内容（可能 FTP 暂不可用或下载失败）。请稍后重试「重新识别」或检查 FTP 配置。",
                    }
                ), 502
            ext = (row.get("ext") or ".docx").lower()
            if ext not in SIGN_ALLOWED_EXT:
                return jsonify({"ok": False, "error": "不支持的文件类型"}), 400
            mysql_blob = row["file_data"]
        except Exception as e:
            return jsonify({"ok": False, "error": f"MySQL 读取失败: {e}"}), 500
    else:
        _sign_ensure_session_inbox()
        sid = session["sign_inbox_sid"]
        rec = _sign_find_record(file_id)
        if not rec:
            return jsonify({"ok": False, "error": "未找到该文件"}), 404
        ext = (rec.get("ext") or ".docx").lower()
        if ext not in SIGN_ALLOWED_EXT:
            return jsonify({"ok": False, "error": "不支持的文件类型"}), 400
        source_path = _sign_saved_disk_path(sid, file_id, ext)
        if not os.path.isfile(source_path):
            return jsonify({"ok": False, "error": "文件已不存在"}), 404

    tmp_dir = tempfile.mkdtemp(prefix="aiprintword_detect_")
    detect_fut = None
    try:
        in_path = os.path.join(tmp_dir, f"{uuid.uuid4().hex[:8]}_detect{ext}")
        if mysql_blob is not None:
            with open(in_path, "wb") as fp:
                fp.write(mysql_blob)
        else:
            shutil.copy2(source_path, in_path)

        if _sign_using_mysql():
            from sign_handlers.mysql_store import is_valid_uploaded_blob

            with open(in_path, "rb") as _chk:
                on_disk = _chk.read()
            if not is_valid_uploaded_blob(on_disk, ext):
                return jsonify(
                    {
                        "ok": False,
                        "error": (
                            "写入临时文件后校验失败：不是有效的 Office 文档。"
                            "请检查 FTP 上文件是否完整，或在列表中重新上传。"
                        ),
                    }
                ), 502

        from sign_handlers.detect_fields import detect_file

        src_name = ""
        if _sign_using_mysql():
            src_name = (row.get("name") or "").strip() if row else ""
        else:
            src_name = (rec.get("name") or "").strip() if rec else ""

        try:
            from runtime_settings.resolve import get_setting

            op_timeout = int(float(get_setting("SIGN_DETECT_OP_TIMEOUT_SEC")))
        except Exception:
            try:
                op_timeout = int(
                    (os.environ.get("SIGN_DETECT_OP_TIMEOUT_SEC") or "3600").strip()
                    or "3600"
                )
            except ValueError:
                op_timeout = 3600
        op_timeout = max(60, min(op_timeout, 43200))
        try:
            from runtime_settings.resolve import get_setting

            light_timeout = int(float(get_setting("SIGN_DETECT_LIGHT_OP_TIMEOUT_SEC")))
        except Exception:
            light_timeout = max(60, min(op_timeout, 900))
        light_timeout = max(30, min(light_timeout, 7200))
        correction = _get_file_detect_correction(file_id)
        detect_hint = {}
        try:
            from sign_handlers.detect_correction import build_detect_hint

            detect_hint = build_detect_hint(correction)
        except Exception:
            detect_hint = {}

        def _run_detect(mode_name: str):
            mode_name = (mode_name or "auto").strip().lower()
            if mode_name not in ("auto", "full", "light"):
                mode_name = "auto"
            if mode_name == "full":
                mode_name = "auto"
            return detect_file(
                in_path, source_name=src_name, mode=mode_name, detect_hint=detect_hint
            )

        result = None
        used_detect_mode = detect_mode_req
        fallback_from_timeout = False
        try:
            from concurrent.futures import ThreadPoolExecutor, TimeoutError as _FutTimeout

            with ThreadPoolExecutor(max_workers=1) as pool:
                detect_fut = pool.submit(_run_detect, detect_mode_req)
                timeout_use = light_timeout if detect_mode_req == "light" else op_timeout
                result = detect_fut.result(timeout=timeout_use)
        except _FutTimeout:
            allow_fallback_to_light = detect_mode_req in ("auto", "full")
            if allow_fallback_to_light:
                try:
                    from concurrent.futures import (
                        ThreadPoolExecutor,
                        TimeoutError as _FutTimeout2,
                    )

                    with ThreadPoolExecutor(max_workers=1) as pool2:
                        detect_fut = pool2.submit(_run_detect, "light")
                        result = detect_fut.result(timeout=light_timeout)
                        used_detect_mode = "light"
                        fallback_from_timeout = True
                except _FutTimeout2:
                    result = {
                        "ok": False,
                        "error": (
                            f"识别超时（全量>{op_timeout} 秒，轻量>{light_timeout} 秒）。"
                            "系统已自动切换轻量模式（仅优先扫描文档前后页）但仍超时；"
                            "请稍后单独重试，或在系统设置中调大 SIGN_DETECT_OP_TIMEOUT_SEC / "
                            "SIGN_DETECT_LIGHT_OP_TIMEOUT_SEC。"
                        ),
                        "error_code": "detect_timeout",
                    }
                except Exception as e2:
                    result = {"ok": False, "error": _humanize_office_package_error(e2)}
            else:
                t_used = light_timeout if detect_mode_req == "light" else op_timeout
                result = {
                    "ok": False,
                    "error": (
                        f"识别超时（>{t_used} 秒）。该文档可能页数很多（如软件需求规范），"
                        "已中止本次解析；请稍后重试。"
                    ),
                    "error_code": "detect_timeout",
                }
        except Exception as e:
            result = {"ok": False, "error": _humanize_office_package_error(e)}

        if isinstance(result, dict):
            result = dict(result)
            result["detect_mode"] = used_detect_mode
            if fallback_from_timeout:
                result["light_fallback"] = True
            if not result.get("ok") and result.get("error"):
                result["error"] = _humanize_office_package_error(result["error"])
            result["file_id"] = file_id
            result["source_name"] = src_name
            _corr_wrong = str((correction or {}).get("wrong_description") or "").strip()
            _corr_slot = (correction or {}).get("expected_slot_layout")
            _corr_has_slot = isinstance(_corr_slot, dict) and bool(_corr_slot)
            try:
                from sign_handlers.sign_document_role_rules import apply_document_role_rules

                result = apply_document_role_rules(result, src_name)
            except Exception:
                pass
            try:
                from sign_handlers.detect_correction import apply_detect_correction

                if correction:
                    result = apply_detect_correction(
                        result, correction, source_name=src_name
                    )
            except Exception:
                pass
            blob_for_hash = mysql_blob
            if blob_for_hash is None and source_path and os.path.isfile(source_path):
                try:
                    with open(source_path, "rb") as hf:
                        blob_for_hash = hf.read()
                except OSError:
                    blob_for_hash = None
            if blob_for_hash is not None:
                import hashlib

                result["content_sha256"] = hashlib.sha256(blob_for_hash).hexdigest()
            if result.get("ok"):
                try:
                    from sign_handlers import ROLE_ID_TO_KEYWORD

                    from sign_handlers.config import canonical_sign_role_id

                    detect_ext = str(os.path.splitext(in_path)[1] or ext or "").lower()
                    kind = str(result.get("kind") or "").lower()
                    if kind == "docx":
                        detect_ext = ".docx"
                    elif kind == "xlsx":
                        detect_ext = ".xlsx"
                    if detect_ext not in (".docx", ".xlsx"):
                        detect_ext = ".docx" if ext in (".doc", ".docx") else ".xlsx"
                    probe_bytes = blob_for_hash
                    if not isinstance(probe_bytes, (bytes, bytearray)):
                        with open(in_path, "rb") as pf:
                            probe_bytes = pf.read()

                    role_ids_for_layout = _collect_layout_candidate_roles(result)
                    role_ids_for_layout = [
                        rid
                        for rid in role_ids_for_layout
                        if rid in ROLE_ID_TO_KEYWORD
                    ]
                    if role_ids_for_layout:
                        try:
                            from sign_handlers.signature_layout import (
                                analyze_signature_layout,
                            )

                            layout_info = analyze_signature_layout(
                                in_path,
                                detect_ext,
                                role_ids_for_layout,
                                source_name=src_name or "",
                            )
                            result["signature_layout"] = layout_info
                            result = _reconcile_detect_roles_with_layout(
                                result, layout_info, source_name=src_name or ""
                            )
                        except Exception as _layout_exc:
                            result["signature_layout"] = {
                                "ok": False,
                                "error": str(_layout_exc),
                                "role_layouts": {},
                            }
                    role_ids = []
                    seen = set()
                    for rr in (result.get("roles") or []):
                        if not isinstance(rr, dict):
                            continue
                        rid = canonical_sign_role_id(str(rr.get("id") or "").strip())
                        if not rid or rid in seen or rid not in ROLE_ID_TO_KEYWORD:
                            continue
                        seen.add(rid)
                        role_ids.append(rid)
                    if role_ids:
                        probe_plan = _build_placement_plan_from_detect(result, role_ids)
                        result["slot_probe"] = _probe_detect_slot_placement(
                            bytes(probe_bytes),
                            detect_ext,
                            src_name or (file_id + detect_ext),
                            role_ids,
                            placement_plan=probe_plan,
                        )
                except Exception:
                    pass
            if result.get("ok"):
                try:
                    if _sign_using_mysql():
                        from sign_handlers import mysql_store

                        mysql_store.set_file_detect_snapshot(file_id, result)
                    else:
                        _sign_ensure_session_inbox()
                        sid = session["sign_inbox_sid"]
                        from sign_handlers.sign_library_local import (
                            set_file_detect_snapshot as local_set_detect,
                        )

                        local_set_detect(SIGN_INBOX_ROOT, sid, file_id, result)
                except Exception:
                    pass
        return jsonify(result)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        try:
            if detect_fut is None or detect_fut.done():
                shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass


@app.route("/api/sign/server-status", methods=["GET"])
def api_sign_server_status():
    """轻量状态：列表类接口不受重任务槽位限制；供前端判断服务器是否繁忙。"""
    sem = _sign_heavy_semaphore()
    avail = getattr(sem, "_value", None)
    dsem = _sign_detect_semaphore()
    detect_avail = getattr(dsem, "_value", None)
    return jsonify(
        {
            "ok": True,
            "heavy_slots": _HEAVY_SEM_SIZE,
            "heavy_slots_available": avail,
            "detect_slots": _DETECT_SEM_SIZE,
            "detect_slots_available": detect_avail,
            "mysql": _sign_using_mysql(),
            "threaded": True,
        }
    )


@app.route("/api/sign/runtime-config", methods=["GET"])
def api_sign_runtime_config():
    """前端启动时读取的运行时配置（非敏感）。

    所有项均可在「系统设置」页面调整后立即对新打开页面生效，无需重启服务。
    取值优先级：数据库 > 环境变量 > 注册表默认值。
    """
    try:
        from runtime_settings.resolve import get_setting

        def _clamp_int(key: str, low: int, high: int, default: int) -> int:
            try:
                v = int(get_setting(key))
            except Exception:
                v = default
            if v < low:
                v = low
            if v > high:
                v = high
            return v

        return jsonify(
            {
                "ok": True,
                "sign_detect_timeout_ms": _clamp_int(
                    "SIGN_DETECT_TIMEOUT_MS", 30000, 86400000, 43200000
                ),
                "sign_detect_retry_times": _clamp_int(
                    "SIGN_DETECT_RETRY_TIMES", 0, 3, 1
                ),
                "sign_archive_upload_timeout_ms": _clamp_int(
                    "SIGN_ARCHIVE_UPLOAD_TIMEOUT_MS", 300000, 43200000, 43200000
                ),
                "sign_batch_file_timeout_ms": _clamp_int(
                    "SIGN_BATCH_FILE_TIMEOUT_MS", 60000, 43200000, 3600000
                ),
            }
        )
    except Exception as e:
        return jsonify(
            {
                "ok": True,
                "sign_detect_timeout_ms": 43200000,
                "sign_detect_retry_times": 1,
                "sign_archive_upload_timeout_ms": 43200000,
                "sign_batch_file_timeout_ms": 3600000,
                "error_hint": str(e),
            }
        )


@app.route("/api/sign", methods=["POST"])
@sign_heavy_route
def api_sign():
    """
    ??????????PNG ??????PNG??    ???????    - file_id?????????????????
    - file??????????????????    ?? roles (JSON)??????sig_{id}?date_{id}??    """
    file_id = (request.form.get("file_id") or "").strip()
    upload = request.files.get("file")

    ext = None
    base_name = None
    source_path = None
    mysql_blob = None

    if file_id:
        if not _SIGN_FILE_ID_RE.match(file_id):
            return jsonify({"ok": False, "error": "??????id"}), 400
        if _sign_using_mysql():
            try:
                from sign_handlers import mysql_store

                mysql_store.ensure_sign_mysql()
                row = mysql_store.get_file_row(file_id)
                if not row or not row.get("file_data"):
                    err = (row.get("file_load_error") or "").strip() if row else ""
                    return jsonify(
                        {"ok": False, "error": err or "未找到该文件或无法读取内容"}
                    ), 404
                ext = (row.get("ext") or ".docx").lower()
                if ext not in SIGN_ALLOWED_EXT:
                    return jsonify({"ok": False, "error": "????????"}), 400
                base_name = _safe_display_filename_keep_unicode(row.get("name") or "document") or "document"
                if not base_name.lower().endswith(ext):
                    base_name = os.path.splitext(base_name)[0] + ext
                mysql_blob = row["file_data"]
            except Exception as e:
                return jsonify({"ok": False, "error": f"MySQL ????: {e}"}), 500
        else:
            _sign_ensure_session_inbox()
            sid = session["sign_inbox_sid"]
            rec = _sign_find_record(file_id)
            if not rec:
                return jsonify({"ok": False, "error": "??????????????"}), 404
            ext = (rec.get("ext") or ".docx").lower()
            if ext not in SIGN_ALLOWED_EXT:
                return jsonify({"ok": False, "error": "????????"}), 400
            source_path = _sign_saved_disk_path(sid, file_id, ext)
            if not os.path.isfile(source_path):
                return jsonify({"ok": False, "error": "????????????"}), 404
            base_name = _safe_display_filename_keep_unicode(rec.get("name") or "document") or "document"
            if not base_name.lower().endswith(ext):
                base_name = os.path.splitext(base_name)[0] + ext
            mysql_blob = None
    else:
        if not upload or not upload.filename:
            return jsonify({"ok": False, "error": "????????????????????"}), 400
        ext = os.path.splitext(upload.filename)[1].lower()
        if ext not in SIGN_ALLOWED_EXT:
            return jsonify({"ok": False, "error": "????.docx ??.xlsx"}), 400

    roles_raw = request.form.get("roles") or "[]"
    try:
        roles = json.loads(roles_raw)
    except Exception:
        return jsonify({"ok": False, "error": "?? roles ???? JSON"}), 400
    if not isinstance(roles, list) or not roles:
        return jsonify({"ok": False, "error": "???????????"}), 400

    from sign_handlers import ROLE_ID_TO_KEYWORD
    from sign_handlers.config import role_display_name

    sign_source = (request.form.get("sign_source") or "canvas").strip().lower()
    if sign_source not in ("canvas", "library"):
        sign_source = "canvas"

    sig_map = {}
    date_map = {}
    apply_report = {"applied": [], "skipped": []}
    if sign_source == "library":
        if not file_id:
            return jsonify({"ok": False, "error": "库映射模式仅支持对“已保存到列表”的文件生成"}), 400
        if not _sign_using_mysql():
            return jsonify({"ok": False, "error": "库映射模式需要启用 MySQL（MYSQL_HOST）"}), 400
        try:
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            from sign_handlers.config import normalize_role_signer_map

            mapping = normalize_role_signer_map(
                mysql_store.get_file_role_signer_map(file_id) or {}
            )
            for rid in roles:
                if rid not in ROLE_ID_TO_KEYWORD:
                    return jsonify({"ok": False, "error": f"无效角色 id: {rid}"}), 400
                pair = mapping.get(rid) or {}
                sig_id = (pair.get("sig") if isinstance(pair, dict) else None) or None
                date_id = (pair.get("date") if isinstance(pair, dict) else None) or None
                dm = (pair.get("date_mode") if isinstance(pair, dict) else None) or None
                diso = (pair.get("date_iso") if isinstance(pair, dict) else None) or None
                if mysql_store.is_composite_date_mode(dm):
                    # 拼接日期模式需要：签名素材 + 日历日期；若条件不齐，退回“能签就签”的普通素材模式
                    if not sig_id or not diso:
                        apply_report["skipped"].append(
                            {
                                "role_id": rid,
                                "role": role_display_name(rid),
                                "what": "composite_date",
                                "reason": "拼接日期条件不齐（需绑定签名并选择日历日期），已退回普通素材模式",
                            }
                        )
                        dm = None
                    if dm:
                        srow = mysql_store.get_stroke_item_row(sig_id)
                        if not srow or not srow.get("png"):
                            apply_report["skipped"].append(
                                {
                                    "role_id": rid,
                                    "role": role_display_name(rid),
                                    "what": "sig",
                                    "reason": "签名素材缺失",
                                }
                            )
                            dm = None
                    if dm:
                        sid0 = (srow.get("signer_id") or "").strip()
                        if not sid0:
                            apply_report["skipped"].append(
                                {
                                    "role_id": rid,
                                    "role": role_display_name(rid),
                                    "what": "composite_date",
                                    "reason": "无法解析签署人",
                                }
                            )
                            dm = None
                    if dm:
                        try:
                            lay = mysql_store.composite_mode_to_layout(dm)
                            dbytes, _lbl = mysql_store.compose_date_piece_png(
                                sid0, str(diso).strip(), lay
                            )
                        except Exception as e:
                            apply_report["skipped"].append(
                                {
                                    "role_id": rid,
                                    "role": role_display_name(rid),
                                    "what": "composite_date",
                                    "reason": f"日期拼接失败：{e}",
                                }
                            )
                            dm = None
                    if dm:
                        sig_map[rid] = srow["png"]
                        date_map[rid] = dbytes
                        apply_report["applied"].append(
                            {
                                "role_id": rid,
                                "role": role_display_name(rid),
                                "sig": True,
                                "date": True,
                                "date_mode": "composite",
                            }
                        )
                        continue
                applied_sig = False
                applied_date = False
                if sig_id:
                    srow = mysql_store.get_stroke_item_row(sig_id)
                    if srow and srow.get("png"):
                        sig_map[rid] = srow["png"]
                        applied_sig = True
                if date_id:
                    drow = mysql_store.get_stroke_item_row(date_id)
                    if drow and drow.get("png"):
                        date_map[rid] = drow["png"]
                        applied_date = True
                if applied_sig or applied_date:
                    apply_report["applied"].append(
                        {
                            "role_id": rid,
                            "role": role_display_name(rid),
                            "sig": applied_sig,
                            "date": applied_date,
                            "date_mode": "item",
                        }
                    )
                else:
                    apply_report["skipped"].append(
                        {
                            "role_id": rid,
                            "role": role_display_name(rid),
                            "what": "sig_date",
                            "reason": "未绑定或素材缺失（签名/日期均无可用数据）",
                        }
                    )
        except Exception as e:
            return jsonify({"ok": False, "error": f"库映射读取失败: {e}"}), 500
    else:
        for rid in roles:
            if rid not in ROLE_ID_TO_KEYWORD:
                return jsonify({"ok": False, "error": f"无效角色 id: {rid}"}), 400
            sig_raw = request.form.get(f"sig_{rid}") or ""
            date_raw = request.form.get(f"date_{rid}") or ""
            sig_bytes = _decode_png_data_url_or_b64(sig_raw)
            date_bytes = _decode_png_data_url_or_b64(date_raw)
            applied_sig = False
            applied_date = False
            if sig_bytes:
                sig_map[rid] = sig_bytes
                applied_sig = True
            if date_bytes:
                date_map[rid] = date_bytes
                applied_date = True
            if applied_sig or applied_date:
                apply_report["applied"].append(
                    {
                        "role_id": rid,
                        "role": role_display_name(rid),
                        "sig": applied_sig,
                        "date": applied_date,
                        "date_mode": "canvas",
                    }
                )
            else:
                apply_report["skipped"].append(
                    {
                        "role_id": rid,
                        "role": role_display_name(rid),
                        "what": "sig_date",
                        "reason": "该角色签名与日期画布均为空，已跳过",
                    }
                )

    # 至少要有一个角色提供签名或日期；否则不生成（避免输出与原文档一致却误以为“已签”）
    if (not sig_map) and (not date_map):
        return (
            jsonify(
                {
                    "ok": False,
                    "error": "未选择任何可用素材：请至少为一个角色提供签名或日期（画布手写/库映射二选一）。",
                }
            ),
            400,
        )

    if _sign_using_mysql():
        try:
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            _max_s = _sign_mysql_max_signed()
            if mysql_store.count_signed_outputs() >= _max_s:
                return jsonify(
                    {
                        "ok": False,
                        "error": (
                            f"已签名记录已满（最多 {_max_s} 条），"
                            "请在本页「已签名文档」中删除旧记录后再试"
                        ),
                    }
                ), 400
        except Exception as e:
            return jsonify({"ok": False, "error": f"MySQL 检查失败: {e}"}), 500

    if not base_name and upload and upload.filename:
        base_name = _safe_display_filename_keep_unicode(os.path.basename(upload.filename)) or "document"
        if not base_name.lower().endswith(ext):
            base_name = os.path.splitext(base_name)[0] + ext

    if file_id:
        if mysql_blob is not None:
            file_bytes = mysql_blob
        else:
            with open(source_path, "rb") as fp:
                file_bytes = fp.read()
    else:
        upload.stream.seek(0)
        file_bytes = upload.read()
        if not base_name:
            base_name = _safe_display_filename_keep_unicode(os.path.basename(upload.filename)) or "document"
            if not base_name.lower().endswith(ext):
                base_name = base_name + ext

    placement_plan = None
    if file_id and _sign_using_mysql():
        try:
            from sign_handlers import mysql_store

            det = mysql_store.get_file_detect_snapshot(file_id)
            placement_plan = _build_placement_plan_from_detect(det or {}, roles)
        except Exception:
            placement_plan = None

    res = _sign_process_document_bytes(
        file_bytes,
        ext,
        base_name,
        roles,
        sig_map,
        date_map,
        file_id or None,
        batch_id=uuid.uuid4().hex,
        placement_plan=placement_plan,
        require_role_placement=False,
    )
    if not res.get("ok"):
        return jsonify({"ok": False, "error": res.get("error", "失败")}), 500

    dl_name = res["dl_name"]
    out_bytes = res["out_bytes"]
    resp = send_file(
        io.BytesIO(out_bytes),
        as_attachment=True,
        download_name=dl_name,
        mimetype=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if ext == ".docx"
            else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ),
    )
    sid_hdr = res.get("signed_id")
    if sid_hdr:
        resp.headers["X-Signed-Record-Id"] = sid_hdr
    try:
        # 更短更稳：只回传摘要（避免头部过大/截断导致前端解码失败）
        ap = apply_report.get("applied") or []
        sk = apply_report.get("skipped") or []
        ap_txt = "；".join(
            [
                f"{x.get('role') or x.get('role_id')}："
                + ("签名" if x.get("sig") else "")
                + ("+" if x.get("sig") and x.get("date") else "")
                + ("日期" if x.get("date") else "")
                + ("（拼接）" if x.get("date_mode") == "composite" else "")
                for x in ap
                if isinstance(x, dict)
            ]
        )
        sk_txt = "；".join(
            [
                f"{x.get('role') or x.get('role_id')}：{x.get('reason') or '跳过'}"
                for x in sk
                if isinstance(x, dict)
            ]
        )
        summary = {
            "applied_n": len(ap) if isinstance(ap, list) else 0,
            "skipped_n": len(sk) if isinstance(sk, list) else 0,
            "applied": ap_txt[:1200],
            "skipped": sk_txt[:1200],
        }
        rep_s = json.dumps(summary, ensure_ascii=False)
        rep_b64 = base64.b64encode(rep_s.encode("utf-8")).decode("ascii")
        resp.headers["X-Sign-Apply-Summary-B64"] = rep_b64
    except Exception:
        pass
    return resp


def _mysql_api_error_message(exc: Exception) -> str:
    msg = str(exc or "").strip()
    low = msg.lower()
    if "timed out" in low or "timeout" in low or "10060" in msg or "10061" in msg:
        host = (os.environ.get("MYSQL_HOST") or "").strip() or "（未配置 MYSQL_HOST）"
        port = (os.environ.get("MYSQL_PORT") or "3306").strip()
        return (
            f"无法连接 MySQL（{host}:{port}）：连接超时或被拒绝。"
            "请确认数据库服务已启动、本机网络/VPN 可达，且 .env 中 MYSQL_HOST/PORT 正确。"
        )
    if msg:
        return f"MySQL 操作失败：{msg}"
    return "MySQL 操作失败"


@app.route("/api/sign/signers", methods=["GET"])
def api_sign_signers_list():
    """签署人库列表（MySQL 多机共享；否则存于当前会话目录）。"""
    brief = (request.args.get("brief") or "").strip().lower() in ("1", "true", "yes")
    compact = (request.args.get("compact") or "").strip().lower() in (
        "1",
        "true",
        "yes",
    )
    try:
        if _sign_using_mysql():
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            if brief:
                rows = mysql_store.list_signers_brief()
            elif compact:
                rows = mysql_store.list_signers_compact()
            else:
                rows = mysql_store.list_signers()
            return jsonify(
                {
                    "ok": True,
                    "db_share": True,
                    "brief": brief,
                    "compact": compact,
                    "signers": rows,
                }
            )
        _sign_ensure_session_inbox()
        sid = session["sign_inbox_sid"]
        from sign_handlers.sign_library_local import list_signers as local_list_signers

        return jsonify(
            {"ok": True, "db_share": False, "signers": local_list_signers(SIGN_INBOX_ROOT, sid)}
        )
    except Exception as e:
        err = _mysql_api_error_message(e) if _sign_using_mysql() else str(e)
        return jsonify({"ok": False, "error": err}), 500


@app.route("/api/sign/signers/stroke-options", methods=["GET"])
def api_sign_signer_stroke_options():
    """签署人签名/日期笔迹 id 列表（供工作台素材下拉，与 brief 签署人列表配合）。"""
    try:
        if _sign_using_mysql():
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            items = mysql_store.list_signer_stroke_options()
            return jsonify({"ok": True, "db_share": True, "items": items})
        _sign_ensure_session_inbox()
        sid = session["sign_inbox_sid"]
        from sign_handlers.sign_library_local import list_signer_stroke_options as local_opts

        return jsonify(
            {"ok": True, "db_share": False, "items": local_opts(SIGN_INBOX_ROOT, sid)}
        )
    except Exception as e:
        err = _mysql_api_error_message(e) if _sign_using_mysql() else str(e)
        return jsonify({"ok": False, "error": err}), 500


def _parse_signer_names_from_json(data: dict) -> list[str]:
    """支持 names 数组，或 name 中用中英文逗号/分号/换行分隔的多个姓名。"""
    import re as _re

    out: list[str] = []
    seen: set[str] = set()
    names_field = data.get("names")
    if isinstance(names_field, list):
        for x in names_field:
            n = str(x).strip()[:128]
            if not n:
                continue
            key = n.casefold()
            if key not in seen:
                seen.add(key)
                out.append(n)
    raw = (data.get("name") or "").strip()
    if raw:
        for part in _re.split(r"[,，;；\r\n]+", raw):
            n = part.strip()[:128]
            if not n:
                continue
            key = n.casefold()
            if key not in seen:
                seen.add(key)
                out.append(n)
    return out


@app.route("/api/sign/signers", methods=["POST"])
def api_sign_signers_create():
    data = request.get_json(silent=True) or {}
    names = _parse_signer_names_from_json(data)
    if not names:
        return jsonify({"ok": False, "error": "请填写至少一个签署人名称（可用逗号分隔多个）"}), 400
    if len(names) > 50:
        return jsonify({"ok": False, "error": "一次最多添加 50 个签署人"}), 400
    try:
        if _sign_using_mysql():
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            created = []
            for name in names:
                nid = mysql_store.insert_signer(name)
                created.append({"id": nid, "name": name})
            return jsonify(
                {
                    "ok": True,
                    "added": len(created),
                    "created": created,
                    "signers": mysql_store.list_signers(),
                }
            )
        _sign_ensure_session_inbox()
        sid = session["sign_inbox_sid"]
        from sign_handlers.sign_library_local import insert_signer as local_insert_signer
        from sign_handlers.sign_library_local import list_signers as local_list_signers

        created = []
        for name in names:
            nid = local_insert_signer(SIGN_INBOX_ROOT, sid, name)
            created.append({"id": nid, "name": name})
        return jsonify(
            {
                "ok": True,
                "added": len(created),
                "created": created,
                "signers": local_list_signers(SIGN_INBOX_ROOT, sid),
            }
        )
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/signers/<signer_id>", methods=["DELETE"])
def api_sign_signers_delete(signer_id):
    if not _SIGN_FILE_ID_RE.match(signer_id or ""):
        return jsonify({"ok": False, "error": "无效的签署人 id"}), 400
    try:
        if _sign_using_mysql():
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            n = mysql_store.delete_signer(signer_id)
            if not n:
                return jsonify({"ok": False, "error": "未找到该签署人"}), 404
            return jsonify({"ok": True, "signers": mysql_store.list_signers()})
        _sign_ensure_session_inbox()
        sid = session["sign_inbox_sid"]
        from sign_handlers.sign_library_local import delete_signer as local_delete_signer
        from sign_handlers.sign_library_local import list_signers as local_list_signers

        n = local_delete_signer(SIGN_INBOX_ROOT, sid, signer_id)
        if not n:
            return jsonify({"ok": False, "error": "未找到该签署人"}), 404
        return jsonify({"ok": True, "signers": local_list_signers(SIGN_INBOX_ROOT, sid)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/signers/<signer_id>/strokes", methods=["PUT"])
def api_sign_signers_strokes_put(signer_id):
    if not _SIGN_FILE_ID_RE.match(signer_id or ""):
        return jsonify({"ok": False, "error": "无效的签署人 id"}), 400
    sig_raw = request.form.get("sig") or ""
    date_raw = request.form.get("date") or ""
    locale = (request.form.get("locale") or "zh").strip().lower()
    sig_b = _decode_png_data_url_or_b64(sig_raw) if str(sig_raw).strip() else None
    date_b = _decode_png_data_url_or_b64(date_raw) if str(date_raw).strip() else None
    if sig_b is None and date_b is None:
        return jsonify({"ok": False, "error": "请至少提交签名或日期笔迹之一"}), 400
    try:
        from sign_handlers import mysql_store as _ms_tmp

        is_temp = _ms_tmp._parse_is_temporary(
            request.form.get("is_temporary") or request.form.get("temporary")
        )
    except Exception:
        _tv = str(request.form.get("is_temporary") or request.form.get("temporary") or "").strip().lower()
        is_temp = _tv in ("1", "true", "yes", "on")
    try:
        if _sign_using_mysql():
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            sig_res = None
            date_res = None
            if sig_b is not None:
                sig_res = mysql_store.upsert_signer_stroke_item(
                    signer_id, "sig", sig_b, locale=locale, is_temporary=is_temp
                )
            if date_b is not None:
                date_res = mysql_store.upsert_signer_stroke_item(
                    signer_id, "date", date_b, locale=locale, is_temporary=is_temp
                )
            # 兼容旧返回字段：stroke_set_id 仍返回（若两者都提交则也写入旧 set）
            stroke_set_res = None
            if sig_b is not None and date_b is not None:
                try:
                    stroke_set_res = mysql_store.upsert_signer_strokes(signer_id, sig_b, date_b, locale=locale)
                except Exception:
                    stroke_set_res = None
            signer_nm = mysql_store.get_signer_display_name(signer_id) or ""
            return jsonify(
                {
                    "ok": True,
                    "signer_id": signer_id,
                    "signer_name": signer_nm,
                    "sig_item_id": (sig_res or {}).get("stroke_item_id"),
                    "date_item_id": (date_res or {}).get("stroke_item_id"),
                    "stroke_set_id": (stroke_set_res or {}).get("stroke_set_id"),
                    "overwritten": bool((sig_res or {}).get("overwritten") or (date_res or {}).get("overwritten") or (stroke_set_res or {}).get("overwritten")),
                }
            )
        _sign_ensure_session_inbox()
        sid = session["sign_inbox_sid"]
        from sign_handlers.sign_library_local import upsert_strokes as local_upsert

        # 会话模式：同样拆分保存（仍兼容 stroke_set）
        from sign_handlers.sign_library_local import upsert_stroke_item as local_upsert_item

        sig_res = None
        date_res = None
        if sig_b is not None:
            sig_res = local_upsert_item(
                SIGN_INBOX_ROOT, sid, signer_id, "sig", sig_b, locale=locale, is_temporary=is_temp
            )
        if date_b is not None:
            date_res = local_upsert_item(
                SIGN_INBOX_ROOT, sid, signer_id, "date", date_b, locale=locale, is_temporary=is_temp
            )
        stroke_set_res = None
        if sig_b is not None and date_b is not None:
            try:
                stroke_set_res = local_upsert(SIGN_INBOX_ROOT, sid, signer_id, sig_b, date_b, locale=locale)
            except Exception:
                stroke_set_res = None
        signer_nm = ""
        try:
            from sign_handlers.sign_library_local import list_signers as local_list_signers

            for row in local_list_signers(SIGN_INBOX_ROOT, sid) or []:
                if isinstance(row, dict) and row.get("id") == signer_id:
                    signer_nm = (row.get("name") or "").strip()
                    break
        except Exception:
            signer_nm = ""
        return jsonify(
            {
                "ok": True,
                "signer_id": signer_id,
                "signer_name": signer_nm,
                "sig_item_id": (sig_res or {}).get("stroke_item_id"),
                "date_item_id": (date_res or {}).get("stroke_item_id"),
                "stroke_set_id": (stroke_set_res or {}).get("stroke_set_id"),
                "overwritten": bool((sig_res or {}).get("overwritten") or (date_res or {}).get("overwritten") or (stroke_set_res or {}).get("overwritten")),
            }
        )
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/signers/<signer_id>/stroke-piece", methods=["PUT"])
@sign_heavy_route
def api_sign_signer_stroke_piece_put(signer_id):
    """录入英文点分日期笔迹元件：数字 0-9、月份 pm01..pm12、连接符 pdot（locale 固定 en）。"""
    if not _SIGN_FILE_ID_RE.match(signer_id or ""):
        return jsonify({"ok": False, "error": "无效的签署人 id"}), 400
    if not _sign_using_mysql():
        return jsonify({"ok": False, "error": "笔迹元件需启用 MySQL（MYSQL_HOST）"}), 400
    piece_kind = (request.form.get("piece_kind") or request.form.get("kind") or "").strip()
    png_raw = request.form.get("png") or ""
    png_b = _decode_png_data_url_or_b64(png_raw) if str(png_raw).strip() else None
    if not png_b:
        return jsonify({"ok": False, "error": "请提交 png 笔迹图"}), 400
    try:
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        is_temp = mysql_store._parse_is_temporary(
            request.form.get("is_temporary") or request.form.get("temporary")
        )
        res = mysql_store.upsert_signer_stroke_piece(
            signer_id, piece_kind, png_b, is_temporary=is_temp
        )
        return jsonify({"ok": True, **res})
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/signers/<signer_id>/stroke-pieces", methods=["PUT"])
@sign_heavy_route
def api_sign_signer_stroke_pieces_batch_put(signer_id):
    """批量录入笔迹元件：JSON body { items: [ { piece_kind, png }, ... ] }。"""
    if not _SIGN_FILE_ID_RE.match(signer_id or ""):
        return jsonify({"ok": False, "error": "无效的签署人 id"}), 400
    if not _sign_using_mysql():
        return jsonify({"ok": False, "error": "笔迹元件需启用 MySQL（MYSQL_HOST）"}), 400
    data = request.get_json(silent=True) or {}
    items = data.get("items")
    if not isinstance(items, list) or not items:
        return jsonify({"ok": False, "error": "请求体需包含非空 items 数组"}), 400
    if len(items) > 80:
        return jsonify({"ok": False, "error": "单次最多提交 80 条元件"}), 400
    try:
        from sign_handlers import mysql_store
        from sign_handlers.date_piece_compose import piece_kind_label

        mysql_store.ensure_sign_mysql()
        overwrite = True
        try:
            overwrite = bool(data.get("overwrite", True))
        except Exception:
            overwrite = True
        is_temp = mysql_store._parse_is_temporary(data.get("is_temporary") or data.get("temporary"))
        results: list = []
        seen_kind = set()
        to_write: list = []
        for it in items:
            if not isinstance(it, dict):
                results.append({"piece_kind": "", "ok": False, "error": "条目须为 JSON 对象"})
                continue
            pk = (it.get("piece_kind") or it.get("kind") or "").strip()
            pk_norm = pk.lower()
            if pk_norm in seen_kind:
                results.append(
                    {
                        "piece_kind": pk,
                        "ok": False,
                        "error_code": "duplicate_in_request",
                        "error": "本次队列中该元件类别重复，请保留一条后重试",
                    }
                )
                continue
            seen_kind.add(pk_norm)
            png_raw = it.get("png") or ""
            png_b = _decode_png_data_url_or_b64(png_raw) if str(png_raw).strip() else None
            if not png_b:
                results.append({"piece_kind": pk, "ok": False, "error": "缺少 png"})
                continue
            to_write.append((pk, png_b))
        if to_write:
            batch_results = mysql_store.batch_upsert_signer_stroke_pieces(
                signer_id, to_write, overwrite=overwrite, is_temporary=is_temp
            )
            for br in batch_results:
                pk = (br.get("piece_kind") or "").strip()
                _pk_l = pk.lower()
                _kind_h = piece_kind_label(_pk_l) if _pk_l else "元件"
                if br.get("error_code") == "exists":
                    results.append(
                        {
                            "piece_kind": pk,
                            "ok": False,
                            "error_code": "exists",
                            "error": f"该签署人的该元件已存在：{_kind_h}（{_pk_l}），请确认是否覆盖",
                        }
                    )
                else:
                    results.append(br)
        return jsonify({"ok": True, "results": results})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/signers/<signer_id>/composite-date-preview", methods=["GET"])
def api_sign_composite_date_preview(signer_id):
    """按签署人笔迹元件预览拼接 PNG。iso=YYYY-MM-DD；layout=zh_ymd|en_space|en_dot（默认 en_dot）。"""
    if not _SIGN_FILE_ID_RE.match(signer_id or ""):
        return jsonify({"ok": False, "error": "无效的签署人 id"}), 400
    if not _sign_using_mysql():
        return jsonify({"ok": False, "error": "需启用 MySQL"}), 400
    iso = (request.args.get("iso") or "").strip()
    layout = (request.args.get("layout") or "en_dot").strip().lower()
    if layout not in ("zh_ymd", "en_space", "en_dot"):
        return jsonify({"ok": False, "error": "layout 须为 zh_ymd / en_space / en_dot"}), 400
    try:
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        png_b, label = mysql_store.compose_date_piece_png(signer_id, iso, layout)
        resp = send_file(io.BytesIO(png_b), mimetype="image/png")
        resp.headers["X-Composite-Date-Label"] = label
        return resp
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 400


@app.route("/api/sign/signers/<signer_id>/stroke/<kind>", methods=["GET"])
def api_sign_signer_stroke_get(signer_id, kind):
    if not _SIGN_FILE_ID_RE.match(signer_id or ""):
        return jsonify({"ok": False, "error": "无效的签署人 id"}), 400
    if kind not in ("sig", "date"):
        return jsonify({"ok": False, "error": "kind 须为 sig 或 date"}), 400
    try:
        blob = None
        loc_q = (request.args.get("locale") or "").strip().lower()
        if loc_q not in ("zh", "en"):
            loc_q = ""
        if _sign_using_mysql():
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            blob = mysql_store.get_signer_stroke_png_resolved(
                signer_id, kind, loc_q or None
            )
        else:
            _sign_ensure_session_inbox()
            sid = session["sign_inbox_sid"]
            from sign_handlers.sign_library_local import get_signer_stroke_bytes_for_kind as local_stroke_kind

            loc_use = loc_q if loc_q in ("zh", "en") else "zh"
            blob = local_stroke_kind(SIGN_INBOX_ROOT, sid, signer_id, kind, loc_use)
        if not blob:
            return Response(status=404)
        return Response(blob, mimetype="image/png")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/stroke-sets/<set_id>/stroke/<kind>", methods=["GET"])
def api_sign_stroke_set_stroke_get(set_id, kind):
    if not _SIGN_FILE_ID_RE.match(set_id or ""):
        return jsonify({"ok": False, "error": "无效的笔迹套 id"}), 400
    if kind not in ("sig", "date"):
        return jsonify({"ok": False, "error": "kind 须为 sig 或 date"}), 400
    try:
        blob = None
        loc_q = (request.args.get("locale") or "").strip().lower()
        if loc_q not in ("zh", "en"):
            loc_q = ""
        if _sign_using_mysql():
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            blob = mysql_store.get_stroke_set_stroke_png_resolved(
                set_id, kind, loc_q or None
            )
        else:
            _sign_ensure_session_inbox()
            sid = session["sign_inbox_sid"]
            from sign_handlers.sign_library_local import get_stroke_set_stroke_bytes_for_kind as local_set_kind

            loc_use = loc_q if loc_q in ("zh", "en") else "zh"
            blob = local_set_kind(SIGN_INBOX_ROOT, sid, set_id, kind, loc_use)
        if not blob:
            return Response(status=404)
        return Response(blob, mimetype="image/png")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/stroke-items/<item_id>/png", methods=["GET"])
def api_sign_stroke_item_get(item_id):
    if not _SIGN_FILE_ID_RE.match(item_id or ""):
        return jsonify({"ok": False, "error": "无效的笔迹素材 id"}), 400
    try:
        if _sign_using_mysql():
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            row = mysql_store.get_stroke_item_row(item_id)
            blob = (row or {}).get("png")
        else:
            _sign_ensure_session_inbox()
            sid = session["sign_inbox_sid"]
            from sign_handlers.sign_library_local import get_stroke_item_bytes as local_get_item

            blob = local_get_item(SIGN_INBOX_ROOT, sid, item_id)
        if not blob:
            return Response(status=404)
        return Response(blob, mimetype="image/png")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


def _load_role_material_export_context(
    file_id: str, map_override, doc_date_fallback: str = ""
):
    """加载文件字节、扩展名、role-map、detect 快照与 export 参数。"""
    file_bytes = None
    ext = ".docx"
    det = {}
    mapping: dict = {}
    export_kw: dict = {}
    base_name = "document"
    if _sign_using_mysql():
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        row = mysql_store.get_file_row(file_id)
        if not row or not row.get("file_data"):
            err = (row.get("file_load_error") or "").strip() if row else ""
            return None, (err or "未找到该文件", 404)
        file_bytes = row["file_data"]
        ext = (row.get("ext") or ".docx").lower()
        base_name = _safe_display_filename_keep_unicode(row.get("name") or "document") or "document"
        base_map = mysql_store.get_file_role_signer_map(file_id) or {}
        if isinstance(map_override, dict):
            from sign_handlers.material_png_export import _merge_role_material_maps

            mapping = _merge_role_material_maps(base_map, map_override)
        else:
            from sign_handlers.config import normalize_role_signer_map

            mapping = normalize_role_signer_map(base_map)
        det = mysql_store.get_file_detect_snapshot(file_id) or {}
        export_kw = {"using_mysql": True, "mysql_store": mysql_store}
    else:
        _sign_ensure_session_inbox()
        sid = session["sign_inbox_sid"]
        rec = _sign_find_record(file_id)
        if not rec:
            return None, ("未找到该文件", 404)
        ext = (rec.get("ext") or ".docx").lower()
        source_path = _sign_saved_disk_path(sid, file_id, ext)
        if not os.path.isfile(source_path):
            return None, ("文件内容缺失", 404)
        with open(source_path, "rb") as fp:
            file_bytes = fp.read()
        base_name = _safe_display_filename_keep_unicode(rec.get("name") or "document") or "document"
        from sign_handlers.sign_library_local import get_file_role_map as local_get_map
        from sign_handlers.sign_library_local import get_stroke_item_bytes as local_get_item_fn
        from sign_handlers.sign_library_local import get_file_detect_snapshot as local_get_det

        base_map = local_get_map(SIGN_INBOX_ROOT, sid, file_id) or {}
        if isinstance(map_override, dict):
            from sign_handlers.material_png_export import _merge_role_material_maps

            mapping = _merge_role_material_maps(base_map, map_override)
        else:
            from sign_handlers.config import normalize_role_signer_map

            mapping = normalize_role_signer_map(base_map)
        det = local_get_det(SIGN_INBOX_ROOT, sid, file_id) or {}
        export_kw = {
            "using_mysql": False,
            "mysql_store": None,
            "local_get_item": local_get_item_fn,
            "inbox_root": SIGN_INBOX_ROOT,
            "sid": sid,
        }
    if ext not in SIGN_ALLOWED_EXT:
        return None, ("不支持的文件类型", 400)
    if doc_date_fallback:
        patched = {}
        for rk, rv in (mapping or {}).items():
            if isinstance(rv, dict):
                p = dict(rv)
                if not (p.get("date_iso") or "").strip():
                    p["date_iso"] = doc_date_fallback[:10]
                patched[str(rk)] = p
            else:
                patched[str(rk)] = rv
        mapping = patched
    return (
        {
            "file_bytes": file_bytes,
            "ext": ext,
            "mapping": mapping,
            "det": det,
            "export_kw": export_kw,
            "base_name": base_name,
        },
        None,
    )


@app.route("/api/sign/files/<file_id>/role-materials-zip", methods=["GET", "POST"])
def api_sign_file_role_materials_zip(file_id):
    """一次打包下载当前行所有已有签字素材（编签/编日/审签/审日/批签/批日）。"""
    if not _SIGN_FILE_ID_RE.match(file_id or ""):
        return jsonify({"ok": False, "error": "无效的文件 id"}), 400
    map_override = None
    doc_date_fallback = ""
    if request.method == "GET":
        doc_date_fallback = str(request.args.get("doc_date") or "").strip()
    else:
        data = request.get_json(silent=True) or {}
        if "map" in data:
            map_override = data.get("map")
        doc_date_fallback = str(data.get("doc_date") or "").strip()
    try:
        ctx, err = _load_role_material_export_context(
            file_id, map_override, doc_date_fallback
        )
        if err:
            msg, code = err
            return jsonify({"ok": False, "error": msg}), code
        from sign_handlers import ROLE_ID_TO_KEYWORD
        from sign_handlers.material_png_export import export_role_materials_zip

        roles = list(ROLE_ID_TO_KEYWORD.keys())
        placement_plan = _build_placement_plan_from_detect(ctx["det"], roles)
        zip_b, zerr, count = export_role_materials_zip(
            file_bytes=ctx["file_bytes"],
            ext=ctx["ext"],
            role_map=ctx["mapping"],
            placement_plan=placement_plan,
            doc_date_fallback=doc_date_fallback,
            **ctx["export_kw"],
        )
        if zerr:
            return jsonify({"ok": False, "error": zerr}), 400
        if not zip_b:
            return jsonify({"ok": False, "error": "无法生成 ZIP"}), 400
        base = os.path.splitext(os.path.basename(ctx["base_name"] or "document"))[0]
        safe_base = re.sub(r'[\\/:*?"<>|]', "_", base) or "document"
        fname = f"{safe_base}_签字素材.zip"
        resp = send_file(io.BytesIO(zip_b), mimetype="application/zip", as_attachment=True, download_name=fname)
        resp.headers["X-Material-Count"] = str(count)
        return resp
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/files/<file_id>/role-material-png", methods=["GET", "POST"])
def api_sign_file_role_material_png(file_id):
    """导出与自动签字落位同缩放比例的 PNG，供工作台手动贴图。"""
    if not _SIGN_FILE_ID_RE.match(file_id or ""):
        return jsonify({"ok": False, "error": "无效的文件 id"}), 400
    map_override = None
    doc_date_fallback = ""
    role_id = ""
    kind = "sig"
    if request.method == "GET":
        role_id = str(request.args.get("role") or "").strip()
        kind = str(request.args.get("kind") or "sig").strip().lower()
        doc_date_fallback = str(request.args.get("doc_date") or "").strip()
    else:
        data = request.get_json(silent=True) or {}
        role_id = str(data.get("role") or "").strip()
        kind = str(data.get("kind") or "sig").strip().lower()
        if "map" in data:
            map_override = data.get("map")
        doc_date_fallback = str(data.get("doc_date") or "").strip()
    from sign_handlers import ROLE_ID_TO_KEYWORD

    if role_id not in ROLE_ID_TO_KEYWORD:
        return jsonify({"ok": False, "error": f"无效角色 id: {role_id}"}), 400
    if kind not in ("sig", "date"):
        return jsonify({"ok": False, "error": "kind 须为 sig 或 date"}), 400
    try:
        ctx, err = _load_role_material_export_context(
            file_id, map_override, doc_date_fallback
        )
        if err:
            msg, code = err
            return jsonify({"ok": False, "error": msg}), code
        pair = (ctx["mapping"] or {}).get(role_id) if isinstance(ctx["mapping"], dict) else {}
        if not isinstance(pair, dict):
            pair = {}
        from sign_handlers.material_png_export import _pair_with_doc_date

        pair = _pair_with_doc_date(pair, doc_date_fallback)
        placement_plan = _build_placement_plan_from_detect(ctx["det"], [role_id])
        from sign_handlers.material_png_export import export_role_material_png

        png_b, err_msg = export_role_material_png(
            file_bytes=ctx["file_bytes"],
            ext=ctx["ext"],
            role_id=role_id,
            kind=kind,
            pair=pair,
            placement_plan=placement_plan,
            **ctx["export_kw"],
        )
        if err_msg:
            return jsonify({"ok": False, "error": err_msg}), 400
        if not png_b:
            return jsonify({"ok": False, "error": "无法生成图片"}), 400
        resp = send_file(io.BytesIO(png_b), mimetype="image/png")
        resp.headers["X-Insert-Scale"] = "docx" if ctx["ext"] in (".docx", ".doc") else "xlsx"
        return resp
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/files/<file_id>/role-map", methods=["GET"])
def api_sign_file_role_map_get(file_id):
    if not _SIGN_FILE_ID_RE.match(file_id or ""):
        return jsonify({"ok": False, "error": "无效的文件 id"}), 400
    try:
        if _sign_using_mysql():
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            m = mysql_store.get_file_role_signer_map(file_id)
            return jsonify({"ok": True, "map": m})
        _sign_ensure_session_inbox()
        sid = session["sign_inbox_sid"]
        from sign_handlers.sign_library_local import get_file_role_map as local_get_map

        return jsonify({"ok": True, "map": local_get_map(SIGN_INBOX_ROOT, sid, file_id)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/files/<file_id>/role-map", methods=["PUT"])
def api_sign_file_role_map_put(file_id):
    if not _SIGN_FILE_ID_RE.match(file_id or ""):
        return jsonify({"ok": False, "error": "无效的文件 id"}), 400
    data = request.get_json(silent=True) or {}
    m = data.get("map")
    if not isinstance(m, dict):
        return jsonify({"ok": False, "error": "请求体需包含 map 对象"}), 400
    clean = {}
    for k, v in m.items():
        if not k or not v:
            continue
        if isinstance(v, dict):
            clean[str(k)] = {
                "sig": v.get("sig"),
                "date": v.get("date"),
                "date_mode": v.get("date_mode"),
                "date_iso": v.get("date_iso"),
            }
        else:
            clean[str(k)] = str(v)
    try:
        if _sign_using_mysql():
            from sign_handlers import mysql_store

            mysql_store.ensure_sign_mysql()
            mysql_store.set_file_role_signer_map(file_id, clean)
            return jsonify({"ok": True, "map": mysql_store.get_file_role_signer_map(file_id)})
        _sign_ensure_session_inbox()
        sid = session["sign_inbox_sid"]
        from sign_handlers.sign_library_local import set_file_role_map as local_set_map
        from sign_handlers.sign_library_local import get_file_role_map as local_get_map

        local_set_map(SIGN_INBOX_ROOT, sid, file_id, clean)
        return jsonify({"ok": True, "map": local_get_map(SIGN_INBOX_ROOT, sid, file_id)})
    except ValueError as e:
        # set_file_role_signer_map：拼接日期缺少签名/日期、或日期格式非法
        return jsonify({"ok": False, "error": str(e)}), 400
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/batch", methods=["POST"])
@sign_heavy_route
def api_sign_batch():
    """
    按每个文件已保存的 role-map，从签署人库取笔迹批量生成已签名文档。
    仅写入 sign_signed_output（MySQL）；不逐个返回文件流。
    """
    if not _sign_using_mysql():
        return jsonify({"ok": False, "error": "批量签名需启用 MySQL（MYSQL_HOST）"}), 400
    data = request.get_json(silent=True) or {}
    source = (data.get("source") or "library").strip().lower()
    apply_person = bool(data.get("apply_person", True))
    apply_date = bool(data.get("apply_date", True))
    if (not apply_person) and (not apply_date):
        return jsonify({"ok": False, "error": "apply_person 与 apply_date 不能同时为 false"}), 400
    if source not in ("library", "canvas"):
        source = "library"
    file_ids = data.get("file_ids")
    if not isinstance(file_ids, list) or not file_ids:
        return jsonify({"ok": False, "error": "请提供 file_ids 数组"}), 400
    for fid in file_ids:
        if not _SIGN_FILE_ID_RE.match(str(fid or "")):
            return jsonify({"ok": False, "error": f"无效的文件 id: {fid}"}), 400

    from sign_handlers import ROLE_ID_TO_KEYWORD
    from sign_handlers import mysql_store

    mysql_store.ensure_sign_mysql()
    _max_s = _sign_mysql_max_signed()
    cur_n = mysql_store.count_signed_outputs()
    if cur_n + len(file_ids) > _max_s:
        return jsonify(
            {
                "ok": False,
                "error": (
                    f"已签名记录空间不足（当前 {cur_n}，本次 {len(file_ids)}，"
                    f"上限 {_max_s}），请先删除旧记录"
                ),
            }
        ), 400

    roles_req = data.get("roles")
    if roles_req is not None and (not isinstance(roles_req, list) or not roles_req):
        return jsonify({"ok": False, "error": "roles 必须为非空数组"}), 400
    if isinstance(roles_req, list):
        roles_req = [str(x) for x in roles_req if str(x or "").strip()]
        roles_req = list(dict.fromkeys(roles_req))
        for rr in roles_req:
            if rr not in ROLE_ID_TO_KEYWORD:
                return jsonify({"ok": False, "error": f"无效角色 id: {rr}"}), 400

    canvas_sig_map_raw = data.get("sig_map") if isinstance(data, dict) else None
    canvas_date_map_raw = data.get("date_map") if isinstance(data, dict) else None
    canvas_sig_map: dict = {}
    canvas_date_map: dict = {}
    lib_canvas_sig: dict = {}
    lib_canvas_date: dict = {}
    if source == "library":
        if isinstance(canvas_sig_map_raw, dict):
            for rid, v in canvas_sig_map_raw.items():
                rid_s = str(rid)
                if rid_s in ROLE_ID_TO_KEYWORD:
                    b = _decode_png_data_url_or_b64(v or "")
                    if b:
                        lib_canvas_sig[rid_s] = b
        if isinstance(canvas_date_map_raw, dict):
            for rid, v in canvas_date_map_raw.items():
                rid_s = str(rid)
                if rid_s in ROLE_ID_TO_KEYWORD:
                    b = _decode_png_data_url_or_b64(v or "")
                    if b:
                        lib_canvas_date[rid_s] = b
    if source == "canvas":
        if not isinstance(canvas_sig_map_raw, dict) or not isinstance(canvas_date_map_raw, dict):
            return jsonify({"ok": False, "error": "画布模式需提供 sig_map/date_map 对象"}), 400
        for rid, v in canvas_sig_map_raw.items():
            rid_s = str(rid)
            if rid_s in ROLE_ID_TO_KEYWORD:
                b = _decode_png_data_url_or_b64(v or "")
                if b:
                    canvas_sig_map[rid_s] = b
        for rid, v in canvas_date_map_raw.items():
            rid_s = str(rid)
            if rid_s in ROLE_ID_TO_KEYWORD:
                b = _decode_png_data_url_or_b64(v or "")
                if b:
                    canvas_date_map[rid_s] = b
        if roles_req:
            for rr in roles_req:
                if rr not in canvas_sig_map or rr not in canvas_date_map:
                    return jsonify({"ok": False, "error": f"角色 {rr} 缺少画布签名或日期"}), 400

    results = []
    # 允许前端分文件多次提交同一批（用于进度展示）；须为 32 位 hex，否则每次请求自动生成
    _bid = (data.get("batch_id") or "").strip().lower()
    batch_id = _bid if _bid and _SIGN_FILE_ID_RE.match(_bid) else uuid.uuid4().hex
    for fid in file_ids:
        try:
            row = mysql_store.get_file_row(fid)
            if not row or not row.get("file_data"):
                fe = (row.get("ftp_last_error") if isinstance(row, dict) else None) if row else None
                if fe:
                    results.append({"file_id": fid, "ok": False, "error": f"无法读取文件内容：{fe}"})
                else:
                    results.append({"file_id": fid, "ok": False, "error": "无法读取文件内容（可能 FTP 暂不可用）"})
                continue
            ext = (row.get("ext") or ".docx").lower()
            if ext not in SIGN_ALLOWED_EXT:
                results.append({"file_id": fid, "ok": False, "error": "不支持的扩展名"})
                continue
            base_name = _safe_display_filename_keep_unicode(row.get("name") or "document") or "document"
            if not base_name.lower().endswith(ext):
                base_name = os.path.splitext(base_name)[0] + ext
            from sign_handlers.config import normalize_role_signer_map

            mapping = normalize_role_signer_map(
                mysql_store.get_file_role_signer_map(fid) or {}
            )
            if roles_req:
                roles = _normalize_sign_role_ids(roles_req)
            else:
                roles = _resolve_actual_roles_for_file(
                    fid, [r for r in mapping.keys() if r in ROLE_ID_TO_KEYWORD]
                )
            if not roles:
                results.append({"file_id": fid, "ok": False, "error": "未配置角色-签署人映射"})
                continue
            sig_map = {}
            date_map = {}
            applied = []
            skipped = []
            if source == "canvas":
                for rid in roles:
                    sb = canvas_sig_map.get(rid)
                    db = canvas_date_map.get(rid)
                    if sb and apply_person:
                        sig_map[rid] = sb
                    if db and apply_date:
                        date_map[rid] = db
                    if (sb and apply_person) or (db and apply_date):
                        applied.append(
                            {
                                "role_id": rid,
                                "sig": bool(sb and apply_person),
                                "date": bool(db and apply_date),
                                "date_mode": "canvas",
                            }
                        )
                    else:
                        skipped.append(
                            {
                                "role_id": rid,
                                "reason": "该角色签名与日期画布均为空，已跳过",
                            }
                        )
            else:
                for rid in roles:
                    pair = mapping.get(rid) or {}
                    sig_id = (pair.get("sig") if isinstance(pair, dict) else None) or None
                    date_id = (pair.get("date") if isinstance(pair, dict) else None) or None
                    dm = (pair.get("date_mode") if isinstance(pair, dict) else None) or None
                    diso = (pair.get("date_iso") if isinstance(pair, dict) else None) or None
                    if mysql_store.is_composite_date_mode(dm):
                        if not apply_date:
                            dm = None
                        # 与单文件 POST /api/sign 一致：拼接失败不整角色放弃，退回「素材签名/整张日期图」能签则签
                        if (not sig_id) or (not diso) or (not apply_person):
                            skipped.append(
                                {
                                    "role_id": rid,
                                    "reason": "拼接日期条件不齐（需签名素材+日历），已改试整张日期/仅签名",
                                }
                            )
                            dm = None
                        srow = None
                        if dm:
                            srow = mysql_store.get_stroke_item_row(sig_id)
                            if not srow or not srow.get("png"):
                                skipped.append(
                                    {
                                        "role_id": rid,
                                        "reason": "签名素材缺失，无法拼接；已改试整张日期/仅签名",
                                    }
                                )
                                dm = None
                        if dm:
                            sid0 = (srow.get("signer_id") or "").strip()
                            if not sid0:
                                skipped.append(
                                    {
                                        "role_id": rid,
                                        "reason": "无法解析签署人，跳过拼接；已改试整张日期/仅签名",
                                    }
                                )
                                dm = None
                        if dm:
                            try:
                                lay = mysql_store.composite_mode_to_layout(dm)
                                dbytes, _lbl = mysql_store.compose_date_piece_png(
                                    sid0, str(diso).strip(), lay
                                )
                            except Exception as e:
                                skipped.append(
                                    {
                                        "role_id": rid,
                                        "reason": f"日期拼接失败：{e}；已改试整张日期/仅签名",
                                    }
                                )
                                dm = None
                            else:
                                if apply_person:
                                    sig_map[rid] = srow["png"]
                                if apply_date:
                                    date_map[rid] = dbytes
                                applied.append(
                                    {
                                        "role_id": rid,
                                        "sig": bool(apply_person),
                                        "date": bool(apply_date),
                                        "date_mode": "composite",
                                    }
                                )
                                continue
                    sb = None
                    if sig_id:
                        srow = mysql_store.get_stroke_item_row(sig_id)
                        if srow and srow.get("png"):
                            sb = srow["png"]
                            if apply_person:
                                sig_map[rid] = sb
                    db = None
                    if date_id:
                        drow = mysql_store.get_stroke_item_row(date_id)
                        if drow and drow.get("png"):
                            db = drow["png"]
                            if apply_date:
                                date_map[rid] = db
                    ov_sb = lib_canvas_sig.get(rid)
                    ov_db = lib_canvas_date.get(rid)
                    if ov_sb and apply_person:
                        sig_map[rid] = ov_sb
                        sb = ov_sb
                    if ov_db and apply_date:
                        date_map[rid] = ov_db
                        db = ov_db
                    if (sb and apply_person) or (db and apply_date):
                        applied.append(
                            {
                                "role_id": rid,
                                "sig": bool(sb and apply_person),
                                "date": bool(db and apply_date),
                                "date_mode": "canvas_override" if (ov_sb or ov_db) else "item",
                            }
                        )
                    else:
                        skipped.append(
                            {
                                "role_id": rid,
                                "reason": "未绑定或素材缺失（签名/日期均无可用数据）",
                            }
                        )
            # 至少要有一个角色提供签名或日期；否则该文件不生成
            if (not sig_map) and (not date_map):
                results.append(
                    {
                        "file_id": fid,
                        "ok": False,
                        "error": "未选择任何可用素材：请至少为一个角色提供签名或日期后再生成",
                        "applied_n": 0,
                        "skipped_n": len(skipped),
                        "skipped": "；".join(
                            [
                                f"{x.get('role_id')}：{x.get('reason') or '跳过'}"
                                for x in skipped
                                if isinstance(x, dict)
                            ]
                        )[:1200],
                    }
                )
                continue
            det_snapshot = mysql_store.get_file_detect_snapshot(fid) or {}
            placement_plan = _build_placement_plan_from_detect(det_snapshot, roles)
            res = _sign_process_document_bytes(
                row["file_data"],
                ext,
                base_name,
                roles,
                sig_map,
                date_map,
                fid,
                batch_id=batch_id,
                placement_plan=placement_plan,
                require_role_placement=True,
            )
            if not res.get("ok"):
                results.append(
                    {
                        "file_id": fid,
                        "ok": False,
                        "error": res.get("error", "失败"),
                        "applied_n": len(applied),
                        "skipped_n": len(skipped),
                        "missing_roles": res.get("missing_roles") or [],
                        "tail_appended_roles": res.get("tail_appended_roles") or [],
                        "fallback_roles": res.get("fallback_roles") or [],
                        "per_role_results": res.get("per_role_results") or {},
                    }
                )
            else:
                ap_txt = "；".join(
                    [
                        f"{x.get('role_id')}："
                        + ("签名" if x.get("sig") else "")
                        + ("+" if x.get("sig") and x.get("date") else "")
                        + ("日期" if x.get("date") else "")
                        + ("（拼接）" if x.get("date_mode") == "composite" else "")
                        for x in applied
                        if isinstance(x, dict)
                    ]
                )
                sk_txt = "；".join(
                    [
                        f"{x.get('role_id')}：{x.get('reason') or '跳过'}"
                        for x in skipped
                        if isinstance(x, dict)
                    ]
                )
                results.append(
                    {
                        "file_id": fid,
                        "ok": True,
                        "partial": bool(res.get("partial")),
                        "signed_id": res.get("signed_id"),
                        "name": res.get("dl_name"),
                        "applied_n": len(applied),
                        "skipped_n": len(skipped),
                        "applied": ap_txt[:1200],
                        "skipped": sk_txt[:1200],
                        "missing_roles": res.get("missing_roles") or [],
                        "tail_appended_roles": res.get("tail_appended_roles") or [],
                        "placed_roles": res.get("placed_roles") or [],
                        "fallback_roles": res.get("fallback_roles") or [],
                        "per_role_results": res.get("per_role_results") or {},
                    }
                )
        except Exception as e:
            results.append({"file_id": fid, "ok": False, "error": str(e) or type(e).__name__})
    try:
        ok_n = sum(1 for it in results if isinstance(it, dict) and it.get("ok"))
        fail_n = max(0, len(results) - ok_n)
        summary_items = []
        for it in results:
            if not isinstance(it, dict):
                continue
            summary_items.append(
                {
                    "file_id": str(it.get("file_id") or ""),
                    "name": str(it.get("name") or ""),
                    "ok": bool(it.get("ok")),
                    "partial": bool(it.get("partial")),
                    "error": str(it.get("error") or ""),
                    "applied_n": int(it.get("applied_n") or 0),
                    "skipped_n": int(it.get("skipped_n") or 0),
                    "applied": str(it.get("applied") or ""),
                    "skipped": str(it.get("skipped") or ""),
                    "missing_roles": list(it.get("missing_roles") or []),
                    "tail_appended_roles": list(it.get("tail_appended_roles") or []),
                    "fallback_roles": list(it.get("fallback_roles") or []),
                    "per_role_results": dict(it.get("per_role_results") or {}),
                    "signed_id": str(it.get("signed_id") or ""),
                }
            )
        mysql_store.upsert_sign_batch_result(
            batch_id,
            {
                "batch_id": batch_id,
                "created_at": datetime.now().isoformat(sep=" ", timespec="seconds"),
                "total": len(results),
                "ok": ok_n,
                "failed": fail_n,
                "items": summary_items,
            },
        )
    except Exception:
        pass
    return jsonify({"ok": True, "results": results, "batch_id": batch_id})


@app.route("/api/sign/signed", methods=["GET"])
def api_sign_signed_list():
    """List signed outputs from MySQL shared storage（分页、按文件名搜索）。"""
    if not _sign_using_mysql():
        return jsonify(
            {
                "ok": True,
                "items": [],
                "total": 0,
                "page": 1,
                "page_size": 10,
                "db_share": False,
            }
        )
    try:
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        q = (request.args.get("q") or "").strip()
        page = max(1, int(request.args.get("page") or 1))
        page_size = max(1, min(100, int(request.args.get("page_size") or 10)))
        items, total = mysql_store.list_signed_outputs_page(
            q=q, page=page, page_size=page_size
        )
        return jsonify(
            {
                "ok": True,
                "items": items,
                "total": total,
                "page": page,
                "page_size": page_size,
                "db_share": True,
            }
        )
    except Exception as e:
        return jsonify({"ok": False, "error": f"MySQL ????: {e}"}), 500


@app.route("/api/sign/signed/<signed_id>", methods=["GET"])
def api_sign_signed_download(signed_id):
    if not _sign_using_mysql():
        return jsonify({"ok": False, "error": "????MySQL ????"}), 404
    if not _SIGN_FILE_ID_RE.match(signed_id or ""):
        return jsonify({"ok": False, "error": "????id"}), 400
    try:
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        row = mysql_store.get_signed_row(signed_id)
        if not row or not row.get("file_data"):
            err = (row or {}).get("file_load_error") or "记录不存在或文件内容不可用"
            return jsonify({"ok": False, "error": err}), 404
        ext = (row.get("ext") or ".docx").lower()
        name = row.get("name") or ("signed" + ext)
        return send_file(
            io.BytesIO(row["file_data"]),
            as_attachment=True,
            download_name=name,
            mimetype=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                if ext == ".docx"
                else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ),
        )
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/signed/<signed_id>", methods=["DELETE"])
def api_sign_signed_delete(signed_id):
    if not _sign_using_mysql():
        return jsonify({"ok": False, "error": "????MySQL ????"}), 400
    if not _SIGN_FILE_ID_RE.match(signed_id or ""):
        return jsonify({"ok": False, "error": "????id"}), 400
    try:
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        n = mysql_store.delete_signed_output(signed_id)
        if not n:
            return jsonify({"ok": False, "error": "??????"}), 404
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/signed-batches", methods=["GET"])
def api_sign_signed_batches():
    """按批次列出已签名输出（分页、按批次/文件名搜索）。"""
    if not _sign_using_mysql():
        return jsonify(
            {
                "ok": True,
                "batches": [],
                "total": 0,
                "page": 1,
                "page_size": 10,
                "db_share": False,
            }
        )
    try:
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        q = (request.args.get("q") or "").strip()
        page = max(1, int(request.args.get("page") or 1))
        page_size = max(1, min(100, int(request.args.get("page_size") or 10)))
        batches, total, legacy_total = mysql_store.list_signed_batches_page(
            q=q, page=page, page_size=page_size
        )
        for b in batches:
            try:
                bid = str((b or {}).get("batch_id") or "").strip()
                if not bid:
                    continue
                br = mysql_store.get_sign_batch_result(bid) or {}
                if br:
                    b["result_total"] = int(br.get("total") or 0)
                    b["result_ok"] = int(br.get("ok") or 0)
                    b["result_failed"] = int(br.get("failed") or 0)
            except Exception:
                continue
        return jsonify(
            {
                "ok": True,
                "batches": batches,
                "legacy_total": legacy_total,
                "total": total,
                "page": page,
                "page_size": page_size,
                "db_share": True,
            }
        )
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


def _format_signed_roles_json(roles_json: str) -> str:
    if not roles_json:
        return ""
    try:
        ids = json.loads(roles_json)
        if isinstance(ids, list):
            return "、".join(str(x) for x in ids if x)
    except Exception:
        pass
    return str(roles_json).strip()


def _build_signed_zip_summary_text(
    items,
    pack_lines: list,
    *,
    batch_id: Optional[str] = None,
    legacy: bool = False,
    batch_result: Optional[dict] = None,
) -> str:
    """生成批次/历史 zip 内的「签字结果说明.txt」正文。"""
    lines = ["批量签字结果说明", ""]
    if batch_id:
        lines.append(f"批次 ID：{batch_id}")
    elif legacy:
        lines.append("范围：历史记录（无批次号）")
    lines.append(f"说明生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")
    total = len(items or [])
    packed_n = sum(1 for p in pack_lines if p.get("packed"))
    partial_n = sum(
        1
        for p in pack_lines
        if p.get("packed") and "_部分成功" in str(p.get("zip_name") or p.get("output_name") or "")
    )
    lines.append(f"共 {total} 条记录，ZIP 内已包含成品 {packed_n} 个")
    if partial_n:
        lines.append(
            f"其中「部分成功」{partial_n} 个：仅部分签字角色落位成功，文件名带 `_部分成功` 后缀，详见下方明细。"
        )
    if total > packed_n:
        lines.append(f"未能打包 {total - packed_n} 个（见下方明细）")
    if isinstance(batch_result, dict):
        br_total = int(batch_result.get("total") or 0)
        br_ok = int(batch_result.get("ok") or 0)
        br_fail = int(batch_result.get("failed") or 0)
        if br_total > 0:
            lines.append(f"本次批签执行记录：成功 {br_ok} / {br_total}，失败 {br_fail}")
    lines.append("")
    lines.append("=" * 48)
    for p in pack_lines:
        title = p.get("output_name") or p.get("name") or p.get("id") or "?"
        lines.append(f"【{title}】")
        if p.get("source_name"):
            lines.append(f"  源文件：{p['source_name']}")
        if p.get("created_at"):
            lines.append(f"  签字时间：{p['created_at']}")
        roles_l = p.get("roles_label")
        if roles_l:
            lines.append(f"  签字角色：{roles_l}")
        if p.get("packed"):
            lines.append(f"  打包：已包含 → {p.get('zip_name') or title}")
        else:
            lines.append(f"  打包：未包含 — {p.get('error') or '无法读取文件内容'}")
        lines.append("")
    lines.append(
        "【说明】上列为本批次/历史范围内各文件的打包情况；"
        "若某条「未包含」，多为 FTP 读取失败或历史 BLOB 已清空，"
        "可在服务端运行 FTP 校验补传后重试下载。"
    )
    if isinstance(batch_result, dict):
        failed_items = []
        for it in (batch_result.get("items") or []):
            if not isinstance(it, dict) or it.get("ok"):
                continue
            nm = str(it.get("name") or it.get("file_id") or "未知文件").strip() or "未知文件"
            err = str(it.get("error") or "失败").strip() or "失败"
            detail = []
            miss = [str(x) for x in (it.get("missing_roles") or []) if str(x).strip()]
            tail = [str(x) for x in (it.get("tail_appended_roles") or []) if str(x).strip()]
            per = it.get("per_role_results") if isinstance(it.get("per_role_results"), dict) else {}
            if miss:
                detail.append("未落位角色：" + "、".join(miss))
            if tail:
                detail.append("文末补签：" + "、".join(tail))
            if per:
                role_lines = []
                for rid, one in per.items():
                    if not isinstance(one, dict) or one.get("placed"):
                        continue
                    why = str(one.get("failure_reason") or "").strip()
                    placed_by = str(one.get("placed_by") or "").strip() or "not_found"
                    chain = one.get("attempt_chain") if isinstance(one.get("attempt_chain"), list) else []
                    chain_txt = " -> ".join(str(x) for x in chain if str(x).strip())[:400]
                    line = f"{rid}（{placed_by}"
                    if why:
                        line += f"，{why}"
                    if chain_txt:
                        line += f"，链路={chain_txt}"
                    line += "）"
                    role_lines.append(line)
                if role_lines:
                    detail.append("识别失败原因：" + "；".join(role_lines[:8]))
            failed_items.append((nm, err, detail))
        if failed_items:
            lines.append("")
            lines.append("=" * 48)
            lines.append("【批签失败文件明细】")
            for nm, err, detail in failed_items:
                lines.append(f"- {nm}：{err}")
                for d in detail:
                    lines.append(f"  · {d}")
    return "\ufeff" + "\n".join(lines) + "\n"


def _signed_outputs_zip_response(
    items,
    download_name: str,
    *,
    batch_id: Optional[str] = None,
    legacy: bool = False,
    batch_result: Optional[dict] = None,
):
    """将已签名记录打成 zip 返回；用内存缓冲，避免临时文件在 send_file 前被删导致下载失败。"""
    from sign_handlers import mysql_store

    buf = io.BytesIO()
    added = 0
    seen_names: set[str] = set()
    pack_lines: list = []
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for it in items or []:
            sid = it.get("id")
            meta_name = it.get("name") or it.get("output_name") or ""
            source_name = it.get("source_name") or ""
            created_at = it.get("created_at")
            roles_label = _format_signed_roles_json(it.get("roles_json") or "")
            pack_entry = {
                "id": sid,
                "name": meta_name,
                "output_name": meta_name,
                "source_name": source_name,
                "created_at": created_at,
                "roles_label": roles_label,
                "packed": False,
                "zip_name": "",
                "error": "",
            }
            if not sid:
                pack_entry["error"] = "记录 id 缺失"
                pack_lines.append(pack_entry)
                continue
            row = mysql_store.get_signed_row(str(sid))
            if not row or not row.get("file_data"):
                pack_entry["error"] = (row or {}).get("file_load_error") or "无法读取文件内容"
                if it.get("ftp_last_error"):
                    pack_entry["error"] += f"（登记错误：{it['ftp_last_error']}）"
                pack_lines.append(pack_entry)
                continue
            nm = (meta_name or row.get("name") or (str(sid) + (it.get("ext") or row.get("ext") or ""))) or "signed"
            nm = os.path.basename(nm)
            base = nm
            k = 2
            while nm in seen_names:
                root, ext = os.path.splitext(base)
                nm = f"{root}({k}){ext}"
                k += 1
            seen_names.add(nm)
            zf.writestr(nm, row["file_data"])
            added += 1
            pack_entry["packed"] = True
            pack_entry["zip_name"] = nm
            pack_lines.append(pack_entry)
        summary = _build_signed_zip_summary_text(
            items,
            pack_lines,
            batch_id=batch_id,
            legacy=legacy,
            batch_result=batch_result,
        )
        zf.writestr(
            "签字结果说明.txt",
            summary,
            compress_type=zipfile.ZIP_DEFLATED,
        )
    if added <= 0 and not items:
        return jsonify({"ok": False, "error": "无匹配文件"}), 404
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=download_name,
        mimetype="application/zip",
    )


@app.route("/api/sign/signed-batch/<batch_id>", methods=["GET"])
def api_sign_signed_batch_items(batch_id):
    """列出某批次内的已签名文件（支持按文件名搜索）。"""
    if not _sign_using_mysql():
        return jsonify({"ok": False, "error": "未启用 MySQL"}), 404
    if not _SIGN_FILE_ID_RE.match(batch_id or ""):
        return jsonify({"ok": False, "error": "无效批次 id"}), 400
    try:
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        q = (request.args.get("q") or "").strip()
        items = mysql_store.list_signed_outputs_by_batch(batch_id=batch_id, q=q)
        return jsonify({"ok": True, "items": items})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/signed-batch/<batch_id>/zip", methods=["GET"])
def api_sign_signed_batch_zip(batch_id):
    """下载某批次的 zip 包。"""
    if not _sign_using_mysql():
        return jsonify({"ok": False, "error": "未启用 MySQL"}), 404
    if not _SIGN_FILE_ID_RE.match(batch_id or ""):
        return jsonify({"ok": False, "error": "无效批次 id"}), 400
    try:
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        items = mysql_store.list_signed_outputs_by_batch(batch_id=batch_id, q="")
        if not items:
            return jsonify({"ok": False, "error": "该批次无文件"}), 404
        dl = f"signed_batch_{batch_id[:8]}.zip"
        batch_result = mysql_store.get_sign_batch_result(batch_id)
        return _signed_outputs_zip_response(
            items, dl, batch_id=batch_id, batch_result=batch_result
        )
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/signed-legacy", methods=["GET"])
def api_sign_signed_legacy():
    """历史已签名记录（无 batch_id）：分页、搜索。"""
    if not _sign_using_mysql():
        return jsonify(
            {
                "ok": True,
                "items": [],
                "total": 0,
                "page": 1,
                "page_size": 50,
                "db_share": False,
            }
        )
    try:
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        q = (request.args.get("q") or "").strip()
        page = max(1, int(request.args.get("page") or 1))
        page_size = max(1, min(200, int(request.args.get("page_size") or 50)))
        items, total = mysql_store.list_signed_legacy_page(q=q, page=page, page_size=page_size)
        return jsonify(
            {
                "ok": True,
                "items": items,
                "total": total,
                "page": page,
                "page_size": page_size,
                "db_share": True,
            }
        )
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/signed-legacy/zip", methods=["GET"])
def api_sign_signed_legacy_zip():
    """下载历史（无 batch_id）记录 zip 包（按搜索过滤）。"""
    if not _sign_using_mysql():
        return jsonify({"ok": False, "error": "未启用 MySQL"}), 404
    try:
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        q = (request.args.get("q") or "").strip()
        # 限制最多打包 300 个，避免 zip 过大
        items, total = mysql_store.list_signed_legacy_page(q=q, page=1, page_size=300)
        if not items:
            return jsonify({"ok": False, "error": "无匹配历史文件"}), 404
        return _signed_outputs_zip_response(items, "signed_legacy.zip", legacy=True)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/stroke-items", methods=["GET"])
def api_sign_stroke_items_list():
    """已入库签字 PNG 素材列表（分页、按签署人搜索）。"""
    if not _sign_using_mysql():
        return jsonify(
            {
                "ok": True,
                "items": [],
                "total": 0,
                "page": 1,
                "page_size": 10,
                "db_share": False,
            }
        )
    try:
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        q = (request.args.get("q") or "").strip()
        cat = (request.args.get("cat") or "").strip()
        temporary = (request.args.get("temporary") or "").strip()
        page = max(1, int(request.args.get("page") or 1))
        page_size = max(1, min(100, int(request.args.get("page_size") or 10)))
        items, total = mysql_store.list_stroke_items_page(
            q=q, page=page, page_size=page_size, cat=cat, temporary=temporary
        )
        return jsonify(
            {
                "ok": True,
                "items": items,
                "total": total,
                "page": page,
                "page_size": page_size,
                "db_share": True,
            }
        )
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/stroke-items/purge-temporary", methods=["POST"])
def api_sign_stroke_items_purge_temporary():
    """按当前筛选条件批量删除临时素材（仅 is_temporary=1）。"""
    if not _sign_using_mysql():
        return jsonify({"ok": False, "error": "未启用 MySQL"}), 400
    data = request.get_json(silent=True) or {}
    q = (data.get("q") or request.args.get("q") or "").strip()
    cat = (data.get("cat") or request.args.get("cat") or "").strip()
    try:
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        n = mysql_store.delete_temporary_stroke_items(q=q, cat=cat)
        return jsonify({"ok": True, "deleted": n})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/sign/stroke-items/<item_id>", methods=["DELETE"])
def api_sign_stroke_item_delete(item_id):
    if not _sign_using_mysql():
        return jsonify({"ok": False, "error": "未启用 MySQL"}), 400
    if not _SIGN_FILE_ID_RE.match(item_id or ""):
        return jsonify({"ok": False, "error": "无效 id"}), 400
    try:
        from sign_handlers import mysql_store

        mysql_store.ensure_sign_mysql()
        n = mysql_store.delete_stroke_item(item_id)
        if not n:
            return jsonify({"ok": False, "error": "记录不存在"}), 404
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


if __name__ == "__main__":
    _dp = _resolved_dotenv_path()
    print(
        "[aiprintword] build=%s app.py=%s"
        % (AIPRINTWORD_WEB_BUILD, os.path.abspath(__file__))
    )
    print(
        "[aiprintword] .env=%s exists=%s token_loaded=%s"
        % (
            _dp,
            os.path.isfile(_dp),
            bool((os.environ.get("AIPRINTWORD_ADMIN_TOKEN") or "").strip()),
        )
    )
    print(
        "[aiprintword] 若 token_loaded=False 或浏览器仍见极短 503：关闭其它占用 5050 的窗口后只保留本进程，并确认 .env 在同目录"
    )
    # 启动自检：端口被占用时，直接提示并退出，避免“看似启动成功但页面打不开/超时”
    _port = 5050
    try:
        _sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        _sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
        _sock.bind(("127.0.0.1", _port))
        _sock.close()
    except OSError:
        print("")
        print("[aiprintword] 启动自检失败：端口 %s 已被占用。" % _port)
        print("[aiprintword] 已有其它 app.py 在跑，先关掉它，再重新启动本服务。")
        if os.name == "nt":
            print("[aiprintword] Windows 排查/关闭示例：")
            print("  netstat -ano | findstr \":%s\"" % _port)
            print("  tasklist /FI \"PID eq <PID>\" /FO LIST")
            print("  taskkill /PID <PID> /F")
        else:
            print("[aiprintword] Linux/macOS 排查/关闭示例：")
            print("  lsof -nP -iTCP:%s -sTCP:LISTEN" % _port)
            print("  kill -9 <PID>")
        raise SystemExit(2)
    # threaded=True：批量上传/保存笔迹时，其它页面仍可拉取签署人列表，避免整站「请求超时」
    app.run(host="0.0.0.0", port=5050, debug=False, threaded=True)
